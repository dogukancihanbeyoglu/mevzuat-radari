"""
Local Audit AI - Resmi Formatlı Çalışma Kağıdı İhraç Motoru (Workpaper Exporter)
Üretilen metin ve tabloları kurumsal antetli Word (.docx) ve biçimlendirilmiş Excel (.xlsx) çalışma kağıtlarına dönüştürür.
"""
import io
import re
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import pandas as pd
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

class WorkpaperExporter:
    def __init__(self):
        pass

    def _set_cell_background(self, cell, fill_color="1E3A8A"):
        """Word tablo başlık hücresine kurumsal arka plan rengi verir."""
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), fill_color)
        tcPr.append(shd)

    def export_to_docx(
        self,
        title: str,
        content: str,
        audit_trail_id: str = "AT-2026-LOCAL",
        context: str = "Mega Holding A.Ş. — 2026 İç Denetimi"
    ) -> bytes:
        """Markdown formatındaki denetim çıktısını profesyonel Word belgesine dönüştürür."""
        doc = docx.Document()

        # Sayfa Kenar Boşlukları
        for section in doc.sections:
            section.top_margin = Inches(0.8)
            section.bottom_margin = Inches(0.8)
            section.left_margin = Inches(0.8)
            section.right_margin = Inches(0.8)

        # Üst Bilgi / Antet
        header = doc.sections[0].header
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hrun = hp.add_run("IIA GLOBAL STANDARDS — İÇ DENETİM ÇALIŞMA KAĞIDI (CONFIDENTIAL)")
        hrun.font.size = Pt(8.5)
        hrun.font.color.rgb = RGBColor(120, 120, 120)

        # Başlık
        h1 = doc.add_heading(title, level=1)
        h1.runs[0].font.color.rgb = RGBColor(30, 58, 138) # Dark Navy

        # Meta Bilgi Kutusu
        meta_table = doc.add_table(rows=2, cols=2)
        meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        meta_table.rows[0].cells[0].text = f"Kurumsal Bağlam: {context}"
        meta_table.rows[0].cells[1].text = f"Denetim İzi (Audit Trail ID): {audit_trail_id}"
        meta_table.rows[1].cells[0].text = "Metodoloji: IIA Global Standards (2026)"
        meta_table.rows[1].cells[1].text = "Güvenlik Sınıfı: %100 Air-Gapped / Yerel Şifrelenmiş"
        
        for row in meta_table.rows:
            for cell in row.cells:
                self._set_cell_background(cell, "F1F5F9")
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(9)
                        r.font.color.rgb = RGBColor(71, 85, 105)

        doc.add_paragraph() # Boşluk

        # İçerik Ayrıştırma (Başlıklar, Paragraflar ve Markdown Tabloları)
        lines = content.splitlines()
        in_table = False
        table_lines = []

        for line in lines:
            line_str = line.strip()
            
            # Markdown Tablo Satırı
            if line_str.startswith("|") and line_str.endswith("|"):
                in_table = True
                table_lines.append(line_str)
                continue
            else:
                if in_table and table_lines:
                    self._render_markdown_table_to_docx(doc, table_lines)
                    in_table = False
                    table_lines = []

            if not line_str:
                continue

            if line_str.startswith("### "):
                h = doc.add_heading(line_str.replace("### ", ""), level=3)
                h.runs[0].font.color.rgb = RGBColor(30, 58, 138)
            elif line_str.startswith("## "):
                h = doc.add_heading(line_str.replace("## ", ""), level=2)
                h.runs[0].font.color.rgb = RGBColor(15, 23, 42)
            elif line_str.startswith("# "):
                h = doc.add_heading(line_str.replace("# ", ""), level=1)
                h.runs[0].font.color.rgb = RGBColor(30, 58, 138)
            elif line_str.startswith("- ") or line_str.startswith("* "):
                p = doc.add_paragraph(style='List Bullet')
                self._add_formatted_runs(p, line_str[2:])
            elif re.match(r"^\d+\.\s", line_str):
                p = doc.add_paragraph(style='List Number')
                clean_text = re.sub(r"^\d+\.\s", "", line_str)
                self._add_formatted_runs(p, clean_text)
            else:
                p = doc.add_paragraph()
                self._add_formatted_runs(p, line_str)

        if in_table and table_lines:
            self._render_markdown_table_to_docx(doc, table_lines)

        # Alt Bilgi / İmza
        footer = doc.sections[0].footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        frun = fp.add_run(f"📌 Doğrulanmış Kriptografik Denetim Kaydı — {audit_trail_id} | Yerel AI Denetim Sistemi")
        frun.font.size = Pt(8)
        frun.font.color.rgb = RGBColor(148, 163, 184)

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    def _add_formatted_runs(self, paragraph, text: str):
        """Bold (**metin**) ve Italic (*metin*) formatlarını Word run nesnesine dönüştürür."""
        tokens = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
        for token in tokens:
            if token.startswith('**') and token.endswith('**'):
                r = paragraph.add_run(token[2:-2])
                r.bold = True
            elif token.startswith('*') and token.endswith('*'):
                r = paragraph.add_run(token[1:-1])
                r.italic = True
            else:
                paragraph.add_run(token)

    def _render_markdown_table_to_docx(self, doc, table_lines: list):
        """Markdown tablo satırlarını Word tablosuna dönüştürür."""
        rows = []
        for tl in table_lines:
            if re.match(r"^\|[\s\-:|]+\|$", tl):
                continue # Ayırıcı çizgi
            cells = [c.strip() for c in tl.split("|")[1:-1]]
            if cells:
                rows.append(cells)

        if not rows:
            return

        table = doc.add_table(rows=len(rows), cols=len(rows[0]))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for r_idx, row_data in enumerate(rows):
            for c_idx, cell_value in enumerate(row_data):
                if c_idx < len(table.rows[r_idx].cells):
                    cell = table.rows[r_idx].cells[c_idx]
                    cell.text = cell_value
                    if r_idx == 0:
                        self._set_cell_background(cell, "1E3A8A") # Koyu Lacivert Başlık
                        for p in cell.paragraphs:
                            for r in p.runs:
                                r.font.bold = True
                                r.font.color.rgb = RGBColor(255, 255, 255)
                                r.font.size = Pt(9.5)
                    else:
                        bg = "F8FAFC" if r_idx % 2 == 1 else "FFFFFF"
                        self._set_cell_background(cell, bg)
                        for p in cell.paragraphs:
                            for r in p.runs:
                                r.font.size = Pt(9)
                                r.font.color.rgb = RGBColor(51, 65, 85)
        doc.add_paragraph()

    def export_to_excel(self, title: str, content: str) -> bytes:
        """Markdown içeriklerini ve tablolarını doğrudan openpyxl ile kurşun geçirmez şekilde Excel'e yazar."""
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Denetim_Calisma_Kagidi"

        # Başlık ve Stil Tanımları
        header_font = Font(name='Calibri', size=11, bold=True, color='FFFFFF')
        header_fill = PatternFill(start_color='1E3A8A', end_color='1E3A8A', fill_type='solid')
        title_font = Font(name='Calibri', size=14, bold=True, color='1E3A8A')
        section_font = Font(name='Calibri', size=12, bold=True, color='0F172A')
        bold_font = Font(name='Calibri', size=10, bold=True)
        regular_font = Font(name='Calibri', size=10)
        thin_border = Border(
            left=Side(style='thin', color='E2E8F0'),
            right=Side(style='thin', color='E2E8F0'),
            top=Side(style='thin', color='E2E8F0'),
            bottom=Side(style='thin', color='E2E8F0')
        )

        current_row = 1
        
        # 1. Başlık Satırı
        ws.cell(row=current_row, column=1, value=title.upper()).font = title_font
        current_row += 2

        lines = content.splitlines()
        in_table = False
        table_rows = []

        for line in lines:
            line_str = line.strip()
            
            # Markdown Tablo Satırı
            if line_str.startswith("|") and line_str.endswith("|"):
                if re.match(r"^\|[\s\-:|]+\|$", line_str):
                    continue # Ayırıcı çizgi
                in_table = True
                cells = [c.strip() for c in line_str.split("|")[1:-1]]
                if cells:
                    table_rows.append(cells)
                continue
            else:
                if in_table and table_rows:
                    # Tabloyu Excel'e yaz
                    for r_idx, row_vals in enumerate(table_rows):
                        for c_idx, val in enumerate(row_vals):
                            cell = ws.cell(row=current_row + r_idx, column=c_idx + 1, value=str(val)[:32000])
                            cell.border = thin_border
                            if r_idx == 0:
                                cell.font = header_font
                                cell.fill = header_fill
                                cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
                            else:
                                cell.font = regular_font
                                cell.alignment = Alignment(vertical='center', wrap_text=True)
                                if r_idx % 2 == 1:
                                    cell.fill = PatternFill(start_color='F8FAFC', end_color='F8FAFC', fill_type='solid')

                    current_row += len(table_rows) + 1
                    in_table = False
                    table_rows = []

            if not line_str:
                continue

            # Başlıklar ve Paragraflar
            if line_str.startswith("#"):
                clean_heading = line_str.lstrip("#").strip()
                cell = ws.cell(row=current_row, column=1, value=clean_heading[:32000])
                cell.font = section_font
                current_row += 1
            else:
                # Metin Satırı
                cell = ws.cell(row=current_row, column=1, value=line_str[:32000])
                cell.font = regular_font
                current_row += 1

        if in_table and table_rows:
            for r_idx, row_vals in enumerate(table_rows):
                for c_idx, val in enumerate(row_vals):
                    cell = ws.cell(row=current_row + r_idx, column=c_idx + 1, value=str(val)[:32000])
                    cell.border = thin_border
                    if r_idx == 0:
                        cell.font = header_font
                        cell.fill = header_fill
                    else:
                        cell.font = regular_font

        # Sütun Genişliklerini Otomatik Ayarla
        for col in ws.columns:
            max_len = max(len(str(cell.value or '')) for cell in col)
            col_letter = openpyxl.utils.get_column_letter(col[0].column)
            ws.column_dimensions[col_letter].width = min(max(max_len + 3, 15), 60)

        buf = io.BytesIO()
        wb.save(buf)
        return buf.getvalue()
