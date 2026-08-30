"""
Local Audit AI - Ingestion Module
Yerel dosya ayrıştırıcıları (Word, Excel, PDF, TXT)
"""
import os
from typing import Dict, Any, List

def parse_docx(file_path: str) -> str:
    """Word (.docx) dosyasından metin ve tabloları çıkarır."""
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"Dosya bulunamadı: {file_path}")
    
    try:
        import docx
        doc = docx.Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text.strip())
        
        # Tabloları da metin olarak ekle
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                table_data.append(" | ".join(row_text))
            if table_data:
                full_text.append("\n[TABLO]:\n" + "\n".join(table_data))
                
        return "\n\n".join(full_text)
    except Exception as e:
        # Fallback to zip parser if python-docx fails
        import zipfile
        import xml.etree.ElementTree as ET
        with zipfile.ZipFile(file_path) as docx_zip:
            tree = ET.fromstring(docx_zip.read('word/document.xml'))
            namespaces = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
            paragraphs = []
            for p in tree.iterfind('.//w:p', namespaces):
                p_text = ''.join(node.text for node in p.iterfind('.//w:t', namespaces) if node.text)
                if p_text.strip():
                    paragraphs.append(p_text.strip())
            return "\n\n".join(paragraphs)

def parse_excel(file_path: str, max_rows: int = 50) -> str:
    """Excel (.xlsx, .xls) dosyasından sekmeleri ve örnek verileri çıkarır."""
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"Dosya bulunamadı: {file_path}")
    
    import pandas as pd
    excel_file = pd.ExcelFile(file_path)
    sheet_summaries = []
    
    for sheet_name in excel_file.sheet_names:
        df = pd.read_excel(file_path, sheet_name=sheet_name, nrows=max_rows)
        columns_str = ", ".join([str(c) for c in df.columns])
        sample_str = df.head(10).to_string(index=False)
        sheet_summaries.append(
            f"--- Sekme: {sheet_name} (Toplam {len(df)} satır örnek) ---\n"
            f"Sütunlar: {columns_str}\n"
            f"Örnek Veri:\n{sample_str}"
        )
    return "\n\n".join(sheet_summaries)

def parse_pdf(file_path: str) -> str:
    """PDF dosyasından metin çıkarır."""
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"Dosya bulunamadı: {file_path}")
    
    try:
        import pypdf
        reader = pypdf.PdfReader(file_path)
        pages_text = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text()
            if text and text.strip():
                pages_text.append(f"[Sayfa {i+1}]\n{text.strip()}")
        return "\n\n".join(pages_text)
    except Exception as e:
        return f"PDF ayrıştırma hatası: {str(e)}"

def parse_document(file_path: str) -> str:
    """Uzantıya göre uygun ayrıştırıcıyı çağırır."""
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".docx":
        return parse_docx(file_path)
    elif ext in [".xlsx", ".xls"]:
        return parse_excel(file_path)
    elif ext == ".pdf":
        return parse_pdf(file_path)
    elif ext in [".txt", ".md", ".csv", ".json"]:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    else:
        raise ValueError(f"Desteklenmeyen dosya formatı: {ext}")

def parse_multiple_documents(file_paths: List[str]) -> str:
    """Birden fazla dosyayı ayrıştırıp başlıklarıyla birleştirir."""
    combined_texts = []
    for fpath in file_paths:
        if os.path.exists(fpath):
            fname = os.path.basename(fpath)
            content = parse_document(fpath)
            combined_texts.append(f"=== [BELGE: {fname}] ===\n{content}\n")
    return "\n\n".join(combined_texts)
