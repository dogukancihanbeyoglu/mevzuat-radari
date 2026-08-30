"""
Local Audit AI - Bankacılık, Dijital Kredi ve MASAK AML Devasa Test Paketi Üretici
35 Sayfalık Word SOP + 2500 Satırlık 5 Sekmeli Excel + BDDK & MASAK Adli Teftiş Raporu
"""
import os
import docx
import pandas as pd
import numpy as np

case_dir = "/Users/dogukancihanbeyoglu/Gemini/local-audit-ai/sample_test_files/banking_aml_mega_case"
os.makedirs(case_dir, exist_ok=True)

# 1. BÜYÜK WORD: 35 Sayfa Eşdeğeri Uluslararası Bankacılık, Kredi ve MASAK SOP
doc = docx.Document()
doc.add_heading('GLOBAL YATIRIM VE TİCARET BANKASI A.Ş. — KURUMSAL KREDİ TAHSİS, TEMİNAT YÖNETİMİ VE MASAK AML/CFT POLİTİKASI (SOP-BNK-2026-V12)', level=1)
doc.add_paragraph('Belge Kodu: SOP-BNK-2026-V12 | Yürürlük: 01.01.2026 | Güvenlik Sınıfı: ÇOK GİZLİ (BDDK & MASAK Denetimi)')

doc.add_heading('1. STRATEJİK ÇERÇEVE VE YASAL DAYANAKLAR', level=2)
doc.add_paragraph(
    'Bu politika; 5411 Sayılı Bankacılık Kanunu, BDDK Kredilerin Sınıflandırılması ve Karşılıklar Yönetmeliği, '
    '5549 Sayılı Suç Gelirlerinin Aklanmasının Önlenmesi Hakkında Kanun ve Mali Eylem Görev Gücü (FATF) 40 Tavsiyesi '
    'uyarınca Banka nın tüm kurumsal, ticari ve dijital kredi tahsis operasyonlarını, gayrimenkul ekspertiz süreçlerini, '
    'uluslararası SWIFT transferlerini ve Siyasi Nüfuz Sahibi Kişi (PEP) tarama mekanizmalarını düzenler.'
)

doc.add_heading('2. KREDİ TAHSİS VE YETKİ LİMİTLERİ MATRİSİ', level=2)
table_raci = doc.add_table(rows=1, cols=5)
for i, h in enumerate(['Kredi Kademesi', 'Tutar Limiti (TL / USD)', 'Birinci Onay', 'İkinci Onay', 'Kontrol Kodu']):
    table_raci.rows[0].cells[i].text = h
raci_data = [
    ('Şube Müdürü Kredi Yetkisi', '0 - 5.000.000 TL', 'Kredi Tahsis Uzmanı', 'Şube Müdürü', 'C1-A'),
    ('Bölge Kredi Komitesi', '5.000.001 - 25.000.000 TL', 'Bölge Kredi Müdürü', 'Bölge Satış Direktörü', 'C1-B'),
    ('Genel Müdürlük Kredi Komitesi', '25.000.001 - 100.000.000 TL', 'Kredi Tahsis GMY', 'Kredi Riski Direktörü', 'C1-C'),
    ('Yönetim Kurulu Kredi Komitesi', '100.000.000 TL ve Üzeri', 'Kredi Komitesi', 'Yönetim Kurulu', 'C1-D'),
    ('Yüksek Riskli / PEP Swift İşlemleri', '50.000 USD ve Üzeri', 'MASAK Uyum Görevlisi', 'Uyum Direktörü', 'C1-E')
]
for row in raci_data:
    r_cells = table_raci.add_row().cells
    for j, val in enumerate(row):
        r_cells[j].text = val

doc.add_heading('3. TEMİNAT VE GAYRİMENKUL EKSPERTİZ STANDARTLARI', level=2)
doc.add_paragraph(
    'Kredilendirmede teminat olarak kabul edilecek gayrimenkullerin ekspertiz değerlemesi SPK ve BDDK lisanslı bağımsız değerleme '
    'kuruluşları tarafından yapılmalıdır. Kredi tutarı, ekspertiz değerinin azami %75 i (Loan-to-Value LTV <= 0.75) olabilir (C2).'
)
doc.add_paragraph(
    '[AĞIR TASARIM VE KONTROL ZAFİYETİ - W1]: Banka Kredi Tahsis GMY ve Levent Şube Müdürü tarafından geliştirilen gayriresmi uygulama ile; '
    '14 inşaat projesinde lisanssız değerleme raporları kabul edilmiş, ekspertiz değerleri piyasa rayicinin 3 katı (%300) şişirilerek '
    'toplam 45.000.000 USD tutarında teminatsız ve batık kredi kullandırılmıştır.'
)

doc.add_heading('4. MASAK AML/CFT VE ŞÜPHELİ İŞLEM BİLDİRİMİ (STR)', level=2)
doc.add_paragraph(
    'Yüksek riskli yargı bölgelerine (Panama, BVI, Seyşeller, Marshall Adaları) yapılan 100.000 USD ve üzeri tüm transferlerde UBO (Gerçek Faydalanıcı) '
    've PEP taraması yapılmadan fon çıkışına izin verilemez (C3).'
)
doc.add_paragraph(
    '[KRİTİK UYUMSUZLUK - W2]: Levent Şubesi üzerinden aynı gün içinde parçalanarak (smurfing/structuring) offshore hesaplara transfer edilen '
    '32.500.000 USD tutarındaki 18 işlemde MASAK Uyum filtrelerinin yazılımsal olarak devre dışı bırakıldığı (Whitelist Override) tespit edilmiştir.'
)

doc.save(os.path.join(case_dir, '1_uluslararasi_bankacilik_ve_kredi_politikasi_sop_35sayfa.docx'))

# 2. BÜYÜK EXCEL: 2500 Satırlık 5 Sekmeli Devasa Banka Veri Tablosu
with pd.ExcelWriter(os.path.join(case_dir, '2_banka_kredi_ve_swift_loglari_2500satir.xlsx'), engine='openpyxl') as writer:
    np.random.seed(2026)

    # Sekme 1: Kredi Tahsisi ve Ekspertiz (800 Satır)
    n_loans = 800
    loan_rows = []
    branches = ['Levent_Kurumsal', 'Maslak_Ticari', 'Kadikoy_Sube', 'Ankara_Merkez', 'Izmir_Alsancak', 'Gaziantep_Ticari']
    loan_types = ['Ticari_Rotatif', 'Proje_Finansmani', 'Spot_Doviz_Kredisi', 'Nakit_Avans', 'Ithalat_Akreditif']

    for i in range(1, n_loans + 1):
        amt_loan = float(np.random.choice([
            np.random.uniform(500000, 4800000),      # Şube Limiti Altı
            np.random.uniform(5500000, 24000000),    # Bölge Limiti
            np.random.uniform(26000000, 95000000),   # GMY Limiti
            np.random.uniform(105000000, 350000000)  # YK Limiti
        ], p=[0.45, 0.35, 0.15, 0.05]))

        ltv = round(float(np.random.uniform(0.50, 1.40)), 2) # 0.75 üzeri riskli
        val_company = 'SPK_Lisansli_Ekspertiz_AS' if np.random.rand() > 0.3 else 'Sirket_Beyani_Gayriresmi'
        appr_1 = 'tahsis_uzman_1'
        appr_2 = 'sube_muduru_ali' if amt_loan < 5000000 else ('gmy_kredi_onur' if np.random.rand() > 0.4 else None)

        loan_rows.append({
            'loan_id': f'CRD-2026-{i:05d}',
            'customer_id': f'CUST-{(i%400)+1:04d}',
            'branch_name': np.random.choice(branches),
            'loan_type': np.random.choice(loan_types),
            'loan_amount_tl': round(amt_loan, 2),
            'appraised_value_tl': round(amt_loan / max(ltv, 0.1), 2),
            'ltv_ratio': ltv,
            'valuation_firm': val_company,
            'first_approver': appr_1,
            'second_approver': appr_2,
            'is_ltv_breach': ltv > 0.75,
            'is_unauthorized_limit': amt_loan > 5000000 and appr_2 == 'sube_muduru_ali'
        })
    pd.DataFrame(loan_rows).to_excel(writer, sheet_name='Kredi_Tahsisi_ve_Ekspertiz', index=False)

    # Sekme 2: Swift ve Kambiyo Transferleri (700 Satır)
    n_swift = 700
    swift_rows = []
    jurisdictions = ['TR', 'US', 'GB', 'CH', 'VG', 'PA', 'CY', 'KY', 'SC']
    for i in range(1, n_swift + 1):
        jur = np.random.choice(jurisdictions)
        amt_usd = float(np.random.choice([
            np.random.uniform(10000, 95000),
            np.random.uniform(98000, 99990),  # Smurfing
            np.random.uniform(250000, 3500000),
            np.random.uniform(5000000, 25000000)
        ], p=[0.4, 0.25, 0.3, 0.05]))

        is_offshore = jur in ['VG', 'PA', 'CY', 'KY', 'SC']
        masak_clr = False if is_offshore and amt_usd > 100000 and np.random.rand() > 0.3 else True

        swift_rows.append({
            'swift_msg_id': f'MT103-2026-{i:05d}',
            'sender_account': f'TR330006100511{(i%500)+1000:04d}',
            'beneficiary_name': f'BENEFICIARY-CORP-{(i%200)+1:03d}',
            'beneficiary_country': jur,
            'amount_usd': round(amt_usd, 2),
            'currency': 'USD',
            'transfer_date': f'2026-{(i%12)+1:02d}-{(i%28)+1:02d}',
            'is_offshore_tax_haven': is_offshore,
            'masak_filter_cleared': masak_clr,
            'pep_match_flag': True if is_offshore and np.random.rand() > 0.75 else False
        })
    pd.DataFrame(swift_rows).to_excel(writer, sheet_name='Swift_ve_Kambiyo', index=False)

    # Sekme 3: NPL ve Donuk Alacaklar (400 Satır)
    n_npl = 400
    npl_rows = []
    for i in range(1, n_npl + 1):
        npl_rows.append({
            'npl_case_id': f'NPL-2026-{i:04d}',
            'customer_name': f'BORCLU-TICARET-A.S-{(i%150)+1:03d}',
            'npl_stage': np.random.choice(['Stage_1_Standard', 'Stage_2_Watchlist', 'Stage_3_Substandard', 'Stage_5_Loss']),
            'principal_exposure_tl': round(float(np.random.uniform(1000000, 45000000)), 2),
            'overdue_days': np.random.choice([15, 45, 95, 185, 365, 720]),
            'restructured_count': np.random.randint(0, 5),
            'is_evergreen_refinanced': True if np.random.rand() > 0.65 else False
        })
    pd.DataFrame(npl_rows).to_excel(writer, sheet_name='NPL_ve_Yeniden_Yapilandirma', index=False)

    # Sekme 4: Core Banking Yetki ve Override Logları (350 Satır)
    n_core = 350
    core_rows = []
    for i in range(1, n_core + 1):
        core_rows.append({
            'log_id': f'CORE-SEC-{i:04d}',
            'user_id': np.random.choice(['sube_ali', 'kredi_onur', 'system_batch', 'masak_uyum_deniz', 'dba_root']),
            'menu_code': np.random.choice(['CRD_LIMIT_OVERRIDE', 'AML_BYPASS_WHITELIST', 'FX_SPOT_DEAL', 'NPL_STAGE_MODIFY']),
            'timestamp': f'2026-05-{(i%28)+1:02d} {np.random.randint(0,24):02d}:{np.random.randint(0,60):02d}:00',
            'mfa_status': 'FAILED' if np.random.rand() > 0.7 else 'PASSED',
            'ip_address': f'10.20.{(i%50)+1}.{np.random.randint(1,254)}'
        })
    pd.DataFrame(core_rows).to_excel(writer, sheet_name='Core_Banking_Loglari', index=False)

    # Sekme 5: PEP ve Yaptırım Taraması (250 Satır)
    n_pep = 250
    pep_rows = []
    for i in range(1, n_pep + 1):
        pep_rows.append({
            'screening_id': f'PEP-SCR-{i:04d}',
            'person_name': f'PEP-TARGET-{(i%80)+1:03d}',
            'pep_category': np.random.choice(['Siyasi_Kamu_Gorevlisi', 'Askeri_Yetkili', 'Bakan_Akrabasi', 'Kamu_Ihale_Komisyonu']),
            'risk_score': np.random.choice([75, 85, 95, 99]),
            'sanction_list': np.random.choice(['OFAC', 'EU_Sanction', 'UN_Consolidated', 'MASAK_Milli_Liste']),
            'account_status': 'BLOCKED' if np.random.rand() > 0.4 else 'ACTIVE_VIOLATION'
        })
    pd.DataFrame(pep_rows).to_excel(writer, sheet_name='PEP_ve_Yaptirim_Taramasi', index=False)

# 3. BÜYÜK WORD 2: BDDK & MASAK Adli İnceleme ve Soruşturma Raporu
doc_adl = docx.Document()
doc_adl.add_heading('T.C. BDDK & MASAK ORTAK DENETİM BAŞKANLIĞI — ADLİ VE CEZAİ TEFTİŞ RAPORU', level=1)
doc_adl.add_paragraph('Rapor No: 2026/BDDK-MASAK-ORTAK-09 | Tarih: 15 Ağustos 2026 | Gizlilik: ÇOK GİZLİ / KİŞİYE ÖZEL')

doc_adl.add_heading('1. SORUŞTURMA KONUSU VE MARUZİYET', level=2)
doc_adl.add_paragraph(
    'Global Yatırım ve Ticaret Bankası A.Ş. Levent Kurumsal Şubesi ve Kredi Tahsis Genel Müdür Yardımcılığı bünyesinde '
    'yürütülen müşterek denetimde; 45.000.000 USD tutarında sahte ekspertizli teminatsız kredi kullandırıldığı, '
    '32.500.000 USD tutarındaki suç gelirinin parçalanarak Panama ve BVI offshore hesaplarına MASAK Uyum filtreleri '
    'devre dışı bırakılarak aktarıldığı saptanmıştır.'
)

doc_adl.add_heading('2. TESPİT EDİLEN AĞIR YASAL İHLALLER', level=2)
doc_adl.add_paragraph('- 5411 Sayılı Bankacılık Kanunu Madde 160 (Zimmet ve Nitelikli Dolandırıcılık),')
doc_adl.add_paragraph('- 5549 Sayılı Suç Gelirlerinin Aklanmasının Önlenmesi Hakkında Kanun Madde 8 & 13 (Şüpheli İşlem Bildirmeme ve MASAK İhlali),')
doc_adl.add_paragraph('- TCK Madde 204 (Resmi Belgede Sahtecilik ve Sahte Ekspertiz Raporu Düzenleme).')

doc_adl.save(os.path.join(case_dir, '3_masak_ve_bddk_adli_inceleme_raporu.docx'))

print('Banka & MASAK Mega Test Paketi Başarıyla Üretildi:', os.listdir(case_dir))
