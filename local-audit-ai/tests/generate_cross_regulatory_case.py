"""
Local Audit AI - Çoklu Regülasyon İhlali (Cross-Regulatory Fraud & Compliance) Test Vakası
KVKK Veri Sızıntısı + SPK Insider Trading + BDDK Kredi Zimmeti + MASAK Kara Para Aklama
"""
import os
import docx

case_dir = "/Users/dogukancihanbeyoglu/Gemini/local-audit-ai/sample_test_files/cross_regulatory_mega_case"
os.makedirs(case_dir, exist_ok=True)

doc = docx.Document()
doc.add_heading('ANADOLU FİNANS VE TEKNOLOJİ HOLDİNG A.Ş. — ÇAPRAZ REGÜLASYON VE ADLİ TEFTİŞ RAPORU', level=1)
doc.add_paragraph('Rapor Kodu: AUD-MULTI-REG-2026-09 | Gizlilik: ÇOK GİZLİ (BDDK, MASAK, SPK, KVKK ve Savcılık Müşterek Dosyası)')

doc.add_heading('1. VAKA ÖZETİ VE ÇAPRAZ HUKUKİ MARUZİYET', level=2)
doc.add_paragraph(
    'Holding in bankacılık ve fintech iştirakinde yürütülen iç soruşturmada; '
    'Bilgi Teknolojileri Direktörü ve Hazine Müdürü nün organize işbirliği ile 4 farklı yasal otoritenin '
    'mevzuatını ihlal eden ağır suistimaller zinciri ortaya çıkarılmıştır.'
)

doc.add_heading('2. TESPİT EDİLEN 4 BÜYÜK YASAL İHLAL VE MEVZUAT EŞLEŞMELERİ', level=2)
doc.add_paragraph(
    '1. KVKK & GDPR İhlali (Müşteri Veri Tabanı Sızıntısı): '
    'Bankanın 250.000 adet VIP müşterisine ait TCKN, kredi kartı numaraları ve hesap bakiyeleri, '
    'BT Direktörü tarafından CyberArk PAM sistemi bypass edilerek harici USB belleğe aktarılmış ve Dark Web de satışa sunulmuştur (KVKK Madde 12 ihlali).'
)
doc.add_paragraph(
    '2. SPK & Sermaye Piyasası İhlali (İçeriden Öğrenenlerin Ticareti - Insider Trading): '
    'Holding in halka açık iştirakinin bedelli sermaye artırımı ve dev satın alma haberi kamuya açıklanmadan 5 gün önce, '
    'Hazine Müdürü tarafından 1. derece akrabalarının Borsa İstanbul hesapları üzerinden 450.000 lot hisse alımı yapılmış ve 18.200.000 TL haksız kazanç sağlanmıştır (SPK Madde 106 ihlali).'
)
doc.add_paragraph(
    '3. BDDK & TCK İhlali (Sahte Ekspertiz ve Zimmet Kredisi): '
    'Levent Kurumsal Şube Müdürü ile işbirliği yapılarak, piyasa değeri 10 Milyon TL olan arsa için sahte SPK lisanssız ekspertiz raporuyla 85 Milyon TL (%850 LTV) kredi kullandırılmış ve para paravan şirkete aktarılmıştır (5411 Sayılı Kanun Madde 160 ve TCK 158/204 zimmet ve dolandırıcılık).'
)
doc.add_paragraph(
    '4. MASAK İhlali (Kara Para Aklama ve Parçalama / Smurfing): '
    'Elde edilen suç gelirleri, MASAK Uyum filtreleri sistemden devre dışı bırakılarak (Whitelist Override) '
    'aynı gün içinde 50.000 USD lik parçalar halinde Panama ve BVI offshore şirket hesaplarına transfer edilmiştir (5549 Sayılı Kanun Madde 8 & 13 ihlali).'
)

doc.save(os.path.join(case_dir, 'cross_regulatory_fraud_and_compliance_case.docx'))
print('Çoklu Regülasyon Test Dosyası Başarıyla Üretildi:', os.listdir(case_dir))
