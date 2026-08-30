"""
Local Audit AI - Devasa Endüstriyel Boyutlu Test Dosyaları Üretici
100 Sayfalık Word Dokümanı + 25.000 Satırlık 6 Sekmeli Excel + Kapsamlı Adli İnceleme Raporu
"""
import os
import docx
from docx.shared import Inches, Pt, RGBColor
import pandas as pd
import numpy as np

case_dir = "/Users/dogukancihanbeyoglu/Gemini/local-audit-ai/sample_test_files/industrial_scale_cases"
os.makedirs(case_dir, exist_ok=True)

print("⏳ 1. 100 Sayfalık Devasa Word SOP Dokümanı Üretiliyor...")

# 1. 100 SAYFALIK DEVASA WORD BELGESİ
doc = docx.Document()
doc.add_heading('GLOBAL KONSORSIYUM BANKACILIK VE FINANS GRUBU — KURUMSAL KREDI, HAZINE, ITGC VE MASAK AML/CFT ANA POLITIKA VE PROSEDUR METNI (SOP-CORP-2026-ULTRA)', level=1)
doc.add_paragraph('Belge Kodu: SOP-CORP-2026-ULTRA-V20 | Yürürlük: 01.01.2026 | Sayfa Sayısı: 100+ | Güvenlik: ÇOK GİZLİ')

sections_data = [
    ("GENEL HÜKÜMLER VE YÖNETİŞİM ÇERÇEVESİ", "Bu bölüm bankanın üst düzey risk iştahını, yönetim kurulu komitelerinin yetki ve sorumluluklarını düzenler."),
    ("KURUMSAL VE TİCARİ KREDİ TAHSİS METODOLOJİSİ", "Kredi skorlama modelleri, mali tahlil kriterleri, sektör konsantrasyon limitleri ve grup risk sınırları."),
    ("GAYRİMENKUL VE MENKUL TEMİNAT DEĞERLEME ESASLARI", "SPK ve BDDK lisanslı değerleme zorunlulukları, teminat iskonto oranları ve LTV sınırları."),
    ("DİJİTAL KREDİLER VE YAPAY ZEKA ALGORİTMA DENETİMİ", "Algoritmik kredi onaylama modelleri, önyargı testleri ve model doğrulama süreçleri."),
    ("HAZİNE OPERASYONLARI VE TÜREV SWAP İŞLEMLERİ", "Döviz spot, vadeli, opsiyon, faiz ve emtia swapları için 4-göz ilkesi ve stop-loss limitleri."),
    ("ULUSLARARASI SWIFT VE FON TRANSFER GÜVENLİĞİ", "MT103, MT202 mesaj kontrolleri, çift imza ve Sanction Screening filtreleri."),
    ("MASAK AML / CFT VE ŞÜPHELİ İŞLEM TESPİT POLİTİKASI", "5549 Sayılı Kanun, FATF kırmızı bayrakları, UBO doğrulaması ve STR bildirim süreçleri."),
    ("SİYASİ NÜFUZ SAHİBİ KİŞİLER (PEP) VE YAPTIRIM YÖNETİMİ", "OFAC, AB, BM ve MASAK listeleri ile gerçek zamanlı eşleştirme prosedürü."),
    ("DONUK ALACAKLAR (NPL) VE YENİDEN YAPILANDIRMA", "Aşama 1, Aşama 2 ve Aşama 3 sınıflandırmaları, TFRS 9 beklenen kredi zararı (ECL) hesaplamaları."),
    ("BİLGİ TEKNOLOJİLERİ GENEL KONTROLLERİ (ITGC) VE PAM", "Ayrıcalıklı hesap yönetimi (CyberArk), SSH erişimleri ve MFA uygulama standartları."),
    ("BULUT ALTYAPISI VE DEVOPS GÜVENLİĞİ", "Mikroservis mimarisi, Kubernetes cluster güvenliği ve CI/CD pipeline denetimleri."),
    ("SİBER OLAY MÜDAHALE VE İŞ SÜREKLİLİĞİ (BCP)", "SOC 7/24 izleme, fidye yazılımı (ransomware) protokolleri ve olağanüstü durum merkezi (DRC) testleri."),
    ("İNSAN KAYNAKLARI VE İÇ DOLANDIRICILIK ÖNLEME", "Kritik rol arka plan araştırmaları, zorunlu blok izin (mandatory leave) uygulamaları."),
    ("DIŞ KAYNAK VE TEDARİKÇİ RİSK YÖNETİMİ", "Üçüncü taraf hizmet sağlayıcıların güvenlik denetimleri ve SLA izleme mekanizmaları."),
    ("SÜRDÜRÜLEBİLİRLİK VE ESG KREDİLENDİRME KRİTERLERİ", "Yeşil finansman standartları, AB Taksonomisi ve karbon emisyon limitleri."),
    ("İÇ DENETİM VE GÜVENCE STANDARTLARI (IIA 2026)", "Sürekli denetim metodolojisi, 5C bulgu yazım standartları ve kalite güvence gözden geçirmeleri.")
]

for sec_idx, (sec_title, sec_desc) in enumerate(sections_data, 1):
    doc.add_heading(f'{sec_idx}. {sec_title}', level=2)
    doc.add_paragraph(sec_desc)
    
    # Her ana bölüm için detaylı kurumsal alt maddeler (100 sayfa hacmine ulaşmak için)
    for sub_idx in range(1, 8):
        doc.add_heading(f'{sec_idx}.{sub_idx}. Operasyonel Uygulama Esasları ve Kontrol Kuralları - Madde {sub_idx}', level=3)
        doc.add_paragraph(
            f'Bu alt madde ({sec_idx}.{sub_idx}); bankanın tüm yurtiçi ve yurtdışı şubelerinde görev yapan '
            f'operasyon personeli, risk yöneticileri ve iç kontrol birimleri için bağlayıcıdır. '
            f'İşbu kural uyarınca her işlem sistemsel loglara kaydedilmeli, görevler ayrılığı (SoD) prensibine uyulmalı, '
            f'hiçbir şart altında tek bir çalışan tarafından uçtan uca tamamlanmamalıdır. '
            f'Yetki aşımı veya kural ihlali durumlarında Teftiş Kurulu Başkanlığı na derhal bildirimde bulunulur.'
        )
        # Detaylı Kontrol ve Tolerans Tabloları
        table = doc.add_table(rows=1, cols=4)
        for i, h in enumerate(['Risk Göstergesi', 'Tolerans Sınırı', 'Yetkili Onay Mercii', 'Kontrol Kodu']):
            table.rows[0].cells[i].text = h
        
        tab_rows = [
            (f'Kritik Limit Aşımı - Kademe {sub_idx}', f'{sub_idx * 5}.000.000 TL', 'Kredi Komitesi & GMY', f'CTRL-{sec_idx:02d}-{sub_idx:02d}-A'),
            (f'Ekspertiz Sapma Oranı', f'%{sub_idx * 3}', 'Bağımsız Lisanslı Değerleme', f'CTRL-{sec_idx:02d}-{sub_idx:02d}-B'),
            (f'MASAK Şüpheli İşlem Sinyali', '0 Tolerans (Anında Bildirim)', 'Uyum Görevlisi & MASAK', f'CTRL-{sec_idx:02d}-{sub_idx:02d}-C')
        ]
        for tr in tab_rows:
            r_cells = table.add_row().cells
            for j, val in enumerate(tr):
                r_cells[j].text = val
        doc.add_paragraph()

doc.save(os.path.join(case_dir, 'devasa_kuresel_bankacilik_ve_kredi_proseduru_100sayfa.docx'))
print("✅ 100 Sayfalık Word Dokümanı Hazır!")

print("⏳ 2. 25.000 Satırlık 6 Sekmeli Devasa Excel Tablosu Üretiliyor...")

# 2. 25.000 SATIRLIK DEVASA EXCEL TABLOSU
with pd.ExcelWriter(os.path.join(case_dir, 'devasa_bankacilik_ve_swift_loglari_25000satir.xlsx'), engine='openpyxl') as writer:
    np.random.seed(2026)

    # Sekme 1: Kredi Tahsisleri (8.000 Satır)
    n1 = 8000
    loans = []
    branches = ['Levent_Kurumsal', 'Maslak_Ticari', 'Kadikoy_Ticari', 'Ankara_Merkez', 'Izmir_Ege', 'Bursa_Sanayi', 'Gaziantep_Kobi', 'Antalya_Turizm']
    for i in range(1, n1 + 1):
        amt = float(np.random.choice([
            np.random.uniform(500000, 4900000),
            np.random.uniform(5100000, 24000000),
            np.random.uniform(25000000, 98000000),
            np.random.uniform(100000000, 450000000)
        ], p=[0.4, 0.35, 0.2, 0.05]))
        ltv = round(float(np.random.uniform(0.45, 1.55)), 2)
        val_comp = 'SPK_Lisansli_Degerleme' if np.random.rand() > 0.35 else 'Sirket_Beyani_Lisanssiz'
        loans.append({
            'loan_ref': f'LON-2026-{i:06d}',
            'customer_id': f'CUST-{(i%1500)+1:05d}',
            'branch': np.random.choice(branches),
            'loan_amount_tl': round(amt, 2),
            'appraised_value_tl': round(amt / max(ltv, 0.1), 2),
            'ltv_ratio': ltv,
            'valuation_firm': val_comp,
            'first_approver': 'kredi_uzman_ali',
            'second_approver': 'sube_muduru_hakan' if amt < 5000000 else ('gmy_onur' if np.random.rand() > 0.4 else None),
            'is_ltv_breach': ltv > 0.75,
            'is_unauthorized_limit': amt > 5000000 and np.random.rand() > 0.85
        })
    pd.DataFrame(loans).to_excel(writer, sheet_name='Kredi_Tahsis_ve_Teminat', index=False)

    # Sekme 2: Swift Transferleri (7.000 Satır)
    n2 = 7000
    swifts = []
    juris = ['TR', 'US', 'GB', 'DE', 'CH', 'VG', 'PA', 'CY', 'KY', 'SC', 'AE']
    for i in range(1, n2 + 1):
        j = np.random.choice(juris)
        amt_u = float(np.random.choice([
            np.random.uniform(10000, 95000),
            np.random.uniform(98000, 99990),
            np.random.uniform(250000, 4500000),
            np.random.uniform(5000000, 35000000)
        ], p=[0.4, 0.25, 0.3, 0.05]))
        is_off = j in ['VG', 'PA', 'CY', 'KY', 'SC']
        swifts.append({
            'swift_id': f'SWIFT-2026-{i:06d}',
            'sender_account': f'TR330006100511{(i%1200)+1000:04d}',
            'beneficiary_country': j,
            'amount_usd': round(amt_u, 2),
            'currency': 'USD',
            'transfer_date': f'2026-{(i%12)+1:02d}-{(i%28)+1:02d}',
            'is_offshore_haven': is_off,
            'masak_filter_cleared': False if is_off and amt_u > 100000 and np.random.rand() > 0.3 else True,
            'pep_match_flag': True if is_off and np.random.rand() > 0.8 else False
        })
    pd.DataFrame(swifts).to_excel(writer, sheet_name='Swift_ve_Kambiyo', index=False)

    # Sekme 3: NPL ve Donuk Alacaklar (4.000 Satır)
    n3 = 4000
    npls = []
    for i in range(1, n3 + 1):
        npls.append({
            'npl_id': f'NPL-2026-{i:05d}',
            'customer_name': f'BORCLU-SIRKET-{(i%600)+1:04d}',
            'npl_stage': np.random.choice(['Stage_1', 'Stage_2_Watchlist', 'Stage_3_Substandard', 'Stage_5_Loss']),
            'exposure_tl': round(float(np.random.uniform(500000, 85000000)), 2),
            'overdue_days': np.random.choice([15, 45, 95, 185, 365, 720]),
            'restructured_count': np.random.randint(0, 6),
            'is_evergreen_refinanced': True if np.random.rand() > 0.65 else False
        })
    pd.DataFrame(npls).to_excel(writer, sheet_name='NPL_Donuk_Alacaklar', index=False)

    # Sekme 4: Core Banking Yetki Logları (3.000 Satır)
    n4 = 3000
    logs = []
    for i in range(1, n4 + 1):
        logs.append({
            'log_id': f'LOG-2026-{i:06d}',
            'user_id': np.random.choice(['sube_ali', 'kredi_onur', 'admin_root', 'masak_deniz', 'ops_kemal']),
            'menu_action': np.random.choice(['LIMIT_OVERRIDE', 'AML_WHITELIST', 'FX_SPOT_DEAL', 'NPL_STAGE_CHANGE']),
            'timestamp': f'2026-06-{(i%28)+1:02d} {np.random.randint(0,24):02d}:{np.random.randint(0,60):02d}:00',
            'mfa_status': 'FAILED' if np.random.rand() > 0.75 else 'PASSED'
        })
    pd.DataFrame(logs).to_excel(writer, sheet_name='Core_Banking_Loglari', index=False)

    # Sekme 5: PEP ve Yaptırım Taraması (2.000 Satır)
    n5 = 2000
    peps = []
    for i in range(1, n5 + 1):
        peps.append({
            'pep_id': f'PEP-2026-{i:05d}',
            'target_name': f'PEP-ENTITY-{(i%400)+1:04d}',
            'category': np.random.choice(['Kamu_Gorevlisi', 'Bakan_Akrabasi', 'Askeri_Yetkili', 'Ihale_Baskani']),
            'risk_score': np.random.choice([80, 90, 95, 99]),
            'sanction_source': np.random.choice(['OFAC', 'EU_Sanctions', 'MASAK_Liste']),
            'status': 'BLOCKED' if np.random.rand() > 0.4 else 'ACTIVE_VIOLATION'
        })
    pd.DataFrame(peps).to_excel(writer, sheet_name='PEP_ve_Yaptirimlar', index=False)

    # Sekme 6: Mizan ve Muhasebe Kayıtları (1.000 Satır)
    n6 = 1000
    accs = []
    for i in range(1, n6 + 1):
        accs.append({
            'account_code': f'ACC-{100 + (i%50):03d}',
            'account_name': np.random.choice(['Krediler_Hesabi', 'NPL_Karsiliklari', 'Doviz_Pozisyonu', 'Komisyon_Gelirleri', 'Kanunen_Kabul_Edilmeyen_Giderler']),
            'debit_tl': round(float(np.random.uniform(10000, 50000000)), 2),
            'credit_tl': round(float(np.random.uniform(10000, 50000000)), 2),
            'currency': np.random.choice(['TRY', 'USD', 'EUR'])
        })
    pd.DataFrame(accs).to_excel(writer, sheet_name='Mizan_Muhasebe', index=False)

print("✅ 25.000 Satırlık Devasa Excel Tablosu Hazır!")

# 3. BÜYÜK ADLİ RAPOR
doc_adl = docx.Document()
doc_adl.add_heading('T.C. BDDK & MASAK MÜŞTEREK TEFTİŞ HEYETİ — ADLİ VE CEZAİ SORUŞTURMA RAPORU', level=1)
doc_adl.add_paragraph('Rapor Referansı: 2026/BDDK-MASAK-ORTAK-ULTRA-01 | Tarih: 28 Ağustos 2026 | Gizlilik: ÇOK GİZLİ / KİŞİYE ÖZEL')

doc_adl.add_heading('1. KAPSAM VE SORUŞTURMA BULGULARI', level=2)
doc_adl.add_paragraph(
    'Bankanın 2026 yılı ilk 6 aylık kredi tahsis ve fon transfer operasyonlarında yapılan incelemelerde; '
    'Levent ve Maslak şubeleri üzerinden 8.000 adet kredi dosyasının 1.420 sinde sahte/şişirilmiş gayrimenkul ekspertiz raporları kullanıldığı, '
    'toplam 145.000.000 USD tutarında teminatsız kredi riskinin oluştuğu, '
    '7.000 adet SWIFT transferinden 840 ında MASAK Uyum filtrelerinin yazılımsal olarak devre dışı bırakıldığı '
    've 78.500.000 USD tutarındaki suç gelirinin Panama ve BVI offshore hesaplarına aklandığı tespit edilmiştir.'
)
doc_adl.save(os.path.join(case_dir, 'devasa_masak_ve_bddk_adli_teftis_raporu.docx'))
print("✅ Devasa Adli Teftiş Raporu Hazır!")

print("🎉 Tüm Endüstriyel Devasa Dosyalar Başarıyla Oluşturuldu:", os.listdir(case_dir))
