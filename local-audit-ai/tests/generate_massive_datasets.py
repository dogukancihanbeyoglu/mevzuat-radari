"""
Local Audit AI - Büyük ve Karmaşık Test Dosyaları Üretici
PDF, DOCX, XLSX, CSV ve TXT formatlarında uçtan uca kurumsal denetim veri paketi üretir.
"""
import os
import docx
import pandas as pd
import numpy as np

target_dir = "/Users/dogukancihanbeyoglu/Gemini/local-audit-ai/sample_test_files/massive_test_suite"
os.makedirs(target_dir, exist_ok=True)

# 1. DOCX 1: Hazine ve Türev Ürünler Süreç Prosedürü (Tablolu ve Limitli)
doc_hazine = docx.Document()
doc_hazine.add_heading('Mega Holding A.Ş. — Hazine, Nakit Yönetimi ve Türev İşlemler Prosedürü (TR-HAZ-2026)', level=1)
doc_hazine.add_paragraph('Bu prosedür, Holding ve bağlı 14 iştirakinin banka hesapları, vadeli mevduat, kredi kullanımı, swap/forward ve türev işlem limitlerini belirler.')
doc_hazine.add_heading('1. Günlük Likidite ve Fon Transfer Limitleri', level=2)
doc_hazine.add_paragraph('500.000 USD veya muadili döviz transferlerinde Hazine Direktörü ve Mali İşler GMY çift imzası aranır (C1). 1.000.000 USD üzeri işlemlerde Yönetim Kurulu Finans Komitesi onayı zorunludur. Ancak piyasa oynaklığı durumunda Hazine Uzmanı sözlü talimatla 2.000.000 USD limitli spot döviz alım/satımı yapabilmekte ve teyit mektupları ertesi gün sisteme girilmektedir (W1 - Kritik Piyasa Riski).')

table_limits = doc_hazine.add_table(rows=1, cols=4)
hdr_cells = table_limits.rows[0].cells
hdr_cells[0].text = 'İşlem Türü'
hdr_cells[1].text = 'Yetkili Rol'
hdr_cells[2].text = 'Tekli Onay Limiti'
hdr_cells[3].text = 'Çift Onay Limiti'

limits_data = [
    ('Spot Döviz Alım/Satım', 'Hazine Uzmanı', '50.000 USD', '250.000 USD'),
    ('Forward / Vadeli İşlem', 'Hazine Müdürü', '250.000 USD', '1.000.000 USD'),
    ('Faiz ve Para Swapı', 'Hazine Direktörü', '1.000.000 USD', '5.000.000 USD'),
    ('Yurt Dışı Kredi Kullanımı', 'CFO & İcra Kurulu', '5.000.000 USD', '50.000.000 USD')
]
for row in limits_data:
    row_cells = table_limits.add_row().cells
    for i, val in enumerate(row):
        row_cells[i].text = val

doc_hazine.save(os.path.join(target_dir, 'hazine_ve_turev_islemler_proseduru.docx'))

# 2. DOCX 2: İnsan Kaynakları, Üst Düzey Yönetici Prim ve Bordro Süreci
doc_ik = docx.Document()
doc_ik.add_heading('İnsan Kaynakları, Yönetici Primleri ve Bordro Denetim Kılavuzu', level=1)
doc_ik.add_paragraph('Bordro hesaplamaları her ayın 25 inde SAP HCM modülü üzerinden puantaj verileriyle otomatik yapılır. Fazla mesai onayları departman müdürleri tarafından onaylanır.')
doc_ik.add_paragraph('Üst yönetim yıllık performans primleri Ücretlendirme Komitesi kararıyla belirlenir. Ancak 2025 yılı performans prim havuzundan 3 İcra Kurulu üyesine komite kararı olmaksızın toplam 12.500.000 TL avans/prim ödemesi yapıldığı ve bordro tahakkukunun geriye dönük düzeltildiği mülakatlarda tespit edilmiştir (W2 - Vergisel ve Yönetişim Riski).')
doc_ik.save(os.path.join(target_dir, 'ik_yonetici_prim_ve_bordro_sureci.docx'))

# 3. XLSX 1: 1000 Satırlık Detaylı Mizan ve Hesap Planı
accounts = ['100 Kasa', '102 Bankalar', '120 Alıcılar', '150 İlk Madde Malzeme', '153 Ticari Mallar', '255 Demirbaşlar', '320 Satıcılar', '360 Ödenecek Vergi', '600 Yurtiçi Satışlar', '621 Satılan Mal Maliyeti', '770 Genel Yönetim Gideri', '780 Finansman Gideri']
np.random.seed(42)
mizan_rows = []
for i in range(1, 1001):
    acc = np.random.choice(accounts)
    borc = round(float(np.random.exponential(scale=50000)), 2)
    alacak = round(float(np.random.exponential(scale=45000)), 2)
    bakiye = round(borc - alacak, 2)
    mizan_rows.append({
        'Hesap_Kodu': acc.split()[0],
        'Hesap_Adi': ' '.join(acc.split()[1:]),
        'Fis_No': f'FIS-2026-{i:05d}',
        'Tarih': f'2026-{(i%12)+1:02d}-{(i%28)+1:02d}',
        'Borc_Tutari': borc,
        'Alacak_Tutari': alacak,
        'Bakiye': bakiye,
        'Kayit_Yapan': np.random.choice(['muhasebe_uzm1', 'muhasebe_uzm2', 'stajyer_user', 'sistem_oto']),
        'Onaylayan': np.random.choice(['mali_isler_mdr', 'muhasebe_mdr', None])
    })
df_mizan = pd.DataFrame(mizan_rows)
df_mizan.to_excel(os.path.join(target_dir, 'mizan_ve_muhasebe_kayitlari_1000satir.xlsx'), index=False)

# 4. XLSX 2: Kritik Sistem ve Veritabanı Erişim Logları (SOC / ITGC)
log_rows = []
for i in range(1, 301):
    user = np.random.choice(['root', 'admin_sys', 'db_dba', 'developer_01', 'external_consultant', 'app_service'])
    event = np.random.choice(['LOGIN_SUCCESS', 'LOGIN_FAILED', 'PERMISSION_CHANGE', 'DATABASE_DUMP', 'CONFIG_EDIT', 'PASSWORD_RESET'])
    env = np.random.choice(['PROD_DB_PRIMARY', 'PROD_K8S_CORE', 'DEV_CLUSTER', 'UAT_GATEWAY'])
    hour = np.random.randint(0, 24)
    ip = f'192.168.1.{np.random.randint(10, 250)}' if np.random.rand() > 0.1 else f'185.220.{np.random.randint(1, 255)}.{np.random.randint(1, 255)}'
    
    log_rows.append({
        'log_id': f'LOG-{i:04d}',
        'timestamp': f'2026-07-{(i%28)+1:02d} {hour:02d}:{np.random.randint(0,60):02d}:00',
        'username': user,
        'event_type': event,
        'target_environment': env,
        'source_ip': ip,
        'is_vpn': bool(np.random.choice([True, False])),
        'is_mfa_used': False if user in ['root', 'external_consultant'] and hour in [23, 0, 1, 2, 3, 4] else True,
        'session_duration_min': int(np.random.randint(1, 180))
    })
df_logs = pd.DataFrame(log_rows)
df_logs.to_excel(os.path.join(target_dir, 'siber_guvenlik_ve_pam_erisim_loglari.xlsx'), index=False)

# 5. CSV: E-Ticaret Sipariş, İade ve İptal Anomalileri
csv_rows = []
for i in range(1, 401):
    csv_rows.append({
        'order_id': f'ORD-{10000+i}',
        'customer_id': f'CUST-{np.random.randint(100, 300)}',
        'order_date': f'2026-05-{(i%28)+1:02d}',
        'order_amount_tl': round(float(np.random.uniform(500, 25000)), 2),
        'payment_method': np.random.choice(['CreditCard', 'WireTransfer', 'DigitalWallet', 'CashOnDelivery']),
        'discount_code': np.random.choice(['YONETIM100', 'VIP50', 'YAZ2026', None], p=[0.05, 0.1, 0.3, 0.55]),
        'refund_status': np.random.choice(['NONE', 'REFUNDED', 'CHARGEBACK', 'PARTIAL_REFUND'], p=[0.7, 0.15, 0.05, 0.1]),
        'refund_amount_tl': round(float(np.random.uniform(500, 25000)), 2) if np.random.rand() > 0.7 else 0.0,
        'ip_country': np.random.choice(['TR', 'US', 'DE', 'RU', 'CY', 'AZ'])
    })
df_csv = pd.DataFrame(csv_rows)
df_csv.to_csv(os.path.join(target_dir, 'e_ticaret_siparis_ve_iade_anomalileri.csv'), index=False)

# 6. TXT: Adli Soruşturma Mülakat Tutanakları (Doğal Dil ve Çelişkili İfadeler)
tutanak_text = """T.C. MEGA HOLDİNG A.Ş. İÇ DENETİM VE TEFTİŞ KURULU BAŞKANLIĞI
SORUŞTURMA MÜLAKAT TUTANAĞI — NO: SOR-2026/04

Konu: Ambar Stok Açıkları, Hurda Malzeme Çıkışları ve Usulsüz Gider Kayıtları
Tarih: 18 Temmuz 2026 | Yer: Gebze Üretim Kampüsü Denetim Odası
Mülakat Yapılan: M.K. (Fabrika Müdürü) ve A.T. (Eski Ambar Şefi)

[MÜLAKAT DİYALOĞU]:
Baş Denetçi: 2025 yıl sonu fiili sayımında tespit edilen 8.400.000 TL tutarındaki stok açığının Yönetim Kurulu kararı olmaksızın doğrudan 689 hesaba gider yazılması talimatını kim verdi?
Fabrika Müdürü (M.K.): Üretim hattında aksama olmaması için muhasebe müdürüyle sözlü mutabık kaldık. Şirket hedeflerini tutturmak için hızlı hareket etmemiz gerekiyordu. Yönetim Kurulu Başkanı'na konuyu ayaküstü ilettim, onay verdiğini düşündüm.
Eski Ambar Şefi (A.T.): Bu doğru değil. Gece vardiyasında fabrika dışına çıkarılan 3 kamyon hurda bakır malzeme için Fabrika Müdürü bizzat kantar fişi kesilmemesini ve güvenlik kameralarının bakımda gösterilmesini emretti. Malzeme faturasız olarak X Hurdacılık Ltd. firmasına teslim edildi, bedeli resmi hesaplara girmedi.
Baş Denetçi: Bahsettiğiniz 3.200.000 TL değerindeki malzemenin bedeli kime ve nasıl ödendi?
Eski Ambar Şefi (A.T.): Fabrika Müdürü'nün talimatıyla elden nakit teslim alındı. Elimde teslim tutanağının gayriresmi kopyası ve araç plaka kayıtları mevcuttur, denetim ekibine teslim ediyorum.
"""
with open(os.path.join(target_dir, 'adli_sorusturma_mulakat_tutanaklari_ham.txt'), 'w', encoding='utf-8') as f:
    f.write(tutanak_text)

print('Tüm büyük ve karmaşık test dosyaları üretildi:', os.listdir(target_dir))
