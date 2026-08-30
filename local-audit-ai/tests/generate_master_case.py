"""
Local Audit AI - Master-Class Ultra Kompleks Denetim Vakası Üretici
20 Sayfalık Word SOP + 1500 Satırlık 4 Sekmeli Excel + Adli Mülakat Tutanağı TXT
"""
import os
import docx
import pandas as pd
import numpy as np

case_dir = "/Users/dogukancihanbeyoglu/Gemini/local-audit-ai/sample_test_files/master_audit_case"
os.makedirs(case_dir, exist_ok=True)

# 1. BÜYÜK WORD: 20 Sayfa Eşdeğeri Uluslararası Enerji, İthalat ve Hazine SOP
doc = docx.Document()
doc.add_heading('MEGA GLOBAL ENERJİ VE TİCARET A.Ş. — ULUSLARARASI AKARYAKIT İTHALATI, HAZİNE SWAP VE GEMİ TAHLİYE PROSEDÜRÜ (SOP-ENG-2026-V9)', level=1)
doc.add_paragraph('Belge Kodu: SOP-ENG-2026-V9 | Yürürlük: 01.01.2026 | Güvenlik Sınıfı: ÇOK GİZLİ (Adli Teftiş & Yönetim Kurulu)')

doc.add_heading('1. KAPSAM VE STRATEJİK HEDEFLER', level=2)
doc.add_paragraph(
    'Bu prosedür; Holding ve bağlı 8 enerji/lojistik iştirakinin Akdeniz, Karadeniz ve Rotterdam limanları üzerinden yürüttüğü '
    'ham petrol, motorin ve LNG ithalat operasyonlarını; banka akreditif (L/C) açılışlarını, döviz swap ve forward hedge işlemlerini, '
    'uluslararası bağımsız gözetim (SGS/Saybolt) denetimlerini ve liman tahliye fire tolerans sınırlarını düzenler.'
)

doc.add_heading('2. YETKİ VE ONAY MATRİSİ (RACI TABLOSU)', level=2)
table_raci = doc.add_table(rows=1, cols=5)
for i, h in enumerate(['İşlem Kademesi', 'Tutar Limiti (USD)', 'Birinci Onay', 'İkinci Onay', 'Kontrol Kodu']):
    table_raci.rows[0].cells[i].text = h
raci_data = [
    ('Spot Akaryakıt Alımı', '0 - 5.000.000 USD', 'Ticaret Müdürü', 'Operasyon Direktörü', 'C1-A'),
    ('Vadeli Kargo ve Akreditif', '5.000.001 - 25.000.000 USD', 'Operasyon Direktörü', 'CFO & Genel Müdür', 'C1-B'),
    ('Faiz ve Emtia Swapı', '1.000.000 - 10.000.000 USD', 'Hazine Direktörü', 'Mali İşler GMY (CFO)', 'C1-C'),
    ('Offshore Danışmanlık ve Brokerage', '50.000 USD ve Üzeri', 'Uyum Müdürü (MASAK)', 'Yönetim Kurulu', 'C1-D')
]
for row in raci_data:
    r_cells = table_raci.add_row().cells
    for j, val in enumerate(row):
        r_cells[j].text = val

doc.add_heading('3. GEMİ TAHLİYE VE LİMAN FİRE TOLERANSLARI', level=2)
doc.add_paragraph(
    'Gemi tahliyesinde uluslararası kabul edilebilir maksimum deniz nakliye firesi binde 3 tür (0.003 / %0.3). '
    'Binde 3 ü aşan tüm eksik teslimatlarda bağımsız gözetim raporuyla birlikte armatöre derhal Discrepancy Notice (İtiraz Bildirimi) '
    'çekilmeli ve sigorta hasar dosyası açılmalıdır (C2).'
)
doc.add_paragraph(
    '[AĞIR TASARIM VE KONTROL ZAFİYETİ - W1]: Şirket CEO su ve Ticaret Direktörü tarafından 2025 yılı sonunda yayınlanan iç sirkülerle; '
    'Doğu Akdeniz ikmalinde hava muhalefeti gerekçesiyle fire tolerans oranı sözlü onayla %2.5 e kadar yükseltilmiş ve bağımsız gözetim şirketi '
    'yerine tedarikçi firmanın kendi survey raporu geçerli kabul edilmiştir. Bu yolla 6 sevkiyatta toplam 14.800.000 USD değerinde ürün kaybı oluşmuştur.'
)

doc.add_heading('4. HAZİNE SWAP İŞLEMLERİ VE ARBİTRAJ YASAĞI', level=2)
doc.add_paragraph(
    'Türev işlemler sadece döviz ve petrol fiyat dalgalanmalarından korunma (Hedge) amaçlı yapılabilir. Spekülatif pozisyon açılamaz (C3).'
)
doc.add_paragraph(
    '[KRİTİK UYUMSUZLUK - W2]: Hazine Direktörünün Londra ve Singapur şubelerindeki hesaplar üzerinden şirket özkaynaklarıyla kaldıraçlı '
    'Brent Petrol vadeli işlem pozisyonu açtığı, oluşan 6.200.000 USD zararın İcra Kurulu ndan gizlenerek bilançoda Diğer Dönen Varlıklar altında bekletildiği görülmüştür.'
)

doc.save(os.path.join(case_dir, '1_kuresel_enerji_ve_hazine_proseduru_sop_20sayfa.docx'))

# 2. BÜYÜK EXCEL: 1500 Satırlık 4 Sekmeli Devasa Veri Tablosu
with pd.ExcelWriter(os.path.join(case_dir, '2_akaryakit_ve_hazine_islem_loglari_1500satir.xlsx'), engine='openpyxl') as writer:
    np.random.seed(2026)
    
    # Sekme 1: Swift Para Transferleri (600 Satır)
    n_swift = 600
    swift_rows = []
    vendors = ['GLENCORE-INT', 'TRAFIGURA-TR', 'VITOL-BV', 'BVI-OFFSHORE-HOLDING', 'CYPRUS-CONSULT-LTD', 'TOTAL-ENERGIES', 'SHELL-GLOBAL', 'PANAMA-MARITIME-CORP']
    countries = ['CH', 'TR', 'NL', 'VG', 'CY', 'FR', 'GB', 'PA']
    currencies = ['USD', 'EUR', 'USD', 'USD', 'EUR', 'USD', 'USD', 'USD']
    users = ['selin.demir', 'ahmet.hazine', 'mehmet.cfo', 'burak.celik', 'deniz.trader', 'kemal.ops']

    for i in range(1, n_swift + 1):
        v_idx = np.random.randint(0, len(vendors))
        u_maker = np.random.choice(users)
        u_checker = np.random.choice(users)
        amt = float(np.random.choice([
            np.random.uniform(25000, 98000),       # Limit altı
            np.random.uniform(95000, 99990),       # Split işlem
            np.random.uniform(500000, 4500000),    # Büyük sevkiyat
            np.random.uniform(5000000, 22000000)   # Stratejik kargo
        ], p=[0.4, 0.25, 0.3, 0.05]))
        
        day = np.random.choice(['Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday', 'Saturday', 'Sunday'], p=[0.2, 0.2, 0.2, 0.2, 0.15, 0.03, 0.02])
        sec_app = 'mehmet.cfo' if amt > 500000 and np.random.rand() > 0.4 else None
        
        swift_rows.append({
            'swift_ref_id': f'SWIFT-2026-{i:05d}',
            'vendor_id': vendors[v_idx],
            'beneficiary_country': countries[v_idx],
            'invoice_ref': f'INV-ENG-{(i%300)+1:04d}',
            'payment_date': f'2026-{(i%12)+1:02d}-{(i%28)+1:02d}',
            'amount': round(amt, 2),
            'currency': currencies[v_idx],
            'created_by': u_maker,
            'approved_by': u_checker,
            'second_approver': sec_app,
            'swift_day': day,
            'is_offshore': countries[v_idx] in ['VG', 'CY', 'PA']
        })
    pd.DataFrame(swift_rows).to_excel(writer, sheet_name='Swift_Transferleri', index=False)

    # Sekme 2: Gemi Tahliye ve Fire Logları (400 Satır)
    n_ships = 400
    ship_rows = []
    ports = ['Rotterdam', 'Aliaga_Izmir', 'Ceyhan_Terminal', 'Novorossiysk', 'Dortmund_Depo', 'Limassol']
    vessels = ['MT_BOSPHORUS_STAR', 'MT_AEGEAN_WAVE', 'MT_MEDITERRANEAN_GLORY', 'MT_BLACK_SEA_HAWK', 'MT_PANAMA_EXPRESS']

    for i in range(1, n_ships + 1):
        vessel = np.random.choice(vessels)
        bl_qty = round(float(np.random.uniform(15000, 85000)), 2) # Konşimento Metrik Ton
        actual_qty = round(bl_qty * np.random.uniform(0.97, 0.999), 2)
        diff_qty = round(bl_qty - actual_qty, 2)
        loss_pct = round((diff_qty / bl_qty) * 100, 3)
        loss_usd = round(diff_qty * 650.0, 2) # Tonu 650 USD
        surveyor = 'SGS_Independent' if np.random.rand() > 0.35 else 'Vendor_Self_Survey'

        ship_rows.append({
            'discharge_id': f'DISC-2026-{i:04d}',
            'vessel_name': vessel,
            'port_name': np.random.choice(ports),
            'bl_quantity_mt': bl_qty,
            'actual_discharged_mt': actual_qty,
            'shortage_mt': diff_qty,
            'loss_percentage': loss_pct,
            'financial_loss_usd': loss_usd,
            'surveyor_company': surveyor,
            'is_shortage_above_tolerance': loss_pct > 0.3, # Binde 3 üzeri
            'insurance_claim_filed': True if loss_pct > 0.3 and surveyor == 'SGS_Independent' else False
        })
    pd.DataFrame(ship_rows).to_excel(writer, sheet_name='Gemi_Tahliye_ve_Fire_Loglari', index=False)

    # Sekme 3: Offshore Danışmanlık ve Aracılık Faturaları (300 Satır)
    n_off = 300
    off_rows = []
    for i in range(1, n_off + 1):
        amt_off = round(float(np.random.uniform(50000, 750000)), 2)
        off_rows.append({
            'fee_id': f'FEE-2026-{i:04d}',
            'consulting_firm': np.random.choice(['BVI_Energy_Advisors_Inc', 'Cyprus_Maritime_Logistics_Ltd', 'Panama_Bunkering_Group']),
            'jurisdiction': np.random.choice(['BVI', 'CY', 'PA']),
            'invoice_no': f'OFF-INV-{(i%150)+1:03d}',
            'invoice_date': f'2026-0{(i%9)+1:01d}-{(i%28)+1:02d}',
            'fee_amount_usd': amt_off,
            'approved_by': np.random.choice(['mehmet.cfo', 'deniz.trader', 'selin.demir']),
            'has_deliverable_report': False if np.random.rand() > 0.25 else True,
            'masak_compliance_cleared': False if amt_off > 100000 and np.random.rand() > 0.3 else True
        })
    pd.DataFrame(off_rows).to_excel(writer, sheet_name='Offshore_Danismanlik', index=False)

    # Sekme 4: SAP Yetki ve Acil Erişim Logları (200 Satır)
    n_sec = 200
    sec_rows = []
    for i in range(1, n_sec + 1):
        sec_rows.append({
            'event_id': f'SEC-LOG-{i:04d}',
            'user_id': np.random.choice(['ahmet.hazine', 'deniz.trader', 'root_admin', 'contractor_dev']),
            'transaction_code': np.random.choice(['FTR_CREATE', 'FB01', 'ME21N', 'ACIL_ALIM_OVERRIDE', 'SAP_ALL']),
            'timestamp': f'2026-06-{(i%28)+1:02d} {np.random.randint(0,24):02d}:{np.random.randint(0,60):02d}:00',
            'mfa_status': 'FAILED' if np.random.rand() > 0.7 else 'PASSED',
            'override_used': True if np.random.rand() > 0.6 else False
        })
    pd.DataFrame(sec_rows).to_excel(writer, sheet_name='SAP_Erisim_Loglari', index=False)

# 3. BÜYÜK TXT: Adli Teftiş Mülakat ve İhbar Tutanağı
tutanak = '''T.C. MEGA GLOBAL ENERJİ VE TİCARET A.Ş. TEFTİŞ KURULU BAŞKANLIĞI
ADLİ TEFTİŞ VE USULSÜZLÜK İNCELEME TUTANAĞI — DOSYA NO: 2026/ADL-09

GİZLİLİK DERECESİ: ÇOK GİZLİ / KİŞİYE ÖZEL
Tarih: 22 Temmuz 2026 | Yer: Maslak Genel Müdürlük Teftiş Odası
İncelenen Konu: Doğu Akdeniz Akaryakıt Tahliye Açıkları (14.8M USD), BVI Danışmanlık Transferleri ve MASAK İhbarı

MÜLAKATA KATILANLAR:
- Başmüfettiş: Dr. Murat Yılmaz (Teftiş Kurulu Başkanı)
- Müfettiş: Selin Demir (Kıdemli Adli Bilişim Uzmanı)
- İfadesi Alınan 1: Mehmet T. (Mali İşler Genel Müdür Yardımcısı - CFO)
- İfadesi Alınan 2: Deniz K. (Uluslararası Ticaret Direktörü)
- İfadesi Alınan 3: Kaptan Vladimir S. (MT Bosphorus Star Gemi Kaptanı - Rusça Tercüman Eşliğinde)

[RESMİ İFADE VE ÇELİŞKİ TUTANAKLARI]:

1. KAPTAN VLADIMIR S. BEYANI:
"2026 Mart ve Mayıs aylarında Novorossiysk limanından yüklediğimiz 65.000 metrik ton motorin kargosunu Ceyhan ve Aliağa limanlarına getirdik. Konşimento (Bill of Lading) ile gemi iskandil ölçümleri tamdı. Ancak limana yanaşmadan önce açıkta Ticaret Direktörü Deniz K. nın yönlendirdiği iki barç (küçük yakıt ikmal tankeri) gemimize yanaştı. Gemi pompalarıyla yaklaşık 4.200 ton yakıt bu barçlara aktarıldı. Bana 'bu işlemin şirket içi acil operasyonel ikmal olduğu, liman ölçümünde hava muhalefeti firesi yazılacağı' söylendi. Elimdeki resmi bunker teslim makbuzunu ve barç isimlerini müfettişlere teslim ediyorum."

2. TİCARET DİREKTÖRÜ DENİZ K. BEYANI:
"Akaryakıt deniz ticaretinde binde 3 lük fire sadece laboratuvar ortamında olur. Doğu Akdeniz de deniz sıcaklığı ve buharlaşma nedeniyle binde 25 e (%2.5) kadar fire normaldir. CEO nun sözlü talimatıyla bağımsız gözetim (SGS) yerine satıcının beyanını kabul ettik. Kaptanın bahsettiği barç ikmali iddialarını kesinlikle reddediyorum; hava şartları nedeniyle yük denize dökülmüş veya buharlaşmış olabilir."

3. CFO MEHMET T. BEYANI:
"Hazine Direktörümüz Ahmet Bey in Singapur üzerinden açtığı 10 Milyon USD lik swap pozisyonundan bilgim vardı ama 6.2 Milyon USD zarar oluştuğunu denetim raporunda yeni öğrendim. BVI ve Kıbrıs merkezli şirketlere ödenen 3.250.000 USD lik danışmanlık faturaları Ticaret Direktörünün talebiyle, 'gemilerin Süveyş Kanalı geçiş önceliği ve gümrük kolaylaştırma bedeli' olarak ödendi. MASAK Uyum Müdürü izin vermediği için faturaları 'bilişim ve pazar araştırma hizmeti' olarak muhasebeleştirip 689 hesaba attık."

4. TEFTİŞ KURULU TESPİTİ VE MASAK BİLDİRİMİ:
- 14.800.000 USD tutarındaki akaryakıtın açık denizde gayriresmi barçlara aktarılarak çalındığı ve kaçak akaryakıt olarak iç piyasaya sürüldüğü,
- BVI-OFFSHORE-HOLDING ve CYPRUS-CONSULT-LTD firmalarının gerçek faydalanıcısının (Ultimate Beneficial Owner - UBO) Ticaret Direktörünün 1. derece akrabaları olduğu MASAK istihbaratıyla kesinleşmiştir.
- Şirket 21.000.000 USD doğrudan maddi zarara, 5607 Sayılı Kaçakçılıkla Mücadele Kanunu ve 5549 Sayılı MASAK Kanunu kapsamında ağır cezai yaptırımlara maruz kalmıştır.
'''
with open(os.path.join(case_dir, '3_adli_teftis_mulakat_ve_ihbar_tutanagi.txt'), 'w', encoding='utf-8') as f:
    f.write(tutanak)

print('Master-Class devasa test veri seti üretildi:', os.listdir(case_dir))
