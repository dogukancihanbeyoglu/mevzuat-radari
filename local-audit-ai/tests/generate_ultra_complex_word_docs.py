"""
Local Audit AI - Ultra Kapsamlı ve Gerçekçi Kurumsal Süreç Prosedürleri Üretici (Word - 10-15 Sayfa Eşdeğeri)
"""
import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

enterprise_dir = "/Users/dogukancihanbeyoglu/Gemini/local-audit-ai/sample_test_files/enterprise_audit_pack"
massive_dir = "/Users/dogukancihanbeyoglu/Gemini/local-audit-ai/sample_test_files/massive_test_suite"
os.makedirs(enterprise_dir, exist_ok=True)
os.makedirs(massive_dir, exist_ok=True)

# 1. BÜYÜK DOKÜMAN 1: Küresel Tedarik Zinciri, Satınalma ve Stok Süreç Prosedürü (Ultra Detaylı)
doc_tedarik = docx.Document()
doc_tedarik.add_heading('MEGA HOLDİNG A.Ş. — KÜRESEL TEDARİK ZİNCİRİ, SATINALMA VE STOK YÖNETİMİ STANDART PROSEDÜRÜ (SOP-SCM-2026-V8)', level=1)
doc_tedarik.add_paragraph('Belge Kodu: SOP-SCM-2026-V8 | Revizyon No: 08 | Yürürlük Tarihi: 01.01.2026 | Gizlilik: ÇOK GİZLİ (Kurumsal)')

# Bölüm 1
doc_tedarik.add_heading('1. AMAÇ VE KAPSAM', level=2)
doc_tedarik.add_paragraph(
    '1.1. Bu prosedürün amacı; Mega Holding A.Ş. ve bağlı 14 iştirakinin yurt içi ve yurt dışı operasyonlarında ihtiyaç duyulan her türlü hammadde, '
    'yardımcı madde, ticari emtia, makine/teçhizat yatırımları, danışmanlık ve genel işletme hizmetlerinin tedarik edilmesinden; mal kabul, '
    'depo stok yönetimi, konsinye malzeme takibi, 3 lü eşleştirme, hurda/fire süreçleri ve yıl sonu envanter sayımlarına kadar olan tüm aşamaların '
    'şeffaf, izlenebilir, hesap verebilir ve IIA/SOX iç kontrol standartlarına tam uyumlu olarak icra edilmesini sağlamaktır.'
)
doc_tedarik.add_paragraph(
    '1.2. Bu prosedür; Türkiye, Almanya, Hollanda ve BAE de faaliyet gösteren tüm üretim tesisleri, lojistik merkezleri, mağazalar ve bağlı ortaklıkları '
    'kapsar. İstisnasız tüm çalışanlar bu prosedür hükümlerine uymakla mükelleftir.'
)

# Bölüm 2
doc_tedarik.add_heading('2. TANIMLAR VE KISALTMALAR', level=2)
tanimlar = [
    ('SAP MM (Materials Management)', 'Satınalma, stok ve envanter yönetiminin yürütüldüğü kurumsal ERP modülü.'),
    ('PR (Purchase Requisition)', 'Kullanıcı birimler tarafından sistem üzerinden oluşturulan resmi satınalma talep belgesi.'),
    ('PO (Purchase Order)', 'Tedarikçiye resmi sipariş taahhüdü oluşturan yasal satınalma siparişi.'),
    ('GRN (Goods Receipt Note - 101 Hareketi)', 'Fiili mal kabulün depoya girişinin yapıldığını gösteren sistem belgesi.'),
    ('Three-Way Match (3 lü Eşleştirme)', 'PO (Sipariş), GRN (Mal Kabul İrsaliyesi) ve Tedarikçi Faturasının miktar, birim fiyat ve ödeme vadesi açısından sistem tarafından otomatik doğrulanması.'),
    ('Konsinye Stok', 'Mülkiyeti tedarikçide kalmak kaydıyla şirket depolarında tutulan ve ancak fiilen üretime çekildiğinde faturası kesilen stok modeli.'),
    ('TKT (Teklif Karşılaştırma Tablosu)', 'En az 3 bağımsız teklifin teknik, ticari ve vade şartlarını puanlayan resmi komisyon tutanağı.')
]
for t, d in tanimlar:
    p = doc_tedarik.add_paragraph()
    p.add_run(f'• {t}: ').bold = True
    p.add_run(d)

# Bölüm 3: RACI Matrisi
doc_tedarik.add_heading('3. ROL VE SORUMLULUKLAR (RACI MATRİSİ)', level=2)
table_raci = doc_tedarik.add_table(rows=1, cols=6)
raci_headers = ['Süreç Adımı', 'Talep Eden Birim', 'Satınalma Uzmanı', 'Satınalma Direktörü', 'Mali İşler GMY', 'İç Denetim']
for i, h in enumerate(raci_headers):
    table_raci.rows[0].cells[i].text = h
raci_rows = [
    ('Satınalma Talebi Oluşturma', 'Sorumlu (R)', 'Danışılan (C)', 'Bilgilendirilen (I)', 'Onaylayan (A)', 'Bilgilendirilen (I)'),
    ('Teklif Toplama ve TKT Hazırlama', 'Danışılan (C)', 'Sorumlu (R)', 'Onaylayan (A)', 'Bilgilendirilen (I)', 'İnceleyen (I)'),
    ('Satıcı Akreditasyonu ve Risk Değerleme', 'Bilgilendirilen (I)', 'Sorumlu (R)', 'Onaylayan (A)', 'Onaylayan (A)', 'Denetleyen (A)'),
    ('Mal Kabul ve Kalite Doğrulama', 'Bilgilendirilen (I)', 'Bilgilendirilen (I)', 'Bilgilendirilen (I)', 'Bilgilendirilen (I)', 'Denetleyen (I)'),
    ('Fatura Girişi ve 3 lü Eşleştirme', 'Bilgilendirilen (I)', 'Bilgilendirilen (I)', 'Bilgilendirilen (I)', 'Sorumlu (R)', 'Denetleyen (A)'),
    ('Dönem Sonu Fiili Stok Sayımı', 'Katılımcı (C)', 'Gözlemci (I)', 'Bilgilendirilen (I)', 'Sorumlu (R)', 'Bağımsız Gözlemci (A)')
]
for row in raci_rows:
    r_cells = table_raci.add_row().cells
    for j, val in enumerate(row):
        r_cells[j].text = val

# Bölüm 4: Onay Limitleri ve İstisnalar
doc_tedarik.add_heading('4. SATINALMA TALEBİ VE ONAY YETKİ MATRİSİ', level=2)
doc_tedarik.add_paragraph(
    '4.1. Tüm satınalma talepleri SAP MM modülü üzerinden bütçe kontrolü (Availability Control) yapılarak başlatılır. '
    'Bütçesi bulunmayan veya bütçe aşımına yol açan talepler sistem tarafından otomatik bloke edilir.'
)
doc_tedarik.add_paragraph('4.2. Satınalma Onay Limit Tablosu aşağıdaki gibidir:')

table_auth = doc_tedarik.add_table(rows=1, cols=4)
auth_headers = ['Tutar Aralığı (TL)', 'Yetkili Onay Makamı', 'İmza Sayısı', 'Kontrol Kodu']
for i, h in enumerate(auth_headers):
    table_auth.rows[0].cells[i].text = h
auth_data = [
    ('0 - 50.000 TL', 'Birim Müdürü', 'Tek İmza', 'C1-A'),
    ('50.001 - 250.000 TL', 'Departman Direktörü + Bütçe Müdürü', 'Çift İmza', 'C1-B'),
    ('250.001 - 1.000.000 TL', 'Satınalma Direktörü + Mali İşler GMY (CFO)', 'Çift İmza', 'C1-C'),
    ('1.000.001 TL ve Üzeri', 'Genel Müdür (CEO) + Yönetim Kurulu Murahhas Üyesi', 'Müşterek Çift İmza', 'C1-D')
]
for row in auth_data:
    r_cells = table_auth.add_row().cells
    for j, val in enumerate(row):
        r_cells[j].text = val

doc_tedarik.add_paragraph(
    '4.3. [KRİTİK ZAYIFLIK TESPİTİ - W1]: Operasyonel aciliyet veya fabrika üretim duruşu riski gerekçe gösterilerek, '
    'Satınalma Direktörü tarafından SAP sisteminde "ACİL_ALIM_OVERRIDE" yetki koduyla onay limitleri 5 katına kadar artırılabilmekte '
    've bu istisnai işlemler Mali İşler GMY veya CEO onayına sunulmadan doğrudan siparişe (PO) dönüştürülmektedir. 2025 yılı boyunca bu yetki '
    'toplam 47 işlemde kullanılmış ve 38.200.000 TL tutarında denetimsiz alım gerçekleştirilmiştir.'
)

# Bölüm 5: İhale ve Satıcı Seçimi
doc_tedarik.add_heading('5. İHALE SÜRECİ, TEKLİF TOPLAMA VE SATICI AKREDİTASYONU', level=2)
doc_tedarik.add_paragraph(
    '5.1. 250.000 TL üzerindeki tüm alımlarda İhale Komisyonu gözetiminde en az 3 bağımsız ve akredite satıcıdan kapalı teklif toplanması zorunludur. '
    'Teklifler eş zamanlı olarak dijital satınalma portalına yüklenir ve komisyon huzurunda açılır (C2).'
)
doc_tedarik.add_paragraph(
    '5.2. Satıcı seçiminde mali güç, ISO kalite belgeleri, referanslar ve ilişkili taraf durumu (Bağımsızlık Beyannamesi) sorgulanır.'
)
doc_tedarik.add_paragraph(
    '5.3. [KRİTİK ZAYIFLIK TESPİTİ - W2]: Bilişim, yazılım lisansı ve danışmanlık hizmeti alımlarında "Tek Satıcı İstisnası" (Single Source Waiver) '
    'kuralı suistimal edilerek, piyasada ikamesi bulunan genel danışmanlık hizmetleri için dahi tek satıcıdan e-posta ile teklif alınmış; '
    'satıcı akreditasyon formları ve MASAK şüpheli işlem kontrolleri sözleşme imzalandıktan aylar sonra geriye dönük tamamlanmıştır.'
)

# Bölüm 6: Mal Kabul ve 3'lü Eşleştirme
doc_tedarik.add_heading('6. MAL KABUL, İRSALİYE KONTROLÜ VE 3 LÜ EŞLEŞTİRME', level=2)
doc_tedarik.add_paragraph(
    '6.1. Fabrika ve depolara ulaşan tüm fiziksel mallar ambar personeli ve kalite kontrol mühendisi tarafından kontrol edilerek irsaliye teslim alınır. '
    'Sistemde 103 (Kalite Bloke Girişi) yapılır; laboratuvar kalite onayı alındıktan sonra 101 (Kullanıma Hazır Giriş - GRN) kaydı atılır (C3).'
)
doc_tedarik.add_paragraph(
    '6.2. Fatura geldiğinde SAP sistemi PO, GRN ve Faturayı otomatik eşleştirir. Miktarda %0 tolerans, birim fiyatta maksimum %1.5 tolerans tanımlıdır.'
)
doc_tedarik.add_paragraph(
    '6.3. [KRİTİK ZAYIFLIK TESPİTİ - W3]: Hizmet alımlarında (Danışmanlık, Reklam, Hukuk) fiziksel bir mal teslimatı olmadığı için sistemde "Hizmet Giriş Belgesi (SES - Service Entry Sheet)" '
    'düzenlenmesi zorunlu olmasına rağmen; Satınalma Uzmanlarının amir onayı olmaksızın SES kayıtlarını tek taraflı onayladığı ve hak ediş raporu olmaksızın faturaların muhasebeye sevk edildiği tespit edilmiştir.'
)

# Bölüm 7: Konsinye Stok
doc_tedarik.add_heading('7. KONSİNYE STOK YÖNETİMİ VE TÜKETİM MUTABAKATLARI', level=2)
doc_tedarik.add_paragraph(
    '7.1. Konsinye stoklar özel tanımlı "K" ambarında izlenir. Mülkiyet tedarikçidedir; şirket yalnızca üretime çekip tükettiği miktardan sorumludur.'
)
doc_tedarik.add_paragraph(
    '7.2. Her ayın son iş günü konsinye stok ambarında tedarikçi temsilcisi ve şirket depo şefi nezaretinde fiili sayım yapılır ve karşılıklı ıslak imzalı mutabakat tutanağı düzenlenir (C4).'
)
doc_tedarik.add_paragraph(
    '7.3. [KRİTİK ZAYIFLIK TESPİTİ - W4]: Son 8 aydır konsinye depolarda fiziksel sayım yapılmamış; tedarikçilerden gelen aylık faturalar sistemdeki kuramsal tüketim verileriyle '
    'fiili ambar mevcudu karşılaştırılmaksızın ve tedarikçi mutabakat tutanağı aranmaksızın doğrudan onaylanarak ödenmiştir.'
)

# Bölüm 8: Hurda Satışları, Fireler ve Sayım Farkları
doc_tedarik.add_heading('8. HURDA SATIŞLARI, ÜRETİM FİRELERİ VE DÖNEM SONU ENVANTER SAYIMLARI', level=2)
doc_tedarik.add_paragraph(
    '8.1. Üretim sürecinde ortaya çıkan metal, plastik ve kimyasal hurdaların satışı Hurda Komisyonu (Fabrika Müdürü, Güvenlik Şefi, Mali İşler Uzmanı) '
    'gözetiminde kantar tartım fişi düzenlenerek yapılır. Bedel şirket resmi banka hesabına EFT/Havale ile peşin ödenmeden hiçbir aracın tesis dışına çıkışına izin verilmez (C5).'
)
doc_tedarik.add_paragraph(
    '8.2. Yıllık fiili stok sayımları bağımsız sayım komisyonu tarafından yılda bir kez kör sayım (blind count) yöntemiyle gerçekleştirilir. '
    'Sayım sonuçları sistem kayıtlarıyla karşılaştırılarak Sayım Farkı Tutanağı düzenlenir. Şirket Muhasebe Yönetmeliği Madde 9.4 uyarınca; '
    '1.000.000 TL yi aşan her türlü envanter açığı veya fazlası için Kök Neden Soruşturması açılması ve Yönetim Kurulu Kararı alınması şarttır.'
)
doc_tedarik.add_paragraph(
    '8.3. [AĞIR USULSÜZLÜK VE SUİİSTİMAL TESPİTİ - W5]: 2025 yılı dönem sonu sayımında ana hammadde ambarında 8.400.000 TL tutarında stok açığı tespit edilmiş; '
    'ancak Fabrika Müdürü ve Satınalma Direktörü Yönetim Kuruluna bilgi vermeksizin ve soruşturma açmaksızın bu tutarı 689 hesaba gider yazarak kapatmıştır. '
    'Ayrıca güvenlik kamerası kayıtlarında gece vardiyasında 3 kamyon hurda malzemenin (yaklaşık 3.200.000 TL) kantarsız ve faturasız olarak tesis dışına çıkarıldığı belirlenmiştir.'
)

doc_tedarik.save(os.path.join(enterprise_dir, 'kuresel_tedarik_zinciri_ve_stok_proseduru_kompleks.docx'))
doc_tedarik.save(os.path.join(massive_dir, 'kuresel_tedarik_zinciri_ve_stok_proseduru_kompleks.docx'))

# 2. BÜYÜK DOKÜMAN 2: Hazine, Türev Ürünler ve Swift Prosedürü (Ultra Detaylı)
doc_hazine = docx.Document()
doc_hazine.add_heading('MEGA HOLDİNG A.Ş. — HAZİNE, NAKİT YÖNETİMİ, TÜREV İŞLEMLER VE SWIFT TRANSFER PROSEDÜRÜ (SOP-TRS-2026-V6)', level=1)
doc_hazine.add_paragraph('Belge Kodu: SOP-TRS-2026-V6 | Revizyon No: 06 | Yürürlük Tarihi: 01.01.2026 | Gizlilik: ÇOK GİZLİ (Finansal Yönetişim)')

doc_hazine.add_heading('1. AMAÇ VE TEMEL İLKELER', level=2)
doc_hazine.add_paragraph(
    'Bu prosedür; Holding ve bağlı 14 iştirakinin banka hesapları, nakit akış optimizasyonu, döviz pozisyon riski, faiz/kur türev işlemleri (Swap, Forward, Opsiyon), '
    'kredi limitleri ve uluslararası Swift transferlerinin icra edilmesinde uygulanacak iç kontrol standartlarını, yetki limitlerini ve eskalasyon kurallarını belirler.'
)

doc_hazine.add_heading('2. BANKA HESAPLARI VE YETKİLİ İMZA SİRKÜLERİ KADEMELERİ', level=2)
doc_hazine.add_paragraph(
    '2.1. Şirket adına yeni bir banka hesabı açılması veya kapatılması yalnızca Mali İşler GMY (CFO) ve Hazine Direktörü müşterek imzasıyla gerçekleştirilir. '
    'Tüm hesaplar merkezi Hazine Yönetim Sistemine (TMS) entegre olmak zorundadır (C1).'
)

doc_hazine.add_heading('3. GÜNLÜK FON TRANSFERİ VE SWIFT İŞLEMLERİ (DUAL CONTROL - 4-EYES PRINCIPLE)', level=2)
doc_hazine.add_paragraph(
    '3.1. Bankalar üzerinden yapılacak tüm EFT, Havale ve Uluslararası SWIFT transferlerinde Görevler Ayrılığı (SoD) ve 4-Göz İlkesi zorunludur: '
    'İşlemi hazırlayan (Maker) kullanıcı ile banka sisteminde onaylayan (Checker/Releaser) kullanıcı kesinlikle farklı kişiler olmak zorundadır (C2).'
)
doc_hazine.add_paragraph('3.2. Hazine Onay Limit Tablosu:')

table_haz = doc_hazine.add_table(rows=1, cols=4)
haz_headers = ['Tutar Aralığı (USD/EUR Muadili)', 'Birinci Onaycı (Maker)', 'İkinci Onaycı (Checker)', 'Nihai Onay Makamı']
for i, h in enumerate(haz_headers):
    table_haz.rows[0].cells[i].text = h
haz_data = [
    ('0 - 100.000 USD', 'Hazine Uzmanı', 'Hazine Müdürü', 'Hazine Müdürü'),
    ('100.001 - 500.000 USD', 'Hazine Uzmanı', 'Hazine Müdürü', 'Hazine Direktörü'),
    ('500.001 - 2.000.000 USD', 'Hazine Müdürü', 'Hazine Direktörü', 'Mali İşler GMY (CFO)'),
    ('2.000.001 USD ve Üzeri', 'Hazine Direktörü', 'Mali İşler GMY (CFO)', 'Yönetim Kurulu Finans Komitesi')
]
for row in haz_data:
    r_cells = table_haz.add_row().cells
    for j, val in enumerate(row):
        r_cells[j].text = val

doc_hazine.add_heading('4. TÜREV ÜRÜNLER, PİYASA RİSKLERİ VE SPEKÜLATİF İŞLEM YASAĞI', level=2)
doc_hazine.add_paragraph(
    '4.1. Şirket prensip olarak spekülatif veya arbitraj amaçlı hiçbir türev işlem yapamaz. Türev işlemler (Forward, Vadeli Döviz, Faiz Swapı, Opsiyon) '
    'yalnızca kanıtlanabilir ticari ithalat/ihracat döviz riskini hedge etmek (Riskten Korunma) amacıyla yapılabilir (C3).'
)
doc_hazine.add_paragraph(
    '4.2. [KRİTİK PİYASA VE KONTROL ZAYIFLIĞI - W1]: Prosedürün 4.3 maddesinde yer alan "Piyasa aşırı oynaklığı durumunda Hazine Uzmanı sözlü talimatla '
    '2.000.000 USD limitli spot döviz işlemi yapabilir ve banka teyit mektupları ertesi gün sisteme girilir" istisnası, denetim döneminde kötüye kullanılmış; '
    'Hazine Uzmanının mesai saatleri dışında 6 kez tek başına toplam 11.500.000 USD spot döviz alım-satımı yaptığı ve şirketin 420.000 USD kur zararına uğradığı saptanmıştır.'
)

doc_hazine.add_heading('5. OFFSHORE VE YÜKSEK RİSKLİ YARGI BÖLGESİ TRANSFERLERİ (AML & MASAK UYUMU)', level=2)
doc_hazine.add_paragraph(
    '5.1. FATF kara listesinde, gri listede veya vergi cenneti (BVI, Cayman, Kıbrıs, Panama vb.) sayılan ülkelerdeki banka hesaplarına yapılacak '
    '50.000 USD üzerindeki tüm ödemelerde MASAK ve Uyum Departmanı (Compliance) yazılı onayının alınması zorunludur (C4).'
)
doc_hazine.add_paragraph(
    '5.2. [KRİTİK UYUM İHLALİ - W2]: 2026 yılı 2. çeyreğinde BVI ve Kıbrıs merkezli 3 danışmanlık firmasına toplam 3.250.000 USD tutarında SWIFT transferi '
    'yapılmış; Uyum Müdürü onayı alınmamış ve faturaların arkasında somut danışmanlık teslimat raporu bulunamamıştır.'
)

doc_hazine.save(os.path.join(enterprise_dir, 'hazine_ve_turev_islemler_proseduru.docx'))
doc_hazine.save(os.path.join(massive_dir, 'hazine_ve_turev_islemler_proseduru.docx'))

print('Devasa ve ultra detaylı Word süreç prosedürleri başarıyla üretildi.')
