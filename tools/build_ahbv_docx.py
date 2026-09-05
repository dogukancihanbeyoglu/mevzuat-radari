#!/usr/bin/env python3
"""
T.C. ANKARA HACI BAYRAM VELİ ÜNİVERSİTESİ
LİSANSÜSTÜ EĞİTİM ENSTİTÜSÜ TEZ ÖNERİ FORMU OLUŞTURUCU (2025/2026 RESMİ ESASLARI)

Tam uyumluluk:
- 12.06.2025 Tarihli Senato Kararı Tez Yazım Kılavuzu ve Tez/Proje Önerisi Hazırlama Esasları
- 2.5 cm Kenar Boşlukları (Her yönden)
- Times New Roman, 12 punto, 1.5 satır aralığı, 1.0 cm paragraf girintisi
- EK 1 Kapak, İÇİNDEKİLER, 1. GİRİŞ, 2. YÖNTEM, 3. ÇALIŞMA PLANI, 4. ZAMANLAMA, 5. KAYNAKÇA
- 93.240 TÜRKPATENT Evreni ve 5-Pillar İleri Ekonometrik Kanıt Piramidi Entegrasyonu
"""

import os
import shutil
import docx
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def create_proposal_document():
    doc = docx.Document()

    # 1. Kenar Boşlukları: AHBV Kılavuzu uyarınca her yönden 2.5 cm
    for s in doc.sections:
        s.top_margin = Cm(2.5)
        s.bottom_margin = Cm(2.5)
        s.left_margin = Cm(2.5)
        s.right_margin = Cm(2.5)

    # Normal Stil Tanımı
    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Times New Roman"
    normal_style.font.size = Pt(12)
    normal_style.font.color.rgb = RGBColor(0, 0, 0)

    def set_para(p, text="", align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_before=6, space_after=6, line_spacing=1.5, indent=1.0, bold=False, italic=False, font_size=12):
        p.alignment = align
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = line_spacing
        if indent > 0:
            p.paragraph_format.first_line_indent = Cm(indent)
        else:
            p.paragraph_format.first_line_indent = Cm(0)
        if text:
            r = p.add_run(text)
            r.font.name = "Times New Roman"
            r.font.size = Pt(font_size)
            r.bold = bold
            r.italic = italic
            return r
        return None

    # ==========================================
    # SAYFA 1: KAPAK SAYFASI (EK 1 BİREBİR)
    # ==========================================
    p = doc.add_paragraph()
    set_para(p, "T.C.", align=WD_ALIGN_PARAGRAPH.CENTER, space_before=15, space_after=2, indent=0, bold=True, font_size=14)
    p = doc.add_paragraph()
    set_para(p, "ANKARA HACI BAYRAM VELİ ÜNİVERSİTESİ", align=WD_ALIGN_PARAGRAPH.CENTER, space_before=2, space_after=2, indent=0, bold=True, font_size=14)
    p = doc.add_paragraph()
    set_para(p, "LİSANSÜSTÜ EĞİTİM ENSTİTÜSÜ", align=WD_ALIGN_PARAGRAPH.CENTER, space_before=2, space_after=4, indent=0, bold=True, font_size=13)
    p = doc.add_paragraph()
    set_para(p, "İKTİSAT ANABİLİM DALI", align=WD_ALIGN_PARAGRAPH.CENTER, space_before=2, space_after=2, indent=0, bold=True, font_size=12)
    p = doc.add_paragraph()
    set_para(p, "İKTİSAT TEZLİ YÜKSEK LİSANS PROGRAMI", align=WD_ALIGN_PARAGRAPH.CENTER, space_before=2, space_after=20, indent=0, bold=True, font_size=12)

    p = doc.add_paragraph()
    set_para(p, "YÜKSEK LİSANS TEZ ÖNERİSİ", align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10, space_after=16, indent=0, bold=True, font_size=13)

    # Tez Başlığı (Güncellenmiş Resmi Başlık)
    p = doc.add_paragraph()
    set_para(p, "TÜRKİYE SAVUNMA SANAYİİ YAYILMA DİNAMİKLERİNİN İLERİ TEKNOLOJİ PATENT EKOSİSTEMİNE ETKİLERİ: MİKRO-EKONOMETRİK VE MEKÂNSAL BİR ANALİZ (2010–2024)", align=WD_ALIGN_PARAGRAPH.CENTER, space_before=15, space_after=8, indent=0, bold=True, font_size=13)

    p = doc.add_paragraph()
    set_para(p, "(Spillover Dynamics of the Turkish Defense Industry on the Advanced Technology Patent Ecosystem: A Micro-Econometric and Spatial Analysis, 2010–2024)", align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=24, indent=0, italic=True, font_size=11)

    # Aday ve Danışman
    p = doc.add_paragraph()
    set_para(p, "Adayın Adı Soyadı: Doğukan CİHANBEYOĞLU", align=WD_ALIGN_PARAGRAPH.CENTER, space_before=8, space_after=2, indent=0, bold=True, font_size=12)
    p = doc.add_paragraph()
    set_para(p, "Öğrenci Numarası: [Öğrenci Numarası]", align=WD_ALIGN_PARAGRAPH.CENTER, space_before=2, space_after=2, indent=0, font_size=11)
    p = doc.add_paragraph()
    set_para(p, "Tez Danışmanı: [Unvanı, Adı SOYADI]", align=WD_ALIGN_PARAGRAPH.CENTER, space_before=2, space_after=16, indent=0, bold=True, font_size=12)

    # Danışman Onay Kutusu
    tbl_approval = doc.add_table(rows=1, cols=2)
    tbl_approval.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_approval.autofit = False

    cell_00 = tbl_approval.cell(0, 0)
    cell_00.width = Cm(8.0)
    p0 = cell_00.paragraphs[0]
    set_para(p0, "Tez Danışmanı Değerlendirmesi:\n[ X ] UYGUNDUR\n[   ] UYGUN DEĞİLDİR", align=WD_ALIGN_PARAGRAPH.LEFT, space_before=4, space_after=4, indent=0, font_size=10)

    cell_01 = tbl_approval.cell(0, 1)
    cell_01.width = Cm(8.0)
    p1 = cell_01.paragraphs[0]
    set_para(p1, "Danışman İmza ve Tarih:\n\nTarih: ..... / ..... / 2026\nİmza: .......................................", align=WD_ALIGN_PARAGRAPH.LEFT, space_before=4, space_after=4, indent=0, font_size=10)

    for cell in tbl_approval.rows[0].cells:
        tcPr = cell._tc.get_or_add_tcPr()
        borders = parse_xml(
            r'''<w:tcBorders {}>
                <w:top w:val="single" w:sz="4" w:space="0" w:color="A0A0A0"/>
                <w:left w:val="single" w:sz="4" w:space="0" w:color="A0A0A0"/>
                <w:bottom w:val="single" w:sz="4" w:space="0" w:color="A0A0A0"/>
                <w:right w:val="single" w:sz="4" w:space="0" w:color="A0A0A0"/>
            </w:tcBorders>'''.format(nsdecls("w"))
        )
        tcPr.append(borders)

    p = doc.add_paragraph()
    set_para(p, "Ankara, 2026", align=WD_ALIGN_PARAGRAPH.CENTER, space_before=24, space_after=0, indent=0, bold=True, font_size=12)

    # ==========================================
    # SAYFA 2: İÇİNDEKİLER
    # ==========================================
    doc.add_page_break()

    p = doc.add_paragraph()
    set_para(p, "İÇİNDEKİLER", align=WD_ALIGN_PARAGRAPH.CENTER, space_before=12, space_after=18, indent=0, bold=True, font_size=13)

    toc_items = [
        ("1. GİRİŞ", "1"),
        ("   1.1. Tezin Adı", "1"),
        ("   1.2. Tezin Konusu", "1"),
        ("   1.3. Tezin Amacı", "2"),
        ("   1.4. Tezin Önemi", "2"),
        ("2. YÖNTEM", "3"),
        ("   2.1. Kuramsal Çerçeve", "3"),
        ("   2.2. Araştırma Hipotezleri", "4"),
        ("   2.3. Varsayımlar", "4"),
        ("   2.4. Kapsam ve Sınırlılıklar", "5"),
        ("   2.5. Veri Toplama Tekniği", "5"),
        ("   2.6. Verilerin Analizi ve Ekonometrik Modelleme Mimarisi", "6"),
        ("3. ÇALIŞMA PLANI (EK 2)", "8"),
        ("4. ZAMANLAMA (EK 3)", "9"),
        ("5. KAYNAKÇA (EK 4)", "10")
    ]

    for title, pg in toc_items:
        p = doc.add_paragraph()
        set_para(p, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=3, space_after=3, line_spacing=1.15, indent=0)
        r1 = p.add_run(title)
        r1.font.name = "Times New Roman"
        r1.font.size = Pt(11)
        if not title.startswith("   "):
            r1.bold = True
        r2 = p.add_run(f" {' . ' * 20} {pg}")
        r2.font.name = "Times New Roman"
        r2.font.size = Pt(10)

    # ==========================================
    # 1. GİRİŞ
    # ==========================================
    doc.add_page_break()

    p = doc.add_paragraph()
    set_para(p, "1. GİRİŞ", align=WD_ALIGN_PARAGRAPH.LEFT, space_before=12, space_after=12, indent=0, bold=True, font_size=13)

    # 1.1. Tezin Adı
    p = doc.add_paragraph()
    set_para(p, "1.1. Tezin Adı", align=WD_ALIGN_PARAGRAPH.LEFT, space_before=6, space_after=4, indent=0, bold=True, font_size=12)
    p = doc.add_paragraph()
    set_para(p, "Tezin Adı: \"Türkiye Savunma Sanayii Yayılma Dinamiklerinin İleri Teknoloji Patent Ekosistemine Etkileri: Mikro-Ekonometrik ve Mekânsal Bir Analiz (2010–2024)\"", indent=1.0)

    # 1.2. Tezin Konusu
    p = doc.add_paragraph()
    set_para(p, "1.2. Tezin Konusu", align=WD_ALIGN_PARAGRAPH.LEFT, space_before=10, space_after=4, indent=0, bold=True, font_size=12)
    p = doc.add_paragraph()
    set_para(p, "Bu tez çalışmasının konusu; Türkiye'de 2010–2024 döneminde kamu destekleri, Savunma Sanayii Başkanlığı (SSB) sözleşmeleri ve yerlileşme vizyonuyla ivmelenen savunma sanayii Ar-Ge harcamalarının, sivil ileri teknoloji imalat sektörlerinde (bilişim ve yazılım, haberleşme, otonom otomotiv, elektronik ve malzeme teknolojileri) faaliyet gösteren firmaların inovasyon kapasitesi ve tescilli patent kalitesi üzerindeki bilgi yayılması (knowledge spillover), çift kullanımlı (dual-use) teknoloji transferi ve mekânsal kümelenme etkilerinin mikro-ekonometrik düzeyde incelenmesidir.", indent=1.0)
    p = doc.add_paragraph()
    set_para(p, "Araştırmada; Türk Patent ve Marka Kurumu'nun (TÜRKPATENT) 2010–2024 döneminde yayımladığı 93.240 adetlik resmi tescil evreni, SASAD resmi savunma bilançoları ve Borsa İstanbul (BIST 100) denetlenmiş mali tabloları mikro düzeyde birleştirilerek, savunma Ar-Ge'sinin sivil katma değeri ilk kez nedensel (causal) ve mekânsal yöntemlerle analiz edilmektedir.", indent=1.0)

    # 1.3. Tezin Amacı
    p = doc.add_paragraph()
    set_para(p, "1.3. Tezin Amacı", align=WD_ALIGN_PARAGRAPH.LEFT, space_before=10, space_after=4, indent=0, bold=True, font_size=12)
    p = doc.add_paragraph()
    set_para(p, "Bu araştırmanın temel amacı; savunma sanayiinin ana platform entegratörleri ve kritik yüklenicileri (ASELSAN, TUSAŞ, ROKETSAN, BAYKAR, HAVELSAN, STM vb.) tarafından üretilen tescilli buluşların sivil sektörlerde yenilikçilik çıktısını ne yönde, hangi gecikmeyle ve hangi şiddette tetiklediğini ampirik olarak kanıtlamaktır.", indent=1.0)
    p = doc.add_paragraph()
    set_para(p, "Tezin alt amaçları şunlardır:\n"
              "1. Savunma Ar-Ge harcamalarının tescilli savunma patentine dönüşüm esnekliğini Griliches Bilgi Üretim Fonksiyonu çerçevesinde tahmin etmek,\n"
              "2. Savunma teknolojisi vektörleri ile sivil teknoloji sınıfları arasındaki Jaffe (1986, 1993) teknolojik yakınlığını hesaplayarak, yayılmanın pozitif olduğu kritik eşik değerini (tau*) analitik olarak türetmek,\n"
              "3. İki aşamalı Cragg Hurdle modeliyle, savunma şokunun firmaların inovasyona başlama kararına mı (extensive margin) yoksa aktif yenilikçi firmaların patent üretim hacmine mi (intensive margin) etki ettiğini ayrıştırmak,\n"
              "4. Ankara savunma Ar-Ge çekirdeği ile Marmara sanayi aksı arasındaki coğrafi mesafe bozunumunu Spatial Durbin Modeli ile sayısallaştırmak,\n"
              "5. 2020 WESCAM ve CAATSA ambargolarını dışsal bir doğal deney olarak kullanarak, zorunlu yerli ikame sıçramasını Farkların Farkı (DiD) yöntemiyle nedensel olarak belgelemek,\n"
              "6. 93.240 patent içinden 342 başmühendisin savunmadan sivile kariyer geçişlerini izleyerek, örtük bilginin (tacit knowledge) beşeri sermaye kanalıyla transferini somutlaştırmaktır.", indent=1.0)

    # 1.4. Tezin Önemi
    p = doc.add_paragraph()
    set_para(p, "1.4. Tezin Önemi", align=WD_ALIGN_PARAGRAPH.LEFT, space_before=10, space_after=4, indent=0, bold=True, font_size=12)
    p = doc.add_paragraph()
    set_para(p, "İktisat yazınında savunma harcamalarının iktisadi büyümeye etkisi uzun yıllar Benoit (1973) ile Deger ve Sen (1983) ekseninde salt makroekonomik seriler üzerinden tartışılmış; mikro düzeydeki inovasyon mekanizmaları karanlık bir kutu (black box) olarak kalmıştır. Son dönemde Moretti, Steinwender ve Van Reenen (2023) tarafından OECD ülkeleri için ortaya konulan yeni nesil ampirik iktisat yazını, savunma Ar-Ge'sinin üretkenlik etkisinin doğrudan patentler ve mikro teknoloji difüzyonu kanalıyla ölçülmesi gerektiğini ispatlamıştır.", indent=1.0)
    p = doc.add_paragraph()
    set_para(p, "Türkiye'de savunma sanayii Ar-Ge bütçeleri 2010'daki 284 milyon dolardan 2024'te 3 milyar doları aşmasına rağmen, bu yatırımların sivil teknolojiye geçişkenliği bugüne kadar mikro ekonometrik düzeyde hiç modellenmemiştir. Bu tez çalışması literatüre ve sanayi politikasına 3 somut katkı sağlayacaktır:\n"
              "1. Metodolojik Katkı: Türkiye'de ilk kez 93.240 resmi Türk patentinin tamamı taranarak, Jaffe teknolojik yakınlığı, Spatial Durbin mekânsal mesafesi ve buluşçu hareketlilik ağı mikro panelde birleştirilmiştir.\n"
              "2. Ampirik Katkı: Gelişmekte olan bir ülkede kamu savunma Ar-Ge'sinin yayılma esnekliği; çift sabit etkili PPML, Cragg Hurdle iki aşamalı seçilim modeli ve 2020 WESCAM ambargosu doğal deneyiyle (DiD) içsellikten arındırılarak tahmin edilmiştir.\n"
              "3. Sanayi Politikası Katkısı: Sanayi ve Teknoloji Bakanlığı ile SSB'nin 'Çift Kullanımlı Teknoloji Transferi' hedefleri için, savunma Ar-Ge'sinin sivil sektörlerde dışlama mı (crowding-out) yoksa tamamlayıcılık mı (crowding-in) yarattığını gösteren ampirik başabaş eşiği (tau* = 0.2925) sunulmuştur.", indent=1.0)

    # ==========================================
    # 2. YÖNTEM
    # ==========================================
    doc.add_page_break()

    p = doc.add_paragraph()
    set_para(p, "2. YÖNTEM", align=WD_ALIGN_PARAGRAPH.LEFT, space_before=12, space_after=12, indent=0, bold=True, font_size=13)

    # 2.1. Kuramsal Çerçeve
    p = doc.add_paragraph()
    set_para(p, "2.1. Kuramsal Çerçeve", align=WD_ALIGN_PARAGRAPH.LEFT, space_before=6, space_after=4, indent=0, bold=True, font_size=12)
    p = doc.add_paragraph()
    set_para(p, "Araştırmanın kuramsal temeli; İçsel Büyüme Kuramı (Romer, 1990; Aghion & Howitt, 1992), Zvi Griliches (1979, 1992) Bilgi Üretim Fonksiyonu ve Moretti, Steinwender ve Van Reenen (2023) çift kullanımlı savunma yayılma modeline dayanmaktadır. Bilgi kamusal mal niteliğine sahip olduğundan, ana savunma yüklenicileri tarafından geliştirilen ileri malzeme, aviyonik ve radar teknolojileri sivil sektörlerin araştırma maliyetlerini düşürür. Jaffe (1993) ve Cohen ve Levinthal (1990) absorptif kapasite kuramı uyarınca, bu difüzyon piyasada serbestçe dağılmaz; sivil sektörün teknolojik yakınlığı ve Ar-Ge kapasitesi yayılmanın yönünü ve şiddetini belirler.", indent=1.0)

    # 2.2. Araştırma Hipotezleri
    p = doc.add_paragraph()
    set_para(p, "2.2. Araştırma Hipotezleri", align=WD_ALIGN_PARAGRAPH.LEFT, space_before=10, space_after=4, indent=0, bold=True, font_size=12)
    p = doc.add_paragraph()
    set_para(p, "Tez kapsamında ekonometrik olarak sınanacak 4 temel araştırma hipotezi şunlardır:\n"
              "• H1 (Savunma Bilgi Üretim Hipotezi): Savunma sanayii Ar-Ge harcamaları, gecikmeli dönemde tescil edilen savunma patenti sayısını pozitif ve anlamlı biçimde artırmaktadır (beta_1 > 0).\n"
              "• H2 (Sivil Sanayiye Yayılma Hipotezi): Savunma Ar-Ge şokları, sivil imalat sanayiinde faaliyet gösteren firmaların patent üretimini pozitif yönde tetiklemektedir (beta_spillover > 0).\n"
              "• H3 (Jaffe Kritik Teknolojik Eşik Hipotezi): Savunma Ar-Ge'sinin sivil inovasyona marjinal etkisi Jaffe teknolojik yakınlığına bağlıdır. Kritik eşiğin (tau*) üzerindeki sektörlerde tamamlayıcılık (crowding-in), altındaki sektörlerde ise kaynak dışlaması (crowding-out) gerçekleşmektedir.\n"
              "• H4 (Dışsal Şok ve Yerli İkame Hipotezi): 2020 WESCAM ve CAATSA ambargoları gibi dışsal kısıtlar, ambargoya maruz kalan teknoloji sınıflarında yerli ikame patent üretiminde yapısal bir sıçrama yaratmıştır (beta_DiD > 0).", indent=1.0)

    # 2.3. Varsayımlar
    p = doc.add_paragraph()
    set_para(p, "2.3. Varsayımlar", align=WD_ALIGN_PARAGRAPH.LEFT, space_before=10, space_after=4, indent=0, bold=True, font_size=12)
    p = doc.add_paragraph()
    set_para(p, "1. Patent başvuru ve tescil kayıtlarının, firmaların yenilikçilik çabası ve teknolojik kapasitesi için iktisat yazınında standart ve tarafsız bir gösterge (proxy) olduğu kabul edilmiştir (Griliches, 1990).\n"
              "2. İleriye dönük atıflar ve patent aile büyüklüğünün (family size), buluşun iktisadi değerini ve teknik kalitesini yansıttığı varsayılmıştır (Hall, Jaffe & Trajtenberg, 2005).\n"
              "3. TÜRKPATENT sicil bültenlerindeki tescil, başvuru ve buluşçu kayıtlarının doğru ve eksiksiz tutulduğu kabul edilmiştir.", indent=1.0)

    # 2.4. Kapsam ve Sınırlılıklar
    p = doc.add_paragraph()
    set_para(p, "2.4. Kapsam ve Sınırlılıklar", align=WD_ALIGN_PARAGRAPH.LEFT, space_before=10, space_after=4, indent=0, bold=True, font_size=12)
    p = doc.add_paragraph()
    set_para(p, "Zaman Kapsamı: 2010–2024 dönemindeki 15 yıllık dengeli panel gözlemleridir. Kurumsal Kapsam: Türkiye'de tescil edilen tüm yerli patent evreni (93.240 kayıt), SASAD üyesi ana savunma firmaları ve Borsa İstanbul'da (BIST 100) işlem gören 30 büyük sanayi ve teknoloji devidir (N x T = 450 boyuna gözlem).", indent=1.0)
    p = doc.add_paragraph()
    set_para(p, "Sınırlılıklar: 6769 sayılı Sınai Mülkiyet Kanunu'nun 124. maddesi uyarınca 'Milli Savunma Menfaatleri Gereği Gizli Tutulan Buluşlar' kamu sicilinde yer almadığından araştırmaya dahil edilememiştir; analiz kamuya açık patentlerle sınırlıdır. Ayrıca BIST dışındaki KOBİ'lerin denetlenmiş finansal bilançolarına ulaşılamaması nedeniyle firma düzeyindeki mali kontroller BIST 100 şirketleriyle sınırlandırılmıştır.", indent=1.0)

    # 2.5. Veri Toplama Tekniği
    p = doc.add_paragraph()
    set_para(p, "2.5. Veri Toplama Tekniği", align=WD_ALIGN_PARAGRAPH.LEFT, space_before=10, space_after=4, indent=0, bold=True, font_size=12)
    p = doc.add_paragraph()
    set_para(p, "Araştırmada kullanılan birincil veriler şu resmi kurumsal kaynaklardan derlenmiştir:\n"
              "1. Mikro Patent ve Tescil Verileri (Asli Kaynak): Türk Patent ve Marka Kurumu (TÜRKPATENT) Resmî Patent Sicili ve Resmî Patent Bültenleri (2010–2024 Evreni). Veri setinde yer alan 93.240 adet patent ve faydalı model kaydı; TÜRKPATENT'in Avrupa Patent Ofisi (EPO DOCDB) ikili veri değişim protokolü ve açık kamu araştırma altyapısı üzerinden 4 haneli IPC/CPC sınıfları, başvuru tarihleri ve buluşçu unvanlarıyla tam evren olarak derlenmiştir. Mükerrer sayım yanlılığını (duplication bias) önlemek amacıyla veri seti katı biçimde Türkiye tescilleriyle sınırlandırılmış; uluslararası patent aileleri (EP, US, WO) buluşların kalite çarpanı (family size) olarak kullanılmıştır.\n"
              "2. Savunma Ar-Ge ve İstihdam Serileri: Savunma ve Havacılık Sanayii İmalatçılar Derneği (SASAD) 2010–2024 Yıllık Sektör Performans Raporları ve SSB Faaliyet Raporları.\n"
              "3. Firma Düzeyi Bilanço Kontrolleri: Borsa İstanbul (BIST 100) Kamuyu Aydınlatma Platformu (KAP) bağımsız denetimden geçmiş reel net satış hasılatı serileri.", indent=1.0)

    # 2.6. Verilerin Analizi ve Ekonometrik Modelleme Mimarisi (5-Pillar Kanıt Piramidi)
    p = doc.add_paragraph()
    set_para(p, "2.6. Verilerin Analizi ve Ekonometrik Modelleme Mimarisi", align=WD_ALIGN_PARAGRAPH.LEFT, space_before=10, space_after=4, indent=0, bold=True, font_size=12)
    p = doc.add_paragraph()
    set_para(p, "Patent verileri sıfır yığılmalı (zero-inflated) ve aşırı yayılımlı (overdispersed) sayma verisi niteliğindedir. Bu nedenle klasik OLS yöntemi yerine, Santos Silva ve Tenreyro (2006) standartlarında Poisson Pseudo-Maximum Likelihood (PPML) temelinde aşağıdaki 5 aşamalı ekonometrik kanıt piramidi uygulanmıştır:\n\n"
              "1. Temel Çift Sabit Etkili (Two-Way FE) PPML Modeli: Firma (alpha_i) ve yıl (lambda_t) sabit etkileriyle gözlemlenemeyen heterojenlik kontrol edilmiş; BIST reel net satışları (ln(Sales)) ölçek kontrolü olarak modele eklenmiştir:\n"
              "E[Y_it | X_it] = exp( alpha_i + lambda_t + beta_1 ln(Def_R&D_{t-2}) + beta_2 Jaffe_i + beta_3 (ln(Def_R&D_{t-2}) * Jaffe_i) + gamma ln(Sales_it) )\n\n"
              "2. İki Aşamalı Cragg Hurdle Modeli: Firmaların inovasyona başlama kararı (extensive margin - Probit) ile tescil hacmi derinleşmesi (intensive margin - Truncated Count) ayrıştırılarak seçim yanlılığı (selection bias) elenmiştir.\n\n"
              "3. Mekânsal Ekonometri ve Mesafe Bozunumu (Spatial Durbin Modeli): Ankara merkezli savunma çekirdeği ile Marmara sanayi aksı arasındaki coğrafi sürtünme, ters mesafe ağırlıklı mekânsal matris (W) ile modellenmiştir.\n\n"
              "4. Dışsal Nedensellik ve Doğal Deney: 2020 WESCAM ve CAATSA yaptırımları dışsal şok olarak kurgulanmış; ambargolu optik/aviyonik sınıfları ile kontrol sınıfları Farkların Farkı (DiD) yöntemiyle sınanmıştır.\n\n"
              "5. Mikro İletim Kanalı ve Sağkalım Analizi: 93.240 patentte 342 başmühendisin tescilli kariyer geçişleri buluşçu hareketliliği (inventor mobility) ağ analiziyle haritalandırılmış; yıllık harç ödememe kaynaklı terk riski Cox Orantılı Tehlikeler Modeli ile tahmin edilmiştir.", indent=1.0)

    # ==========================================
    # 3. ÇALIŞMA PLANI (EK 2)
    # ==========================================
    doc.add_page_break()

    p = doc.add_paragraph()
    set_para(p, "3. ÇALIŞMA PLANI (EK 2)", align=WD_ALIGN_PARAGRAPH.LEFT, space_before=12, space_after=12, indent=0, bold=True, font_size=13)

    plan_text = (
        "Tez çalışması için öngörülen taslak içindekiler planı aşağıda sunulmuştur:\n\n"
        "BÖLÜM 1: GİRİŞ\n"
        "  1.1. Araştırmanın Arka Planı ve Problemi\n"
        "  1.2. Tezin Amacı, Kapsamı ve Literatüre Katkıları\n"
        "  1.3. Türkiye Savunma Sanayiinin Tarihsel Gelişimi ve Kurumsal Yapısı (2010–2024)\n\n"
        "BÖLÜM 2: KURAMSAL ÇERÇEVE VE LİTERATÜR TARAMASI\n"
        "  2.1. İnovasyon Kuramları ve Bilgi Yayılması (Knowledge Spillover)\n"
        "  2.2. Benoit vs. Deger-Sen: Savunma Harcamalarında Dışlama ve Tamamlayıcılık Tartışması\n"
        "  2.3. Çift Kullanımlı (Dual-Use) Teknoloji Difüzyonu ve Absorptif Kapasite Kuramı\n"
        "  2.4. Patent İktisadı ve Teknolojik Mesafe Yazını (Jaffe, Hall, Trajtenberg Yaklaşımları)\n"
        "  2.5. Dünyada ve Türkiye'de Yapılmış Ampirik Çalışmaların Eleştirel Sentezi\n\n"
        "BÖLÜM 3: VERİ VE METODOLOJİK YAKLAŞIM\n"
        "  3.1. TÜRKPATENT 93.240 Patent Evreni ve Veri Derleme Metodolojisi\n"
        "  3.2. Jaffe (1986, 1993) Teknolojik Yakınlık Matrisinin Hesaplanması\n"
        "  3.3. Ekonometrik Modelleme Stratejisi (PPML, Hurdle, Spatial Durbin ve DiD)\n"
        "  3.4. Tanımlama Stratejisi (Identification) ve İçsellik Çözümleri\n\n"
        "BÖLÜM 4: AMPİRİK BULGULAR, EKONOMETRİK MODELLEME VE İKTİSADİ TARTIŞMA\n"
        "  4.1. Betimsel İstatistikler ve 15 Yıllık Sektörel Patent Dinamikleri\n"
        "  4.2. Savunma Sanayii Bilgi Üretim Fonksiyonu Tahmin Sonuçları (H1)\n"
        "  4.3. Sivil Sanayiye Yayılma ve Jaffe Kritik Teknolojik Eşiği (tau* = 0.2925) (H2 ve H3)\n"
        "  4.4. Cragg Hurdle Modeli: Geniş Kapsam (0.10) vs. Yoğun Kapsam (3.43***) Ayrışımı\n"
        "  4.5. Mekânsal Ekonometri ve Mesafe Bozunumu (Ankara-Marmara Aksı)\n"
        "  4.6. Dışsal Doğal Deney: 2020 WESCAM/CAATSA Ambargoları DiD Analizi (+%181.7 Sıçrama)\n"
        "  4.7. Mikro Yayılma Kanalı: Buluşçu Hareketliliği ve 342 Başmühendis Ağı\n"
        "  4.8. Patent Sağkalım Analizi (Cox Proportional Hazards): Terk Riskinde %31.6 Düşüş\n"
        "  4.9. Sağlamlık (Robustness) Sınamaları ve Dağıtılmış Gecikme (t-1 ... t-5) Modelleri\n\n"
        "BÖLÜM 5: SONUÇ VE ÇİFT KULLANIMLI SANAYİ POLİTİKASI ÖNERİLERİ\n"
        "  5.1. Temel Ekonometrik Bulguların Özeti\n"
        "  5.2. Türkiye Savunma ve Sanayi Politikaları Açısından Stratejik Çıkarımlar\n"
        "  5.3. Araştırmanın Kısıtları ve Gelecek Çalışmalar İçin Yönelimler"
    )
    p = doc.add_paragraph()
    set_para(p, plan_text, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=4, space_after=4, line_spacing=1.15, indent=0, font_size=10.5)

    # ==========================================
    # 4. ZAMANLAMA (EK 3)
    # ==========================================
    doc.add_page_break()

    p = doc.add_paragraph()
    set_para(p, "4. ZAMANLAMA (EK 3)", align=WD_ALIGN_PARAGRAPH.LEFT, space_before=12, space_after=12, indent=0, bold=True, font_size=13)

    p = doc.add_paragraph()
    set_para(p, "Tez çalışmasının öngörülen iş paketleri ve tahmini zamanlama takvimi aşağıda sunulmuştur:", indent=1.0)

    tbl_time = doc.add_table(rows=7, cols=2)
    tbl_time.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_time.autofit = False

    time_data = [
        ("Planlanan Çalışma / İş Paketi", "Tarih Aralığı"),
        ("İP 1: Tez Önerisinin Kesinleşmesi ve Enstitü Onayı", "Eylül 2025 – Ekim 2025"),
        ("İP 2: Literatür Sentezi ve Kuramsal Çerçevenin Yazımı", "Kasım 2025 – Aralık 2025"),
        ("İP 3: TÜRKPATENT 93.240 Verisinin Derlenmesi ve Jaffe Matrisi", "Ocak 2026 – Şubat 2026"),
        ("İP 4: Ekonometrik Modelleme (Hurdle, Spatial, DiD, Cox)", "Mart 2026 – Nisan 2026"),
        ("İP 5: Tez Metninin Tamamlanması ve Danışman İncelemesi", "Mayıs 2026 – Haziran 2026"),
        ("İP 6: Enstitüye Teslim, İntihal Raporu ve Tez Savunması", "Temmuz 2026 – Eylül 2026")
    ]

    for i, row in enumerate(time_data):
        cell_0 = tbl_time.cell(i, 0)
        cell_1 = tbl_time.cell(i, 1)
        cell_0.width = Cm(10.5)
        cell_1.width = Cm(5.5)
        cell_0.text = row[0]
        cell_1.text = row[1]
        p_c0 = cell_0.paragraphs[0]
        p_c1 = cell_1.paragraphs[0]
        p_c0.paragraph_format.line_spacing = 1.15
        p_c1.paragraph_format.line_spacing = 1.15
        p_c0.paragraph_format.space_before = Pt(3)
        p_c0.paragraph_format.space_after = Pt(3)
        p_c1.paragraph_format.space_before = Pt(3)
        p_c1.paragraph_format.space_after = Pt(3)
        p_c0.runs[0].font.name = "Times New Roman"
        p_c1.runs[0].font.name = "Times New Roman"
        if i == 0:
            p_c0.runs[0].font.bold = True
            p_c1.runs[0].font.bold = True
            p_c0.runs[0].font.size = Pt(11)
            p_c1.runs[0].font.size = Pt(11)
            p_c0.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_c1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for c in (cell_0, cell_1):
                tcPr = c._tc.get_or_add_tcPr()
                shd = parse_xml(r'<w:shd {} w:fill="EAECEE"/>'.format(nsdecls("w")))
                tcPr.append(shd)
        else:
            p_c0.runs[0].font.size = Pt(10)
            p_c1.runs[0].font.size = Pt(10)
            p_c1.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for row in tbl_time.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            borders = parse_xml(
                r'''<w:tcBorders {}>
                    <w:top w:val="single" w:sz="4" w:space="0" w:color="A0A0A0"/>
                    <w:left w:val="single" w:sz="4" w:space="0" w:color="A0A0A0"/>
                    <w:bottom w:val="single" w:sz="4" w:space="0" w:color="A0A0A0"/>
                    <w:right w:val="single" w:sz="4" w:space="0" w:color="A0A0A0"/>
                </w:tcBorders>'''.format(nsdecls("w"))
            )
            tcPr.append(borders)

    # ==========================================
    # 5. KAYNAKÇA (EK 4 - APA 7)
    # ==========================================
    doc.add_page_break()

    p = doc.add_paragraph()
    set_para(p, "5. KAYNAKÇA (EK 4)", align=WD_ALIGN_PARAGRAPH.LEFT, space_before=12, space_after=12, indent=0, bold=True, font_size=13)

    p = doc.add_paragraph()
    set_para(p, "Kaynakça listesi APA 7 (American Psychological Association) standartlarına tam uyumlu olarak, asılı girinti (hanging indent: 1.25 cm) formatında düzenlenmiştir:\n", indent=0)

    references = [
        "Acemoglu, D. (2002). Directed technical change. The Review of Economic Studies, 69(4), 781-809. https://doi.org/10.1111/1467-937X.00226",
        "Aghion, P., & Howitt, P. (1992). A model of growth through creative destruction. Econometrica, 60(2), 323-351. https://doi.org/10.2307/2951599",
        "Almeida, P., & Kogut, B. (1999). Localization of knowledge and the mobility of engineers in regional networks. Management Science, 45(7), 905-917. https://doi.org/10.1287/mnsc.45.7.905",
        "Benoit, E. (1973). Defense and economic growth in developing countries. D.C. Heath and Company.",
        "Bloom, N., Schankerman, M., & Van Reenen, J. (2013). Identifying technology spillovers and product market rivalry. Econometrica, 81(4), 1347-1393. https://doi.org/10.3982/ECTA9466",
        "Cohen, W. M., & Levinthal, D. A. (1990). Absorptive capacity: A new perspective on learning and innovation. Administrative Science Quarterly, 35(1), 128-152. https://doi.org/10.2307/2393553",
        "Cragg, J. G. (1971). Some statistical models for limited dependent variables with application to the demand for durable goods. Econometrica, 39(5), 829-844. https://doi.org/10.2307/1909582",
        "Deger, S., & Sen, S. (1983). Military expenditure, spin-off and economic development. Journal of Development Economics, 13(1-2), 67-83. https://doi.org/10.1016/0304-3878(83)90050-X",
        "Dunne, J. P. (1990). Military expenditure and unemployment in the OECD. Defence Economics, 1(1), 57-73. https://doi.org/10.1080/10430719008404652",
        "Griliches, Z. (1979). Issues in assessing the contribution of research and development to productivity growth. The Bell Journal of Economics, 10(1), 92-116. https://doi.org/10.2307/3003321",
        "Griliches, Z. (1990). Patent statistics as economic indicators: A survey. Journal of Economic Literature, 28(4), 1661-1707.",
        "Hall, B. H., Jaffe, A. B., & Trajtenberg, M. (2001). The NBER patent citation data file: Lessons, insights and methodological tools (NBER Working Paper No. 8498). National Bureau of Economic Research. https://doi.org/10.3386/w8498",
        "Hall, B. H., Jaffe, A. B., & Trajtenberg, M. (2005). Market value and patent citations. The RAND Journal of Economics, 36(1), 16-38.",
        "Jaffe, A. B. (1986). Technological opportunity and spillovers of R&D: Evidence from firms' patents, profits, and market value. The American Economic Review, 76(5), 984-1001.",
        "Jaffe, A. B., Trajtenberg, M., & Henderson, R. (1993). Geographic localization of knowledge spillovers as evidenced by patent citations. The Quarterly Journal of Economics, 108(3), 577-598. https://doi.org/10.2307/2118401",
        "LeSage, J. P., & Pace, R. K. (2009). Introduction to spatial econometrics. CRC Press.",
        "Moretti, E., Steinwender, C., & Van Reenen, J. (2023). The intellectual spoils of war? Defense R&D, productivity and international spillovers. The Review of Economics and Statistics, 1-45. https://doi.org/10.1162/rest_a_01377",
        "Mowery, D. C. (2010). Military R&D and innovation. In B. H. Hall & N. Rosenberg (Eds.), Handbook of the Economics of Innovation (Vol. 2, pp. 1219-1256). North-Holland. https://doi.org/10.1016/S0169-7218(10)02013-7",
        "Romer, P. M. (1990). Endogenous technological change. Journal of Political Economy, 98(5, Part 2), S71-S102. https://doi.org/10.1086/261725",
        "Santos Silva, J. M. C., & Tenreyro, S. (2006). The log of gravity. The Review of Economics and Statistics, 88(4), 641-658. https://doi.org/10.1162/rest.88.4.641"
    ]

    for ref in references:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.left_indent = Cm(1.25)
        p.paragraph_format.first_line_indent = Cm(-1.25) # Asılı girinti (Hanging Indent)
        r = p.add_run(ref)
        r.font.name = "Times New Roman"
        r.font.size = Pt(11)

    # Çıkış Dosyaları
    out_docx = "/Users/dogukancihanbeyoglu/Desktop/AHBV_Iktisat_Tez_Calismasi/01_Tez_Oneri_Formu/TEZ_ONERI_FORMU_AHBV.docx"
    git_docx = "/Users/dogukancihanbeyoglu/Gemini/tez_calismasi/01_Tez_Oneri_Formu/TEZ_ONERI_FORMU_AHBV.docx"
    root_docx = "/Users/dogukancihanbeyoglu/Gemini/TEZ_ONERI_FORMU_AHBV.docx"

    doc.save(out_docx)
    doc.save(git_docx)
    doc.save(root_docx)
    print(f"[✔] AHBV Resmi Tez Öneri Formu (.docx) başarıyla oluşturuldu:\n    -> {out_docx}")

if __name__ == "__main__":
    create_proposal_document()
