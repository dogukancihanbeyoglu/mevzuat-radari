#!/usr/bin/env python3
"""
CONVERT CHAPTER 4 TO OFFICIAL AHBV DOCX FORMAT
Formats tables, headings, and econometric results according to AHBV 2025/2026 standards.
"""

import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

MD_PATH = "/Users/dogukancihanbeyoglu/Desktop/AHBV_Iktisat_Tez_Calismasi/01_Tez_Oneri_Formu/BOLUM_4_AMPIRIK_BULGULAR_VE_TARTISMA.md"
DOCX_PATH = "/Users/dogukancihanbeyoglu/Desktop/AHBV_Iktisat_Tez_Calismasi/01_Tez_Oneri_Formu/BOLUM_4_AMPIRIK_BULGULAR_VE_TARTISMA.docx"
GIT_DOCX_PATH = "/Users/dogukancihanbeyoglu/Gemini/tez_calismasi/01_Tez_Oneri_Formu/BOLUM_4_AMPIRIK_BULGULAR_VE_TARTISMA.docx"

def set_cell_background(cell, fill_color):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_color}"/>')
    tcPr.append(shd)

def main():
    doc = docx.Document()

    # Page Margins: AHBV standard (Sol: 4cm, Diğerleri: 2.5cm)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.57) # ~4 cm
        section.right_margin = Inches(1.0)

    # Styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    font.color.rgb = RGBColor(0, 0, 0)

    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p_title.add_run("T.C. ANKARA HACI BAYRAM VELİ ÜNİVERSİTESİ\nLİSANSÜSTÜ EĞİTİM ENSTİTÜSÜ\nİKTİSAT ANABİLİM DALI\n\n")
    r1.font.size = Pt(13)
    r1.font.bold = True
    
    r2 = p_title.add_run("BÖLÜM 4: AMPİRİK BULGULAR, EKONOMETRİK MODELLEME VE İKTİSADİ TARTIŞMA\n\n")
    r2.font.size = Pt(14)
    r2.font.bold = True

    # Read Markdown
    with open(MD_PATH, "r", encoding="utf-8") as f:
        lines = f.readlines()

    in_table = False
    table_lines = []

    for line in lines:
        sline = line.strip()
        if not sline or sline.startswith("---"):
            continue

        if sline.startswith("|"):
            in_table = True
            table_lines.append(sline)
            continue
        else:
            if in_table:
                # Render table
                render_table(doc, table_lines)
                table_lines = []
                in_table = False

        if sline.startswith("## "):
            p = doc.add_paragraph()
            r = p.add_run(sline[3:])
            r.font.size = Pt(13)
            r.font.bold = True
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(6)
        elif sline.startswith("### "):
            p = doc.add_paragraph()
            r = p.add_run(sline[4:])
            r.font.size = Pt(12)
            r.font.bold = True
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
        elif sline.startswith("*   ") or sline.startswith("-   "):
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.space_after = Pt(3)
            # Basic bold parsing
            text = sline[4:]
            add_formatted_text(p, text)
        else:
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.first_line_indent = Inches(0.4) # AHBV paragraf girintisi 1 cm
            add_formatted_text(p, sline)

    if in_table:
        render_table(doc, table_lines)

    doc.save(DOCX_PATH)
    doc.save(GIT_DOCX_PATH)
    print(f"[✔] Successfully generated AHBV docx: {DOCX_PATH}")

def add_formatted_text(paragraph, text):
    parts = text.split("**")
    is_bold = False
    for part in parts:
        if part:
            r = paragraph.add_run(part)
            if is_bold:
                r.bold = True
        is_bold = not is_bold

def render_table(doc, lines):
    rows = []
    for l in lines:
        if "---" in l:
            continue
        cells = [c.strip() for c in l.split("|")[1:-1]]
        if cells:
            rows.append(cells)
    
    if not rows:
        return

    col_count = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            if j < col_count:
                cell = table.cell(i, j)
                cell.text = cell_text.replace("**", "").replace("*", "")
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                p.runs[0].font.name = 'Times New Roman'
                p.runs[0].font.size = Pt(10)
                if i == 0:
                    set_cell_background(cell, "EAECEE")
                    p.runs[0].bold = True

    doc.add_paragraph() # Spacing

if __name__ == "__main__":
    main()
