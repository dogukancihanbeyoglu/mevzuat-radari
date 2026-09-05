#!/usr/bin/env python3
"""
OFFICIAL AHBV TEMPLATE INJECTION ENGINE (2025/2026)
Uses the authentic official HBV template docx (`tez_yazim_sablonu_2025_.docx`),
populating all its native built-in styles, front matters, metadata tables,
and replacing placeholder chapters with our comprehensive 168-page thesis content.
"""

import os
import re
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

OFFICIAL_TEMPLATE_PATH = "/Users/dogukancihanbeyoglu/Desktop/AHBV_Iktisat_Tez_Calismasi/02_Resmi_AHBV_Kilavuzlar_ve_Sablonlar/tez_yazim_sablonu_2025_.docx"
DESKTOP_OUT_PATH = "/Users/dogukancihanbeyoglu/Desktop/AHBV_Iktisat_Tez_Calismasi/01_Tez_Oneri_Formu/AHBV_IKTISAT_TEZI_TAM_METIN.docx"
GIT_OUT_PATH = "/Users/dogukancihanbeyoglu/Gemini/tez_calismasi/01_Tez_Oneri_Formu/AHBV_IKTISAT_TEZI_TAM_METIN.docx"
MD_PATH = "/Users/dogukancihanbeyoglu/Desktop/AHBV_Iktisat_Tez_Calismasi/01_Tez_Oneri_Formu/AHBV_IKTISAT_TEZI_TAM_METIN.md"

def set_cell_background(cell, fill_color):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_color}"/>')
    tcPr.append(shd)

def clean_xml(text):
    if not text:
        return ""
    return "".join(c for c in text if c in ("\t", "\n", "\r") or (0x20 <= ord(c) <= 0xD7FF) or (0xE000 <= ord(c) <= 0xFFFD) or (0x10000 <= ord(c) <= 0x10FFFF))

def populate_official_template():
    print("Loading official AHBV template...")
    doc = docx.Document(OFFICIAL_TEMPLATE_PATH)

    # 1. Update Title Page / Kapak & Jüri Placeholders
    # Paragraph 2: Tez Adı
    thesis_title_tr = "TÜRKİYE SAVUNMA SANAYİİ YAYILMA DİNAMİKLERİNİN İLERİ TEKNOLOJİ PATENT EKOSİSTEMİNE ETKİLERİ: MİKRO-EKONOMETRİK VE MEKÂNSAL BİR ANALİZ (2010–2024)"
    thesis_title_en = "SPILLOVER DYNAMICS OF THE TURKISH DEFENSE INDUSTRY ON THE ADVANCED TECHNOLOGY PATENT ECOSYSTEM: A MICRO-ECONOMETRIC AND SPATIAL ANALYSIS (2010–2024)"
    author_name = "Doğukan CİHANBEYOĞLU"
    supervisor = "Tez Danışmanı: [Unvanı, Adı SOYADI]"
    department_tr = "İKTİSAT ANA BİLİM DALI"
    program_tr = "İKTİSAT TEZLİ YÜKSEK LİSANS PROGRAMI"
    city_year = "Ankara - 2027"

    # Replace in early paragraphs
    for p in doc.paragraphs[:150]:
        t = p.text.strip()
        if "Tez Adı Baş Harfleri Büyük" in t or "Büyük harflerle ve ortalanmış olarak tez adı" in t:
            p.text = clean_xml(thesis_title_tr)
            for r in p.runs:
                r.font.name = "Times New Roman"
                r.font.size = Pt(14)
                r.bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif t == "Ad SOYAD" or t == "Adı SOYADI" or t == "(Ad ve SOYAD)":
            p.text = clean_xml(author_name)
            for r in p.runs:
                r.font.name = "Times New Roman"
                r.font.size = Pt(12)
                r.bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif "Tez Danışmanı Unvan Ad SOYAD" in t:
            p.text = supervisor
            for r in p.runs:
                r.font.name = "Times New Roman"
                r.font.size = Pt(12)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif "… ANA BİLİM DALI" in t:
            p.text = department_tr
            for r in p.runs:
                r.font.name = "Times New Roman"
                r.font.size = Pt(12)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif "… PROGRAMI" in t:
            p.text = program_tr
            for r in p.runs:
                r.font.name = "Times New Roman"
                r.font.size = Pt(12)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif "YÜKSEK LİSANS TEZİ/DOKTORA" in t:
            p.text = "YÜKSEK LİSANS TEZİ"
            for r in p.runs:
                r.font.name = "Times New Roman"
                r.font.size = Pt(12)
                r.bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif "SAVUNMA YILI" in t:
            p.text = city_year
            for r in p.runs:
                r.font.name = "Times New Roman"
                r.font.size = Pt(12)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif "(Tezin/Projenin Türkçe Başlığı)" in t:
            p.text = clean_xml(thesis_title_tr)
            for r in p.runs:
                r.font.name = "Times New Roman"
                r.font.size = Pt(12)
                r.bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif "(English Title of Thesis/Project)" in t:
            p.text = clean_xml(thesis_title_en)
            for r in p.runs:
                r.font.name = "Times New Roman"
                r.font.size = Pt(12)
                r.bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif "Name and SURNAME" in t:
            p.text = clean_xml(author_name)
            for r in p.runs:
                r.font.name = "Times New Roman"
                r.font.size = Pt(12)
                r.bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif "Supervisor: Academic Title" in t:
            p.text = "Supervisor: [Title, Name SURNAME]"
            for r in p.runs:
                r.font.name = "Times New Roman"
                r.font.size = Pt(12)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif "Department of …" in t:
            p.text = "Department of Economics, Master's Program"
            for r in p.runs:
                r.font.name = "Times New Roman"
                r.font.size = Pt(12)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif "Year, Ankara" in t:
            p.text = "January 2027, Ankara"
            for r in p.runs:
                r.font.name = "Times New Roman"
                r.font.size = Pt(12)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif "Master’s Thesis, Doctoral Thesis" in t:
            p.text = "Master's Thesis"
            for r in p.runs:
                r.font.name = "Times New Roman"
                r.font.size = Pt(12)
                r.bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif "… Ana Bilim Dalı/…Ana Sanat Dalı" in t:
            p.text = "İktisat Ana Bilim Dalı"
            for r in p.runs:
                r.font.name = "Times New Roman"
                r.font.size = Pt(12)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif "Yıl, Ankara" in t:
            p.text = "Ocak 2027, Ankara"
            for r in p.runs:
                r.font.name = "Times New Roman"
                r.font.size = Pt(12)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Update ONAY text
    for p in doc.paragraphs:
        if "tarih ve saatinde yapılan tez savunma sınavında" in p.text:
            p.text = f"İktisat Ana Bilim Dalı İktisat Tezli Yüksek Lisans Programı öğrencisi {author_name} tarafından hazırlanan '{thesis_title_tr}' başlıklı tez çalışması .../01/2027 tarihinde yapılan tez savunma sınavında aşağıdaki jüri tarafından OY BİRLİĞİ ile YÜKSEK LİSANS TEZİ olarak KABUL edilmiştir."
            for r in p.runs:
                r.font.name = "Times New Roman"
                r.font.size = Pt(11)
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Update Özet body text (paragraph 58)
    with open(MD_PATH, "r", encoding="utf-8") as f:
        md_text = f.read()

    # Extract Özet text
    ozet_match = re.search(r"## ÖZET.*?\n\n(.*?)\n\n\*\*Anahtar Kelimeler:\*\*", md_text, re.DOTALL)
    if ozet_match:
        ozet_text = ozet_match.group(1).replace("\n", " ").strip()
        for p in doc.paragraphs[50:70]:
            if "……………………………………………………" in p.text:
                p.text = clean_xml(ozet_text)
                for r in p.runs:
                    r.font.name = "Times New Roman"
                    r.font.size = Pt(11)
                p.paragraph_format.line_spacing = 1.0 # AHBV kuralı: özet tek satır aralığı
                p.paragraph_format.space_after = Pt(6)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                break

    # Extract Keywords TR
    kw_tr_match = re.search(r"\*\*Anahtar Kelimeler:\*\*\s*(.*)", md_text)
    if kw_tr_match:
        for p in doc.paragraphs[55:75]:
            if "Anahtar Kelimeler:" in p.text:
                p.text = "Anahtar Kelimeler: " + kw_tr_match.group(1).strip()
                for r in p.runs:
                    r.font.name = "Times New Roman"
                    r.font.size = Pt(11)
                p.paragraph_format.line_spacing = 1.0
                break

    # Extract Abstract text EN
    abstract_match = re.search(r"## ABSTRACT.*?\n\n(.*?)\n\n\*\*Keywords:\*\*", md_text, re.DOTALL)
    if abstract_match:
        abstract_text = abstract_match.group(1).replace("\n", " ").strip()
        for p in doc.paragraphs[80:100]:
            if "……………………………………………………" in p.text:
                p.text = clean_xml(abstract_text)
                for r in p.runs:
                    r.font.name = "Times New Roman"
                    r.font.size = Pt(11)
                p.paragraph_format.line_spacing = 1.0 # AHBV kuralı: abstract tek satır aralığı
                p.paragraph_format.space_after = Pt(6)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                break

    # Extract Keywords EN
    kw_en_match = re.search(r"\*\*Keywords:\*\*\s*(.*)", md_text)
    if kw_en_match:
        for p in doc.paragraphs[85:105]:
            if "Keywords:" in p.text:
                p.text = "Keywords: " + kw_en_match.group(1).strip()
                for r in p.runs:
                    r.font.name = "Times New Roman"
                    r.font.size = Pt(11)
                p.paragraph_format.line_spacing = 1.0
                break

    # Update Teşekkür / İthaf
    for p in doc.paragraphs[95:110]:
        if "İthaf; herhangi bir şekil" in p.text:
            p.text = "Bu çalışma; Türkiye'nin bağımsız teknoloji hamlesine ömrünü adayan mühendis ve bilim insanlarına ithaf edilmiştir."
            for r in p.runs:
                r.font.name = "Times New Roman"
                r.font.size = Pt(11)
                r.italic = True
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif "Metin, 12 punto büyüklüğünde, tek satır aralığında yazılır" in p.text:
            p.text = "Bu tez çalışmasının hazırlanma sürecinde değerli bilgi, birikim ve tecrübeleriyle yolumu aydınlatan kıymetli tez danışmanıma, araştırmam sırasında kurumsal ve teknik desteklerini esirgemeyen tüm hocalarıma ve her daim yanımda olan aileme en içten teşekkürlerimi sunarım."
            for r in p.runs:
                r.font.name = "Times New Roman"
                r.font.size = Pt(11)
            p.paragraph_format.line_spacing = 1.0
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Populate Kısaltmalar tablosu (Table 1 in doc)
    abbr_data = [
        ("AHBV", "Ankara Hacı Bayram Veli Üniversitesi"),
        ("ASELSAN", "Askeri Elektronik Sanayi A.Ş."),
        ("BIST", "Borsa İstanbul"),
        ("CAATSA", "Countering America's Adversaries Through Sanctions Act"),
        ("CPC", "Cooperative Patent Classification (Ortak Patent Sınıflandırması)"),
        ("DiD", "Difference-in-Differences (Farkların Farkı Yöntemi)"),
        ("DOCDB", "EPO Worldwide Patent Database"),
        ("EPO", "European Patent Office (Avrupa Patent Ofisi)"),
        ("GSYİH", "Gayri Safi Yurtiçi Hasıla"),
        ("IPC", "International Patent Classification (Uluslararası Patent Sınıflandırması)"),
        ("KAP", "Kamuyu Aydınlatma Platformu"),
        ("KPF", "Knowledge Production Function (Bilgi Üretim Fonksiyonu)"),
        ("OLS", "Ordinary Least Squares (En Küçük Kareler Yöntemi)"),
        ("PPML", "Poisson Pseudo-Maximum Likelihood"),
        ("ROKETSAN", "Roket Sanayii ve Ticaret A.Ş."),
        ("SASAD", "Savunma ve Havacılık Sanayii İmalatçılar Derneği"),
        ("SDM", "Spatial Durbin Model (Mekânsal Durbin Modeli)"),
        ("SSB", "Savunma Sanayii Başkanlığı"),
        ("SSİK", "Savunma Sanayii İcra Komitesi"),
        ("TCMB", "Türkiye Cumhuriyet Merkez Bankası"),
        ("TFP", "Total Factor Productivity (Toplam Faktör Verimliliği)"),
        ("TSKGV", "Türk Silahlı Kuvvetlerini Güçlendirme Vakfı"),
        ("TUSAŞ", "Türk Havacılık ve Uzay Sanayii A.Ş."),
        ("TÜBİTAK", "Türkiye Bilimsel ve Teknolojik Araştırma Kurumu"),
        ("TÜİK", "Türkiye İstatistik Kurumu"),
        ("TÜRKPATENT", "Türk Patent ve Marka Kurumu"),
        ("TWFE", "Two-Way Fixed Effects (Çift Yönlü Sabit Etkiler)"),
        ("WIPO", "World Intellectual Property Organization"),
    ]
    if len(doc.tables) > 1:
        abbr_table = doc.tables[1] # Kısaltmalar tablosu
        # clear existing rows except header
        while len(abbr_table.rows) > 1:
            abbr_table._tbl.remove(abbr_table.rows[-1]._tr)
        for abbr, full in abbr_data:
            row = abbr_table.add_row()
            row.cells[0].text = abbr
            row.cells[1].text = full
            for c_idx, cell in enumerate(row.cells):
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.line_spacing = 1.0
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                for r in p.runs:
                    r.font.name = "Times New Roman"
                    r.font.size = Pt(10)
                    if c_idx == 0: r.bold = True

    # Now remove sample placeholder body text from paragraph 150 downwards
    # and inject our full exhaustive thesis chapters using official template styles!
    print("Truncating sample placeholder body paragraphs from official template...")
    # Find paragraph with style 'BÖLÜM BAŞLIKLARI' and text 'GİRİŞ'
    start_body_idx = None
    for idx, p in enumerate(doc.paragraphs):
        if p.style.name == "BÖLÜM BAŞLIKLARI" and "GİRİŞ" in p.text:
            start_body_idx = idx
            break
    
    if start_body_idx is not None:
        # remove all paragraphs from start_body_idx to end
        for p in doc.paragraphs[start_body_idx:]:
            p._element.getparent().remove(p._element)

    # Now parse our comprehensive thesis markdown starting from # BÖLÜM 1: GİRİŞ
    body_md_idx = md_text.find("# BÖLÜM 1: GİRİŞ")
    thesis_body_md = md_text[body_md_idx:]

    print("Injecting comprehensive chapters using authentic AHBV template styles...")
    lines = thesis_body_md.split("\n")
    in_table = False
    table_lines = []

    for line in lines:
        sline = line.strip()

        if sline.startswith("|") and sline.endswith("|"):
            in_table = True
            table_lines.append(sline)
            continue
        else:
            if in_table:
                render_table(doc, table_lines)
                in_table = False
                table_lines = []

        if not sline:
            continue

        if sline.startswith("# BÖLÜM ") or sline.startswith("# KAYNAKÇA"):
            # AHBV Ana Bölüm Başlığı (1. Derece): BÖLÜM BAŞLIKLARI stili
            header_text = sline.replace("# ", "").strip()
            # If it is Bölüm 1: GİRİŞ -> AHBV kuralı: 1. GİRİŞ (alt başlık numarası almaz)
            p = doc.add_paragraph(style='BÖLÜM BAŞLIKLARI')
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(12)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(clean_xml(header_text))
            r.font.name = "Times New Roman"
            r.font.size = Pt(12)
            r.bold = True
        elif sline.startswith("## "):
            # AHBV 2. Derece Başlık: İKİNCİ BAŞLIK stili
            p = doc.add_paragraph(style='İKİNCİ BAŞLIK')
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(clean_xml(sline[3:].strip()))
            r.font.name = "Times New Roman"
            r.font.size = Pt(12)
            r.bold = True
        elif sline.startswith("### "):
            # AHBV 3. Derece Başlık: ÜÇÜNCÜ BAŞLIK stili
            p = doc.add_paragraph(style='ÜÇÜNCÜ BAŞLIK')
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(clean_xml(sline[4:].strip()))
            r.font.name = "Times New Roman"
            r.font.size = Pt(12)
            r.bold = True
            r.italic = True
        elif sline.startswith("$$"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            r = p.add_run(clean_xml(sline.replace("$$", "").strip()))
            r.font.name = "Times New Roman"
            r.font.size = Pt(11)
            r.italic = True
        elif sline.startswith("---"):
            continue
        elif sline.startswith("* ") or sline.startswith("- "):
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.left_indent = Inches(0.4) # 1 cm
            add_formatted_text(p, "• " + sline[2:].strip())
        elif re.match(r"^\d+\.\s", sline):
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.left_indent = Inches(0.4)
            add_formatted_text(p, sline)
        else:
            # Body text per AHBV: 12 pt Times New Roman, 1.5 satır aralığı, 6nk önce-sonra, 1.0 cm girinti, iki yana yaslı
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.first_line_indent = Inches(0.3937) # 1.0 cm
            add_formatted_text(p, sline)

    if in_table:
        render_table(doc, table_lines)

    # Save to both paths
    doc.save(DESKTOP_OUT_PATH)
    doc.save(GIT_OUT_PATH)
    print(f"[✔] Successfully injected into Official Template: {DESKTOP_OUT_PATH}")
    print(f"[✔] Synced to Git path: {GIT_OUT_PATH}")

def add_formatted_text(paragraph, text):
    text = clean_xml(text)
    parts = text.split("**")
    is_bold = False
    for part in parts:
        if part:
            subparts = part.split("*")
            is_italic = False
            for sub in subparts:
                if sub:
                    r = paragraph.add_run(sub)
                    r.font.name = "Times New Roman"
                    r.font.size = Pt(12)
                    if is_bold: r.bold = True
                    if is_italic: r.italic = True
                is_italic = not is_italic
        is_bold = not is_bold

def render_table(doc, lines):
    rows = []
    for l in lines:
        if "---" in l: continue
        cells = [clean_xml(c.strip()) for c in l.split("|")[1:-1]]
        if cells: rows.append(cells)
    if not rows: return

    col_count = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx]
        is_header = (r_idx == 0)
        for c_idx in range(min(col_count, len(row_data))):
            cell = row.cells[c_idx]
            cell.text = row_data[c_idx]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if (is_header or c_idx > 0) else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.line_spacing = 1.0 # AHBV kuralı: tablo tek satır aralığı
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)
                if is_header:
                    run.bold = True
            if is_header:
                set_cell_background(cell, "EAECEE")
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

if __name__ == "__main__":
    populate_official_template()
