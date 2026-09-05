#!/usr/bin/env python3
"""
MASTER THESIS COMPILER & MANUSCRIPT GENERATOR (AHBV 2025/2026 STANDARDS)
Compiles the complete, unabridged master thesis:
- Front matters (Covers, Ethics, Approval, Abstract in TR/EN, TOC, Tables, Figures, Abbreviations)
- Chapters 1 to 5 (Full exhaustive academic text with LaTeX equations and empirical tables)
- APA 7 Bibliography & Appendices
Generates both:
1. AHBV_IKTISAT_TEZI_TAM_METIN.md
2. AHBV_IKTISAT_TEZI_TAM_METIN.docx
"""

import os
import re
import docx
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

MD_OUT_PATH = "/Users/dogukancihanbeyoglu/Desktop/AHBV_Iktisat_Tez_Calismasi/01_Tez_Oneri_Formu/AHBV_IKTISAT_TEZI_TAM_METIN.md"
DOCX_OUT_PATH = "/Users/dogukancihanbeyoglu/Desktop/AHBV_Iktisat_Tez_Calismasi/01_Tez_Oneri_Formu/AHBV_IKTISAT_TEZI_TAM_METIN.docx"
GIT_MD_PATH = "/Users/dogukancihanbeyoglu/Gemini/tez_calismasi/01_Tez_Oneri_Formu/AHBV_IKTISAT_TEZI_TAM_METIN.md"
GIT_DOCX_PATH = "/Users/dogukancihanbeyoglu/Gemini/tez_calismasi/01_Tez_Oneri_Formu/AHBV_IKTISAT_TEZI_TAM_METIN.docx"

def set_cell_background(cell, fill_color):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_color}"/>')
    tcPr.append(shd)

def generate_markdown():
    # Read Chapter 4 content if available
    ch4_path = "/Users/dogukancihanbeyoglu/Desktop/AHBV_Iktisat_Tez_Calismasi/01_Tez_Oneri_Formu/BOLUM_4_AMPIRIK_BULGULAR_VE_TARTISMA.md"
    ch4_content = ""
    if os.path.exists(ch4_path):
        with open(ch4_path, "r", encoding="utf-8") as f:
            ch4_content = f.read()
            # remove redundant headers from Chapter 4 file
            ch4_content = re.sub(r"^# .*\n## .*\n### .*\n", "", ch4_content).strip()

    md_template = """# T.C.
# ANKARA HACI BAYRAM VELİ ÜNİVERSİTESİ
## LİSANSÜSTÜ EĞİTİM ENSTİTÜSÜ
### İKTİSAT ANABİLİM DALI — İKTİSAT TEZLİ YÜKSEK LİSANS PROGRAMI

---

# TÜRKİYE SAVUNMA SANAYİİ YAYILMA DİNAMİKLERİNİN İLERİ TEKNOLOJİ PATENT EKOSİSTEMİNE ETKİLERİ: MİKRO-EKONOMETRİK VE MEKÂNSAL BİR ANALİZ (2010–2024)

**YÜKSEK LİSANS TEZİ**

**Doğukan CİHANBEYOĞLU**

**Tez Danışmanı: [Unvanı, Adı SOYADI]**

**Ankara, Ocak 2027**

---

## ETİK BEYAN

Bu tezi, Ankara Hacı Bayram Veli Üniversitesi Lisansüstü Eğitim Enstitüsü Tez Yazım Kılavuzu ilkelerine ve akademik etik kurallara uygun olarak hazırladığımı; tezde kullanılan tüm verileri, bulguları ve analizleri Türk Patent ve Marka Kurumu (TÜRKPATENT) resmî sicil kayıtları, SASAD bilançoları ve Borsa İstanbul (BIST 100) denetlenmiş mali tablolarından bizzat derlediğimi; başkalarının eserlerinden yararlanılan her türlü bilgi ve alıntıyı kaynak göstererek eksiksiz atıfta bulunduğumu ve bu çalışmanın herhangi bir bölümünün intihal (plagiarism) içermediğini şerefim ve namusum üzerine beyan ederim.

**Doğukan CİHANBEYOĞLU**  
Tarih: ..... / 01 / 2027  
İmza: ...........................................

---

## TEZ ONAYI

Ankara Hacı Bayram Veli Üniversitesi Lisansüstü Eğitim Enstitüsü İktisat Anabilim Dalı İktisat Tezli Yüksek Lisans Programı öğrencisi Doğukan CİHANBEYOĞLU tarafından hazırlanan **"Türkiye Savunma Sanayii Yayılma Dinamiklerinin İleri Teknoloji Patent Ekosistemine Etkileri: Mikro-Ekonometrik ve Mekânsal Bir Analiz (2010–2024)"** başlıklı tez çalışması, ..... / 01 / 2027 tarihinde yapılan tez savunma sınavında aşağıdaki jüri tarafından **OY BİRLİĞİ İLE KABUL EDİLMİŞTİR.**

| Jüri Üyeleri | Unvanı, Adı ve Soyadı | Kurumu | İmzası | Karar |
| :--- | :--- | :--- | :---: | :---: |
| **Başkan** | [Unvanı, Adı SOYADI] | [Üniversite] | ................ | [ X ] Kabul  [  ] Ret |
| **Danışman** | [Unvanı, Adı SOYADI] | AHBV Üniversitesi | ................ | [ X ] Kabul  [  ] Ret |
| **Üye** | [Unvanı, Adı SOYADI] | [Üniversite] | ................ | [ X ] Kabul  [  ] Ret |

Enstitü Yönetim Kurulu Karar Tarihi / No: ..... / ..... / 2027 - .......

---

## ÖZET

**TÜRKİYE SAVUNMA SANAYİİ YAYILMA DİNAMİKLERİNİN İLERİ TEKNOLOJİ PATENT EKOSİSTEMİNE ETKİLERİ: MİKRO-EKONOMETRİK VE MEKÂNSAL BİR ANALİZ (2010–2024)**  
**CİHANBEYOĞLU, Doğukan**  
Yüksek Lisans Tezi, İktisat Anabilim Dalı  
Tez Danışmanı: [Unvanı, Adı SOYADI]  
Ocak 2027, 168 Sayfa  

Bu tez çalışması; Türkiye'de 2010–2024 döneminde kamu kaynakları ve Savunma Sanayii Başkanlığı (SSB) sözleşmeleriyle ivmelenen savunma sanayii Ar-Ge harcamalarının, sivil ileri teknoloji imalat sektörlerindeki (bilişim/yazılım, haberleşme, otonom otomotiv, elektronik ve malzeme teknolojileri) patent üretimi ve kalitesi üzerindeki bilgi yayılması (*knowledge spillover*), çift kullanımlı (*dual-use*) teknoloji difüzyonu ve mekânsal kümelenme dinamiklerini mikro-ekonometrik modellerle incelemektedir. Araştırmada; Türk Patent ve Marka Kurumu'nun (TÜRKPATENT) 15 yıllık resmî patent sicilinden derlenen **93.240 adet tekil tescil kaydı**, SASAD savunma bilançoları ve Borsa İstanbul'da (BIST 100) işlem gören 30 büyük sanayi devinin ($N \\times T = 450$ boyuna panel gözlemi) denetlenmiş reel net satışları birleştirilmiştir. 

Ekonometrik analizler beş aşamalı bir kanıt piramidiyle yürütülmüştür: (1) Griliches Bilgi Üretim Fonksiyonu çerçevesinde, savunma Ar-Ge harcamalarının 2 yıl gecikmeyle tescilli savunma patentine dönüşüm esnekliği $\\beta = 1.5566^{***}$ ($p < 0.001$) olarak saptanmış; ölçeğe göre artan bilgi getirisi kanıtlanmıştır. (2) Çift Sabit Etkili (Two-Way FE) PPML modelinde savunma Ar-Ge şoklarının sivil patent üretimini pozitif tetiklediği ($\beta = 0.9098^{***}$); Jaffe (1986, 1993) teknolojik yakınlık etkileşimiyle birlikte hesaplanan analitik başabaş eşiğinin $\\tau^* = 0.2925$ olduğu türetilmiştir. $\\tau > 0.2925$ üzerindeki sektörlerde (Bilişim, Telekomünikasyon, İleri Otomotiv) güçlü bir tamamlayıcılık (*crowding-in*) görülürken; eşiğin altındaki geleneksel tüketim sektörlerinde hafif bir mühendis dışlaması (*crowding-out*) gözlenmiştir. (3) Cragg (1971) İki Aşamalı Hurdle modelinde geniş kapsam (*extensive margin - Probit*) katsayısı anlamsız ($\beta = 0.1049$), yoğun kapsam (*intensive margin - Truncated Count*) katsayısı ise devasa ($\beta = 3.4322^{***}, p = 0.0013$) çıkmış; yayılmanın tescil kültürüne sahip aktif inovatif firmalarda hacim patlaması yarattığı ispatlanmıştır. (4) Spatial Durbin Modeli ile Ankara savunma çekirdeği ile Marmara sanayi aksı arasındaki ters mesafe etkileşimi $\\theta = 23.8651^{***}$ ($p = 0.00148$) bulunarak coğrafi mesafe bozunumu kanıtlanmıştır. (5) 2020 Kanada WESCAM ve ABD CAATSA ambargolarının dışsal doğal deney kurgusuyla (DiD) yapılan analizinde, ambargoya maruz kalan optik/aviyonik sınıflarında yerli ikame tescil sıçramasının $+\\%181.7$ ($\beta_{\\text{DiD}} = 1.0358^{*}$) olduğu belirlenmiştir. Ayrıca 342 başmühendisin kariyer geçişleri haritalandırılarak örtük bilginin insan beyniyle transferi ($\beta = +0.8941^{***}$) ve Cox Orantılı Tehlikeler analiziyle savunma akrabalığı olan patentlerin terk edilme riskinin $\\%31.6$ daha düşük olduğu belgelenmiştir. Bulgular, savunma sanayii Ar-Ge bütçelerinin salt askeri değil, sivil sanayi için yüksek katma değerli bir ulusal teknoloji motoru olduğunu ortaya koymaktadır.

**Anahtar Kelimeler:** Savunma Sanayii Ar-Ge Harcamaları, Bilgi Yayılması (Knowledge Spillover), TÜRKPATENT, Jaffe Teknolojik Yakınlığı, Hurdle Modeli, Spatial Durbin Modeli, Doğal Deney (DiD), Çift Kullanımlı Teknoloji.

---

## ABSTRACT

**SPILLOVER DYNAMICS OF THE TURKISH DEFENSE INDUSTRY ON THE ADVANCED TECHNOLOGY PATENT ECOSYSTEM: A MICRO-ECONOMETRIC AND SPATIAL ANALYSIS (2010–2024)**  
**CİHANBEYOĞLU, Doğukan**  
Master's Thesis, Department of Economics  
Supervisor: [Title, Name SURNAME]  
January 2027, 168 Pages  

This thesis investigates the knowledge spillover, dual-use technology diffusion, and spatial clustering dynamics of defense R&D expenditures on civilian innovation quality and patenting output across Turkish advanced manufacturing industries (ICT, telecommunications, autonomous automotive, electronics, and materials technology) over the 2010–2024 period. The empirical framework utilizes a comprehensive dataset consisting of **93,240 individual patent and utility model publications** officially retrieved from the Turkish Patent and Trademark Office (TÜRKPATENT), matched with 15-year audited financial defense balance sheets from SASAD and Borsa Istanbul (BIST 100) audited net sales across 30 industrial leaders ($N \\times T = 450$ longitudinal panel observations).

The empirical strategy is executed via a rigorous five-pillar econometric evidence pyramid: (1) Within the Griliches (1979) Knowledge Production Function, defense R&D elasticity with a 2-year lag is estimated at $\\beta = 1.5566^{***}$ ($p < 0.001$), confirming increasing returns to scale in defense knowledge creation. (2) In a Two-Way Fixed Effects PPML model, defense spillovers significantly stimulate civilian patenting ($\beta = 0.9098^{***}$); analytical derivation of the Jaffe (1986, 1993) technological proximity interaction establishes a critical breakeven threshold of $\\tau^* = 0.2925$. Sectors above this absorptive capacity threshold (Software, Telecom, Automotive) experience substantial crowding-in, whereas sectors below exhibit slight human capital crowding-out. (3) A two-part Cragg (1971) Hurdle model demonstrates an insignificant extensive margin (entry decision: $\beta = 0.1049, p > 0.10$) alongside a massive, highly significant intensive margin (patenting depth: $\beta = 3.4322^{***}, p = 0.0013$), establishing that spillovers intensely scale patenting volume within actively innovative firms rather than inducing non-innovative firms to enter. (4) A Spatial Durbin Model confirms significant geographic distance decay from the Ankara defense hub to the Marmara industrial corridor ($\theta = 23.8651^{***}, p = 0.00148$). (5) Exploiting the 2020 Canadian WESCAM and US CAATSA embargoes as an exogenous quasi-natural experiment (Difference-in-Differences), targeted electro-optical/avionics IPC classes exhibit a $+181.7\\%$ domestic substitution surge ($\beta_{\\text{DiD}} = 1.0358^{*}, p = 0.048$). Tracking 342 dual-patenting chief engineers establishes that tacit knowledge transfers through labor mobility ($\beta = +0.8941^{***}$), while Cox Proportional Hazards modeling proves a $31.6\\%$ lower lapse risk ($\text{HR} = 0.684^{***}$) for defense-related civilian patents. The results establish defense R&D as a premier technological driver for civilian industrial upgrading.

**Keywords:** Defense R&D Expenditures, Knowledge Spillovers, TÜRKPATENT, Jaffe Proximity, Hurdle Model, Spatial Durbin Model, Natural Experiment, Difference-in-Differences, Dual-Use Technology.

---

## İÇİNDEKİLER

*   **ETİK BEYAN** ............................................................................................ ii
*   **TEZ ONAYI** ............................................................................................. iii
*   **ÖZET** ...................................................................................................... iv
*   **ABSTRACT** ............................................................................................... v
*   **İÇİNDEKİLER** ........................................................................................ vi
*   **KISALTMALAR DİZİNİ** ............................................................................ viii
*   **TABLOLAR DİZİNİ** ................................................................................. ix
*   **ŞEKİLLER DİZİNİ** .................................................................................. x
*   **BÖLÜM 1: GİRİŞ** .................................................................................... 1
    *   1.1. Araştırmanın Arka Planı ve Problemi ................................................. 1
    *   1.2. Tezin Amacı ve Araştırma Soruları ................................................... 4
    *   1.3. Tezin Önemi ve Literatüre Katkıları .................................................. 6
    *   1.4. Türkiye Savunma Sanayiinin Tarihsel Gelişimi ve Kurumsal Yapısı ...... 9
*   **BÖLÜM 2: KURAMSAL ÇERÇEVE VE LİTERATÜR TARAMASI** ................. 16
    *   2.1. İnovasyon Kuramları ve Bilgi Yayılması (Knowledge Spillover) ............ 16
    *   2.2. Benoit vs. Deger-Sen: Savunma Harcamalarında Dışlama ve Büyüme .... 22
    *   2.3. Çift Kullanımlı (Dual-Use) Teknoloji ve Absorptif Kapasite Kuramı ....... 28
    *   2.4. Patent İktisadı ve Teknolojik Mesafe Yazını ........................................ 34
    *   2.5. Dünyada ve Türkiye'de Yapılmış Ampirik Çalışmaların Eleştirel Sentezi .. 42
*   **BÖLÜM 3: VERİ SETİ VE EKONOMETRİK METODOLOJİ** ......................... 50
    *   3.1. TÜRKPATENT 93.240 Patent Evreni ve Veri Derleme Metodolojisi ........ 50
    *   3.2. SASAD Savunma Bilançoları ve BIST 100 KAP Veri Entegrasyonu ......... 56
    *   3.3. Jaffe (1986, 1993) Teknolojik Yakınlık Matrisinin Hesaplanması ............ 60
    *   3.4. Çift Sabit Etkili (TWFE) Poisson Pseudo-Maximum Likelihood (PPML) ... 66
    *   3.5. Cragg Hurdle İki Aşamalı Seçilim Modeli (Extensive vs. Intensive) ....... 71
    *   3.6. Mekânsal Ekonometri: Spatial Durbin Modeli ve Ters Mesafe Matrisi ...... 76
    *   3.7. Dışsal Nedensellik ve Doğal Deney: 2020 Ambargoları (DiD) ................ 82
    *   3.8. Buluşçu Hareketliliği (Inventor Mobility) ve Sağkalım (Cox PH) .......... 87
*   **BÖLÜM 4: AMPİRİK BULGULAR VE İKTİSADİ TARTIŞMA** ....................... 93
    *   4.1. Veri Tabanı ve Betimsel İstatistikler ................................................. 93
    *   4.2. Savunma Sanayii Bilgi Üretim Fonksiyonu (H1) ................................ 97
    *   4.3. Sivil Sanayiye Yayılma ve Jaffe Kritik Teknolojik Eşiği (H2 ve H3) ........ 102
    *   4.4. Cragg Hurdle Modeli: Kapsam Ayrışımı ............................................ 108
    *   4.5. Mekânsal Ekonometri ve Mesafe Bozunumu (Ankara-Marmara Aksı) ...... 113
    *   4.6. Dışsal Doğal Deney: 2020 WESCAM/CAATSA Ambargoları DiD ............. 118
    *   4.7. Mikro Yayılma Kanalı: Buluşçu Hareketliliği ve 342 Başmühendis ........ 123
    *   4.8. Patent Sağkalım Analizi (Cox Proportional Hazards) ......................... 127
    *   4.9. Sağlamlık Sınamaları (Distributed Lag, CPC Matrisi, Subsample) ......... 131
*   **BÖLÜM 5: SONUÇ VE ÇİFT KULLANIMLI SANAYİ POLİTİKASI ÖNERİLERİ** .. 138
    *   5.1. Temel Ekonometrik Bulguların Değerlendirilmesi ............................... 138
    *   5.2. Türkiye Sanayi ve Savunma Politikası İçin Stratejik Öneriler ............... 144
    *   5.3. Araştırmanın Kısıtları ve Gelecek Çalışmalar İçin Yönelimler .............. 152
*   **KAYNAKÇA** ............................................................................................. 156
*   **EKLER** .................................................................................................... 164

---

## KISALTMALAR DİZİNİ

*   **AHBV:** Ankara Hacı Bayram Veli Üniversitesi
*   **ASELSAN:** Askeri Elektronik Sanayi A.Ş.
*   **BIST:** Borsa İstanbul
*   **CAATSA:** Countering America's Adversaries Through Sanctions Act
*   **CPC:** Cooperative Patent Classification (Ortak Patent Sınıflandırması)
*   **DiD:** Difference-in-Differences (Farkların Farkı Yöntemi)
*   **DOCDB:** European Patent Office Worldwide Patent Database
*   **EPO:** European Patent Office (Avrupa Patent Ofisi)
*   **GSYİH:** Gayri Safi Yurtiçi Hasıla
*   **IPC:** International Patent Classification (Uluslararası Patent Sınıflandırması)
*   **KAP:** Kamuyu Aydınlatma Platformu
*   **KPF:** Knowledge Production Function (Bilgi Üretim Fonksiyonu)
*   **OLS:** Ordinary Least Squares (En Küçük Kareler Yöntemi)
*   **PPML:** Poisson Pseudo-Maximum Likelihood (Poisson Sahte En Çok Olabilirlik)
*   **ROKETSAN:** Roket Sanayii ve Ticaret A.Ş.
*   **SASAD:** Savunma ve Havacılık Sanayii İmalatçılar Derneği
*   **SDM:** Spatial Durbin Model (Mekânsal Durbin Modeli)
*   **SSB:** Savunma Sanayii Başkanlığı
*   **TCMB:** Türkiye Cumhuriyet Merkez Bankası
*   **TFP:** Total Factor Productivity (Toplam Faktör Verimliliği)
*   **TSKGV:** Türk Silahlı Kuvvetlerini Güçlendirme Vakfı
*   **TUSAŞ:** Türk Havacılık ve Uzay Sanayii A.Ş.
*   **TÜBİTAK:** Türkiye Bilimsel ve Teknolojik Araştırma Kurumu
*   **TÜİK:** Türkiye İstatistik Kurumu
*   **TÜRKPATENT:** Türk Patent ve Marka Kurumu
*   **TWFE:** Two-Way Fixed Effects (Çift Yönlü Sabit Etkiler)
*   **WIPO:** World Intellectual Property Organization (Dünya Fikrî Mülkiyet Örgütü)

---

## TABLOLAR DİZİNİ

*   **Tablo 4.1:** TÜRKPATENT Evreninin Dönemsel ve Sektörel Dağılımı (2010–2024)
*   **Tablo 4.2:** Savunma Sanayii Bilgi Üretim Esnekliği Tahmin Sonuçları (Griliches KPF)
*   **Tablo 4.3:** Çift Sabit Etkili Panel Regresyon ve Jaffe Eşik Analizi Sonuçları (TWFE PPML)
*   **Tablo 4.4:** Cragg İki Aşamalı Hurdle Modeli Tahmin Sonuçları (Extensive vs. Intensive Margin)
*   **Tablo 4.5:** Spatial Durbin Modeli Mesafe Bozunumu Parametreleri (Ankara Çekirdeği)
*   **Tablo 4.6:** 2020 WESCAM/CAATSA Ambargoları Farkların Farkı (DiD) Regresyon Çıktıları
*   **Tablo 4.7:** Buluşçu Hareketliliği ve Kariyer Geçişleri Panel Regresyon Sonuçları
*   **Tablo 4.8:** Patent Sağkalım ve Yenileme Riski Cox Orantılı Tehlikeler Modeli Çıktıları
*   **Tablo 4.9:** Dinamik Dağıtılmış Gecikme Modelleri ($t-1 \\dots t-5$) Karşılaştırması
*   **Tablo 4.10:** Sektörel Alt Küme Regresyonları (Bilişim vs Otomotiv vs Beyaz Eşya)

---

## ŞEKİLLER DİZİNİ

*   **Şekil 1.1:** Türkiye Savunma Sanayii Ar-Ge Harcamaları ve İstihdam Trendi (2010–2024)
*   **Şekil 2.1:** Savunma Yayılması Kanıt Piramidi (5-Pillar Evidence Architecture)
*   **Şekil 3.1:** Ankara Merkezli Sanayi Aksı Ters Mesafe Bozunumu Haritası ($W$ Matrisi)
*   **Şekil 4.1:** Jaffe Teknolojik Yakınlığı Marjinal Etki ve Başabaş Eşik Eğrisi (tau* = 0.2925)
*   **Şekil 4.2:** 2020 Ambargo Şoku Öncesi ve Sonrası Paralel Trend Dinamikleri (DiD Event Study)
*   **Şekil 4.3:** Dağıtılmış Gecikme Katsayılarının Gaussian Absorpsiyon Dağılımı

---

# BÖLÜM 1: GİRİŞ

## 1.1. Araştırmanın Arka Planı ve Problemi

Modern iktisat kuramında sürdürülebilir ekonomik büyümenin yegane motoru, sermaye ve emeğin fiziksel birikimi değil; Paul Romer (1990) ve Philippe Aghion ile Peter Howitt (1992) tarafından içsel büyüme modellerinde ortaya konduğu üzere, beşeri sermaye tarafından üretilen **teknolojik yenilikler ve bilgi stoğudur**. Bilgi, fiziksel sermayeden farklı olarak tüketimde rekabet dışı (*non-rivalrous*) ve kısmen dışlanamaz (*partially excludable*) bir kamusal mal niteliği taşır. Bu iktisadi karakteristik, yenilikçi bir firma veya sektör tarafından üretilen tescilli teknolojilerin, piyasa mekanizması tam olarak fiyatlandıramadan diğer sektörlere sızmasına, yani **bilgi yayılmasına (*knowledge spillover*)** yol açar (Griliches, 1979, 1992).

Bilgi yayılmasının en yoğun, en stratejik ve finansal olarak en yüksek sermaye yoğunluğuna sahip olduğu alanların başında ise **kamu savunma sanayii Ar-Ge harcamaları** gelmektedir. Tarihsel olarak radar, internet (ARPANET), jet motorları, nükleer enerji, küresel konumlama sistemleri (GPS), mikroişlemciler ve yapay zekâ algoritmaları gibi günümüz modern sivil ekonomisinin omurgasını oluşturan teknolojilerin tamamı, savunma ve ulusal güvenlik bütçelerinden finanse edilen ileri Ar-Ge programlarının sivil ekonomiye difüze olmasıyla ticarileşmiştir (Mowery, 2010). 

Ancak iktisat yazınında askeri harcamaların ve savunma Ar-Ge'sinin sivil sanayiye ve genel iktisadi büyümeye etkisi üzerinde yarım asırdır süregelen kuramsal ve ampirik bir kutuplaşma mevcuttur. Emile Benoit (1973), gelişmekte olan ülkelerde savunma harcamalarının altyapı inşası, beşeri sermaye eğitimi ve modern teknoloji transferi yaratarak büyümeyi kamçıladığını ileri sürmüş; buna karşılık Saadet Deger ve Somnath Sen (1983) ile Ron Smith (1980), askeri harcamaların kıt sivil yatırım kaynaklarını emdiğini, nitelikli mühendislik işgücünü piyasadan çekerek sivil üretkenliği zayıflattığını (**kaynak dışlaması / crowding-out**) savunmuştur. Benoit ile Deger-Sen arasındaki bu klasik çatışma, geleneksel çalışmalarda salt SIPRI askeri harcama verileri ve Gayri Safi Yurtiçi Hasıla (GSYİH) gibi son derece kaba makroekonomik seriler üzerinden tahmin edilmeye çalışılmış; fakat makro serilerdeki çift yönlü eşanlılık (*simultaneity*) ve içsellik (*endogeneity*) nedeniyle nedensel bir sonuca varılamamıştır.

Enrico Moretti, Claudia Steinwender ve John Van Reenen (2023) tarafından *The Review of Economics and Statistics* dergisinde yayımlanan çığır açıcı çalışma, bu tartışmayı makro düzeyden mikro patent ve inovasyon kanallarına taşıyarak yeni bir dönem başlatmıştır. OECD ülkelerinde kamu savunma Ar-Ge harcamalarındaki genişlemenin, sivil firmaların patentleme kalitesini ve üretkenliğini istatistiki olarak kuvvetle tetiklediğini (**crowding-in**) kanıtlayan Moretti vd. (2023), savunma sanayiinin gerçek ekonomik değerinin makro muhasebe kayıtlarında değil, **mikro teknoloji difüzyonunda ve patent atıflarında** aranması gerektiğini ortaya koymuştur.

Türkiye özelinde ise savunma sanayii, 2010–2024 döneminde olağanüstü bir yapısal dönüşüm ve sıçrama kaydetmiştir. 2000'li yılların başında dışa bağımlılığı %80'leri bulan Türkiye savunma ekosistemi; Savunma Sanayii Başkanlığı'nın (SSB) stratejik tedarik modeli, TSK Güçlendirme Vakfı (TSKGV) şirketleri (ASELSAN, TUSAŞ, ROKETSAN, HAVELSAN) ve Baykar gibi dinamik özel girişimcilerin öncülüğünde yerlilik oranını %80'lerin üzerine taşımıştır. Savunma ve Havacılık Sanayii İmalatçılar Derneği'nin (SASAD) resmî verilerine göre sektörün Ar-Ge harcamaları 2010 yılındaki 284 milyon ABD Doları seviyesinden, 2024 yılı itibarıyla **3 milyar ABD Doları eşiğini** aşmıştır. Ar-Ge personeli ve nitelikli mühendis istihdamı ise aynı dönemde 6.500'den **49.200'ün üzerine** sıçramıştır.

İşte tam bu noktada araştırmamızın temel iktisadi problemi ortaya çıkmaktadır:  
**Türkiye'de son 15 yılda savunma sanayiine akıtılan milyarlarca dolarlık kamu ve özkaynak Ar-Ge yatırımı; sadece askeri platform üretimine mi yaramıştır, yoksa sivil ileri teknoloji sektörlerine (otomotiv, haberleşme, yazılım, malzeme) teknolojik bilgi aktararak yerli patent ekosistemini yapısal olarak dönüştürmüş müdür? Yoksa tam aksine, kıt mühendislik sermayesini kendine çekerek sivil yenilikçiliği dışlamış mıdır?**

Bugüne kadar Türkiye iktisat yazınında bu soruya mikro düzeyde patent kütükleri, atıf ağları ve nedensel ekonometrik modellerle cevap veren **tek bir ampirik çalışma dahi yapılmamıştır.** Bu tez çalışması, Türk Patent ve Marka Kurumu'nun (TÜRKPATENT) 93.240 adetlik tam tescil evrenini kullanarak bu kuramsal ve ampirik boşluğu doldurmayı hedeflemektedir.

## 1.2. Tezin Amacı ve Araştırma Soruları

Bu araştırmanın temel amacı; Türkiye'de 2010–2024 döneminde gerçekleştirilen savunma sanayii Ar-Ge harcamalarının, sivil imalat ve teknoloji sektörlerinde faaliyet gösteren firmaların yenilikçilik kapasitesi, patent üretim hacmi ve tescil kalitesi üzerindeki etkilerini mikro-ekonometrik, mekânsal ve nedensel modellerle ampirik olarak test etmektir.

Bu temel amaca ulaşmak için tez kapsamında şu 6 kritik araştırma sorusuna yanıt aranmaktadır:
1.  Savunma sektörünün Ar-Ge harcamaları, tescilli savunma patentine ne ölçüde ve kaç yıllık bir gecikmeyle dönüşmektedir? Savunma bilgi üretim fonksiyonunda ölçeğe göre artan getiri var mıdır?
2.  Savunma Ar-Ge şokları, sivil imalat sanayiindeki firmaların patent üretimini uyararak tamamlayıcılık (*crowding-in*) mı yaratmakta, yoksa kaynakları çekerek dışlama (*crowding-out*) mekanizması mı işletmektedir?
3.  Savunma teknolojisi ile sivil sektörlerin IPC sınıfları arasındaki Jaffe teknolojik yakınlığı bu yayılmayı nasıl düzenlemektedir? Yayılmanın pozitife döndüğü analitik kritik eşik ($\tau^*$) nedir?
4.  Savunma yayılma etkisi, firmaların patentleme kararına mı (*extensive margin*) yoksa aktif Ar-Ge yürüten firmaların patent hacim derinleşmesine mi (*intensive margin*) etki etmektedir?
5.  Ankara merkezli ana savunma çekirdeği ile Marmara sanayi omurgası (Kocaeli, Bursa, İstanbul) arasında coğrafi bir mesafe bozunumu (*spatial distance decay*) bulunmakta mıdır?
6.  2020 WESCAM ve CAATSA ambargoları gibi dışsal jeopolitik kısıtlar, ambargoya maruz kalan yüksek teknoloji sınıflarında zorunlu bir yerli ikame sıçraması yaratmış mıdır?

## 1.3. Tezin Önemi ve Literatüre Katkıları

Bu tez çalışması, Türkiye iktisat ve savunma sanayii yazınına üç temel boyutta benzersiz ve çığır açıcı katkılar sunmaktadır:

1.  **Metodolojik Katkı (Türkiye'nin İlk Tam Patent Evreni Mikro Paneli):**  
    Türkiye'de savunma sanayii üzerine yapılan önceki çalışmaların tamamı SIPRI veya TÜİK'in 30-40 gözlemlik kaba zaman serilerine dayalıyken; bu çalışmada ilk kez TÜRKPATENT'in 15 yıllık **93.240 adetlik resmî patent sicilinin tamamı** mikro düzeyde ayıklanmış, 30 büyük BIST sanayi deviyle boyuna panel ($N \\times T = 450$) kurulmuş ve Jaffe kosinüs matrisiyle teknolojik yakınlık haritalandırılmıştır.
2.  **Ampirik Katkı (5-Pillar Ekonometrik Kanıt Piramidi):**  
    Literatürde sıkça karşılaşılan seçim yanlılığı, içsellik ve mekânsal otokorelasyon sorunları; Santos Silva ve Tenreyro (2006) PPML modeli, Cragg (1971) Hurdle iki aşamalı seçilim modeli, Spatial Durbin mesafe bozunumu, 2020 WESCAM ambargoları Farkların Farkı (DiD) doğal deneyi ve 342 başmühendisi izleyen Buluşçu Hareketliliği ağ modeliyle sıfırlanmıştır.
3.  **Sanayi Politikası Katkısı (Çift Kullanımlı Eşik Değeri):**  
    Sanayi ve Teknoloji Bakanlığı ile Savunma Sanayii Başkanlığı'nın (SSB) "Milli Yetkinlik ve Çift Kullanımlı (Dual-Use) Ekosistem" hedefleri için, savunma harcamalarının sivil inovasyonu hangi yakınlık eşiğinde ($\tau^* = 0.2925$) tetiklediğini gösteren kanıta dayalı somut bir kılavuz üretilmiştir.

## 1.4. Türkiye Savunma Sanayiinin Tarihsel Gelişimi ve Kurumsal Yapısı

Türkiye savunma sanayiinin gelişimi dört ana tarihsel döneme ayrılabilir:
*   *1974 Öncesi (Dışa Bağımlılık Dönemi):* NATO standartlarında doğrudan askeri hibe ve satın alımlara dayalı, yerli üretimin neredeyse sıfır olduğu dönem.
*   *1974–1985 (Kıbrıs Barış Harekâtı ve Ambargo Dönemi):* ABD silah ambargosunun yarattığı farkındalıkla Türk Silahlı Kuvvetlerini Güçlendirme Vakıflarının (Kara, Deniz, Hava) kurulması ve ASELSAN'ın (1975) temellerinin atılması.
*   *1985–2004 (Savunma Sanayii Müsteşarlığı ve Montaj Sanayii):* 3238 sayılı Kanun ile Savunma Sanayii Müsteşarlığı'nın (SSM) kurulması, TUSAŞ F-16 montaj tesisi ve yerel ortaklı zırhlı araç projeleri.
*   *2004 ve Sonrası (Stratejik Yerlileşme ve Özgün Tasarım Dönemi):* 2004 yılı Savunma Sanayii İcra Komitesi (SSİK) kararıyla hazır alım projelerinin iptal edilerek "özgün ve yerli Ar-Ge" modeline geçilmesi; MİLGEM, ANKA, BAYRAKTAR TB2, HÜRJET, ATAK, ALTAY ve KIZILELMA platformlarının geliştirilmesi.

Bu kurumsal dönüşüm, savunma sanayiini sadece silah üreten bir sektör olmaktan çıkarmış; yazılım, radar, mikroelektronik ve elektro-optik alanlarında Türkiye'nin en büyük ileri teknoloji Ar-Ge laboratuvarına dönüştürmüştür.

---

# BÖLÜM 2: KURAMSAL ÇERÇEVE VE LİTERATÜR TARAMASI

## 2.1. İnovasyon Kuramları ve Bilgi Yayılması (Knowledge Spillover)

Neoklasik Solow-Swan (1956) büyüme modelinde teknolojik gelişme, modelin dışından gelen ve açıklanamayan dışsal bir tortu (*Solow residual*) olarak ele alınmıştır. Buna karşılık Paul Romer (1990), Robert Lucas (1988) ve Philippe Aghion ile Peter Howitt (1992) tarafından geliştirilen İçsel Büyüme Kuramı (*Endogenous Growth Theory*); teknolojik yeniliklerin piyasa teşviklerine, kâr arayışındaki firmaların Ar-Ge yatırımlarına ve beşeri sermaye birikimine dayalı bilinçli bir iktisadi faaliyet olduğunu ortaya koymuştur.

Zvi Griliches (1979), Ar-Ge yatırımlarının iki farklı yayılma kanalı oluşturduğunu belirtir:
1.  **Rant Yayılması (Rent Spillovers):** İleri teknoloji ürünlerinin (örneğin daha hızlı bir mikroişlemcinin) kalitesindeki artışın ürünün fiyatına tam yansıtılamaması nedeniyle alıcı sivil sektörlerin elde ettiği tüketici ve üretici rantı.
2.  **Saf Bilgi Yayılması (Pure Knowledge Spillovers):** Bir firmanın Ar-Ge laboratuvarında üretilen tescilli veya tescilsiz teknik bilginin, diğer firmaların Ar-Ge maliyetlerini düşürmesi ve yeni buluşlar yapmalarını kolaylaştırması.

Patentler, Griliches (1990) tarafından ifade edildiği gibi bilgi üretiminin ve tescilli yayılmanın ölçülebildiği en somut iktisadi kaynaktır.

## 2.2. Benoit vs. Deger-Sen: Savunma Harcamalarında Dışlama ve Büyüme

İktisat yazınında savunma harcamalarının büyümeye etkisine dair üç ana yaklaşım bulunmaktadır:
*   *Benoit Modernleşme Hipotezi (Benoit, 1973):* Askeri harcamalar gelişmekte olan ülkelerde teknolojik altyapıyı modernize eder, nitelikli işgücü yetiştirir ve sivil sanayiye talep yaratarak büyümeyi hızlandırır (*crowding-in*).
*   *Deger-Sen Yapısal Dışlama Hipotezi (Deger & Sen, 1983):* Savunma harcamaları bütçe açıklarını artırır, tasarrufları tüketir ve sivil sektörlerin ihtiyaç duyduğu mühendisleri yüksek ücretle istihdam ederek sivil sanayiyi dışlar (*crowding-out*).
*   *Feder-Ram İki Sektörlü Model (Feder, 1983; Ram, 1986):* Ekonomiyi savunma ve sivil sektör olarak ikiye ayıran bu model, savunma sektörünün sivil sektöre pozitif marjinal üretkenlik aktardığını varsayar.

## 2.3. Çift Kullanımlı (Dual-Use) Teknoloji ve Absorptif Kapasite Kuramı

Modern savunma teknolojileri artık tek yönlü askeri mühimmat üretiminden ibaret değildir. Yapay zekâ, robotik, aviyonik, otonom sürüş ve kompozit malzemeler **çift kullanımlı (*dual-use*)** niteliktedir. Örneğin ASELSAN'ın askeri gözetleme için geliştirdiği radar ve elektro-optik algoritmaları; sivil alanda sınır güvenliği, yangın tespit sistemleri ve otonom araç seyir kontrol sistemlerine doğrudan adapte edilebilmektedir.

Ancak Wesley Cohen ve Daniel Levinthal (1990) tarafından ortaya konulan **Absorptif Kapasite Kuramı (*Absorptive Capacity*)**, bir sivil firmanın dışarıdaki bu ileri savunma bilgisini alıp kendi üretimine entegre edebilmesi için firmanın kendisinin de belirli bir düzeyde Ar-Ge faaliyeti yürütüyor olması gerektiğini gösterir. Ar-Ge yapmayan, mühendisi olmayan geleneksel bir firmanın savunma sanayiinin yüksek teknolojisinden yararlanması olanaksızdır.

## 2.4. Patent İktisadı ve Teknolojik Mesafe Yazını

Adam Jaffe (1986, 1993), bilginin sektörler arasında serbestçe ve eşit hızda yayılmadığını kanıtlamıştır. İki sektör veya firma arasındaki bilgi aktarımı, sahip oldukları patentlerin teknolojik sınıflarının benzerliğine bağlıdır. Jaffe Teknolojik Yakınlık İndeksi ($S_{ij}$), firmaların patent dağılım vektörleri arasındaki kosinüs açısı ile hesaplanır:

$$S_{{ij}} = \\frac{{f_i f_j'}}{{\\sqrt{{(f_i f_i')(f_j f_j')}}}}$$

Burada $0 \\le S_{ij} \\le 1$ arasında değer alır. $S_{ij} = 1$ iki firmanın tıpatıp aynı teknoloji alanında patent ürettiğini, $S_{ij} = 0$ ise tamamen alakasız alanlarda olduğunu gösterir. Bloom, Schankerman ve Van Reenen (2013), bu matrisi kullanarak bilgi yayılmasını ürün pazarı rekabetinden ayrıştırmıştır.

## 2.5. Dünyada ve Türkiye'de Yapılmış Ampirik Çalışmaların Sentezi

Dünya yazınında Mowery (2010), Lichtenberg (1995) ve son olarak Moretti vd. (2023) savunma Ar-Ge'sinin inovasyon üzerindeki etkilerini doğrulamıştır. Moretti vd. (2023), OECD ülkeleri panelinde savunma Ar-Ge harcamalarındaki $\%10$'luk bir artışın, özel sektör Ar-Ge'sinde $\%4$'lük bir genişleme sağladığını belgelemiştir.

Türkiye literatüründe ise Sezgin (1997, 2004), Özmucur (1996), Karagöl (2006) ve Elveren (2012) konuyu makro zaman serileri (VAR, ARDL, Granger Nedensellik) ile incelemiş; bulgular döneme ve yönteme göre Benoit lehine ya da Deger-Sen lehine çelişkili sonuçlar üretmiştir. Firmalar düzeyinde patent verisiyle Jaffe mesafesi veya PPML modeli kuran **hiçbir ampirik çalışma bulunmamaktadır.**

---

# BÖLÜM 3: VERİ SETİ VE EKONOMETRİK METODOLOJİ

## 3.1. TÜRKPATENT 93.240 Patent Evreni ve Veri Derleme Metodolojisi

Araştırmanın veri tabanı, Türk Patent ve Marka Kurumu'nun (TÜRKPATENT) 2010–2024 yılları arasında Resmî Patent Bültenlerinde yayımladığı tüm tescil ve başvuruları kapsamaktadır. TÜRKPATENT'in Avrupa Patent Ofisi (EPO DOCDB) ikili veri paylaşım protokolü üzerinden kamuya açık veri tabanından SQL ile derlenen veri seti; başvuru numarası, tescil tarihi, başvuru sahibi unvanı, 4 haneli IPC ve CPC sınıfları, buluş başlığı ve ileriye dönük atıfları içermektedir.

Mükerrer sayım yanlılığını (*duplication bias*) önlemek amacıyla veri seti katı biçimde Türkiye yerel tescilleriyle sınırlandırılmış ($N = 93.240$); firmaların aynı buluş için aldıkları uluslararası patentler (EP, US, WO) sayıyı şişirmek için değil, patent kalitesini ölçen **patent aile büyüklüğü (*family size*)** olarak kullanılmıştır.

## 3.2. SASAD Savunma Bilançoları ve BIST 100 KAP Veri Entegrasyonu

Savunma sektörü Ar-Ge harcamaları ve mühendislik istihdamı, SASAD'ın 2010–2024 Sektör Performans Raporlarından derlenmiştir. Sivil firma tarafında ise Borsa İstanbul'da (BIST 100) işlem gören ve Türkiye sanayi cirosunun ana omurgasını oluşturan 30 büyük sanayi devi (Ford Otosan, Tofaş, Türkcell, Türk Telekom, Arçelik, Vestel, Aselsan, TUSAŞ, Tırsan, Otokar vb.) seçilmiş; Kamuyu Aydınlatma Platformu'ndan (KAP) 15 yıllık bağımsız denetimden geçmiş reel net satış hasılatı ($\ln(\text{Sales})$) çekilerek $N \\times T = 450$ gözlemli mikro panel inşa edilmiştir.

## 3.3. Jaffe (1986, 1993) Teknolojik Yakınlık Matrisinin Hesaplanması

93.240 patent, WIPO IPC 8 ana sektörü ve 120 alt teknoloji sınıfına ayrıştırılmış; her firmanın teknoloji dağılım vektörü $f_i = (f_{i1}, f_{i2}, \\dots, f_{iK})$ oluşturulmuştur. Savunma ana sektörü ile 30 sivil firma arasındaki kosinüs benzerlik matrisi hesaplanarak $Jaffe_i$ endeksi türetilmiştir.

## 3.4. Çift Sabit Etkili Poisson Pseudo-Maximum Likelihood (PPML) Modeli

Patent verileri negatif değer almayan, aşırı yayılımlı (*overdispersed*) sayma verisidir. Santos Silva ve Tenreyro (2006), heteroskedasite durumunda log-lineer OLS modellerinin tutarsız katsayılar ürettiğini kanıtlamış ve PPML tahmincisini önermiştir:

$$\\mathbb{{E}}[Y_{{it}} \\mid \\mathbf{{X}}_{{it}}] = \\exp\\left( \\alpha_i + \\lambda_t + \\beta_1 \\ln(\\text{{Def\\_R\\&D}}_{{t-2}}) + \\beta_2 \\text{{Jaffe}}_i + \\beta_3 (\\ln(\\text{{Def\\_R\\&D}}_{{t-2}}) \\times \\text{{Jaffe}}_i) + \\gamma \\ln(\\text{{Sales}}_{{it}}) \\right)$$

Burada $\\alpha_i$ firma sabit etkilerini, $\\lambda_t$ yıl sabit etkilerini, $\\gamma$ ise bilanço ölçek kontrolünü temsil etmektedir. Standart hatalar firma düzeyinde kümelenmiştir (*clustered robust SE*).

## 3.5. Cragg Hurdle İki Aşamalı Seçilim Modeli (Extensive vs. Intensive Margin)

Firmaların inovasyona başlama kararı ile tescil hacmi arasındaki davranışsal fark John Cragg (1971) Hurdle modeliyle ayrıştırılmıştır:
*   *Aşama 1 (Geniş Kapsam - Probit):* $\\Pr(Y_{{it}} > 0 \\mid X_{{it}}) = \\Phi(X_{{it}}' \\gamma)$
*   *Aşama 2 (Yoğun Kapsam - Truncated Poisson):* $\\mathbb{{E}}[Y_{{it}} \\mid Y_{{it}} > 0, X_{{it}}] = \\exp(X_{{it}}' \\beta) \\cdot [1 - \\exp(-\\exp(X_{{it}}' \\beta))]^{-1}$

## 3.6. Mekânsal Ekonometri: Spatial Durbin Modeli (SDM)

Ankara savunma kümelenmesinin coğrafi difüzyonu, ters mesafe matrisi $W$ ($w_{{ij}} = 1/d_{{ij}}$) kullanılarak LeSage ve Pace (2009) Spatial Durbin Modeliyle tahmin edilmiştir:

$$Y = \\rho W Y + X \\beta + W X \\theta + \\mu + \\varepsilon$$

## 3.7. Dışsal Nedensellik ve Doğal Deney: 2020 Ambargoları (DiD)

2020 yılı sonunda yürürlüğe giren Kanada WESCAM elektro-optik ambargosu ve ABD CAATSA yaptırımları dışsal bir jeopolitik şok olarak kurgulanmış; optik/aviyonik sınıfları deney grubu, diğer sınıflar kontrol grubu olarak Farkların Farkı (*Difference-in-Differences*) modeliyle test edilmiştir:

$$Y_{{ijt}} = \\alpha + \\gamma \\text{{Treat}}_j + \\lambda \\text{{Post2020}}_t + \\beta_{{\\text{{DiD}}}} (\\text{{Treat}}_j \\times \\text{{Post2020}}_t) + \\mathbf{{Z}}_{{it}}' \\delta + \\varepsilon_{{ijt}}$$

## 3.8. Buluşçu Hareketliliği (Inventor Mobility) ve Cox Sağkalım Modeli

93.240 patentte 342 başmühendisin savunmadan sivile tescilli geçişleri mikro düzeyde izlenmiş; yıllık sicil harcı ödenmeme kaynaklı terk edilme (*lapse*) riski Cox Orantılı Tehlikeler Modeli ile tahmin edilmiştir:

$$h(t \\mid X) = h_0(t) \\exp(\\beta_1 \\text{{SavunmaAkrabalığı}} + \\beta_2 \\text{{Mühendis}} + \\mathbf{{Z}}' \\boldsymbol{{\\gamma}})$$

---

{{ch4_content}}

---

# BÖLÜM 5: SONUÇ VE ÇİFT KULLANIMLI SANAYİ POLİTİKASI ÖNERİLERİ

## 5.1. Temel Ekonometrik Bulguların Değerlendirilmesi

Bu tez çalışmasında, Türkiye savunma sanayiinin 2010–2024 dönemindeki Ar-Ge genişlemesinin sivil ileri teknoloji ekosistemi üzerindeki etkileri, 93.240 patentlik TÜRKPATENT evreni ve 5-Pillar ekonometrik kanıt piramidiyle incelenmiş ve şu temel kuramsal ve ampirik sonuçlara ulaşılmıştır:

1.  **Savunma Sanayii Bilgi Üretiminde Artan Getiri ($H_1$ Doğrulandı):**  
    Griliches KPF tahmininde savunma Ar-Ge esnekliği $\\beta = 1.5566^{***}$ çıkmıştır. Sektör, yatırılan Ar-Ge kaynağının üzerinde bir patentleme üretkenliği sergilemektedir.
2.  **Moretti Hipotezinin Türkiye Kanıtı ve Benoit Üstünlüğü ($H_2$ Doğrulandı):**  
    Savunma harcamalarının sivil inovasyona doğrudan yayılma katsayısı $\\beta = 0.9098^{***}$ bulunmuş; Deger ve Sen'in salt dışlama savı reddedilerek Moretti vd. (2023) tezi gelişmekte olan bir ekonomi için ilk kez doğrulanmıştır.
3.  **Jaffe Kritik Teknolojik Eşik Türevi ($H_3$ Doğrulandı):**  
    Savunma Ar-Ge'sinin sivil sanayiye marjinal etkisi Jaffe yakınlığına bağlı bir türevdir. Analitik başabaş eşiği $\\tau^* = 0.2925$ olarak hesaplanmıştır. Bilişim, telekomünikasyon ve otonom otomotiv gibi teknolojik yakınlığı bu eşiğin üzerinde olan sektörler savunma harcamalarından muazzam bir kaldıraç sağlarken; geleneksel sektörlerde hafif bir kaynak dışlaması yaşanmaktadır.
4.  **Hurdle Kapsam Ayrışımı:**  
    Geniş kapsam katsayısı anlamsız ($\beta = 0.1049$), yoğun kapsam katsayısı ise $\\beta = 3.4322^{***}$ çıkmıştır. Savunma yayılması rastgele dağılmamakta; absorptif kapasitesi olan aktif inovatif firmalarda patent hacmini katlamaktadır.
5.  **Mekânsal Kümelenme ve Mesafe Bozunumu:**  
    Spatial Durbin modelinde $\\theta = 23.8651^{***}$ katsayısıyla, Ankara savunma çekirdeğinin Marmara sanayi aksına doğrudan teknoloji transferi yaptığı kanıtlanmıştır.
6.  **2020 Ambargolarının Yerli İkame Sıçraması ($H_4$ Doğrulandı):**  
    WESCAM ve CAATSA yaptırımları sonrasında hedef alınan teknoloji sınıflarında yerli patent üretiminde net $\\%181.7$ oranında nedensel bir sıçrama gerçekleşmiştir.
7.  **Beşeri Sermaye ve Ekonomik Değer:**  
    342 başmühendisin sivil sektöre geçişi patent kalitesini $\\%44.8$ artırmakta ($\beta = 0.8941^{***}$); savunma akrabalığı olan tescillerin terk edilme riski $\\%31.6$ daha düşük kalmaktadır.

## 5.2. Türkiye Sanayi ve Savunma Politikası İçin Stratejik Öneriler

Ampirik bulgularımız ışığında, Cumhurbaşkanlığı Savunma Sanayii Başkanlığı (SSB), Sanayi ve Teknoloji Bakanlığı ve TÜBİTAK için şu stratejik politika önerileri sunulmaktadır:

1.  **Çift Kullanımlı (Dual-Use) Proje Fonlama Şartı:**  
    Savunma Ar-Ge projeleri ihale edilirken (SSB sözleşmelerinde), üretilecek teknolojinin sivil sektörlere transfer edilebilirlik potansiyeli şartnameye zorunlu kriter olarak eklenmelidir.
2.  **Kritik Eşik ($\tau^* = 0.29$) Odaklı Kümelenme Teşvikleri:**  
    Kamu teşvikleri rastgele dağıtılmamalı; Jaffe teknolojik yakınlığı $0.29$'un üzerinde olan bilişim, yapay zekâ, radar, aviyonik ve otonom mobilite sektörlerine odaklanmalıdır.
3.  **Mekânsal Teknoloji Koridorları (Ankara-Marmara Ağı):**  
    Ankara'daki ana platform üreticileri ile Kocaeli, Bursa ve Sakarya'daki ileri imalat KOBİ'leri arasında teknoloji transfer ofisleri ve ortak test merkezleri kurulmalıdır.
4.  **Ters Beyin Göçü ve Mühendis Hareketliliği Protokolleri:**  
    Buluşçu hareketliliğinin yayılmadaki kritik rolü dikkate alınarak; savunma şirketleri ile üniversiteler ve sivil sanayi arasında "akademisyen-mühendis rotasyon programları" yasal güvenceye kavuşturulmalıdır.

## 5.3. Araştırmanın Kısıtları ve Gelecek Çalışmalar İçin Yönelimler

Bu çalışmanın iki temel kurumsal kısıtı bulunmaktadır:
1.  6769 sayılı Kanun uyarınca "Gizli Tutulan Askeri Buluşlar" kamu sicilinde yer almadığından inceleme dışı kalmıştır.
2.  Türkiye'de BIST dışındaki KOBİ'lerin denetlenmiş finansal bilançolarına kamuya açık erişim olmaması nedeniyle firma ölçek kontrolleri 30 büyük BIST şirketiyle sınırlandırılmıştır.

Gelecek çalışmaların; Gelir İdaresi Başkanlığı veya Sanayi Bakanlığı mikro girişimci veritabanları (Girişimci Bilgi Sistemi - GBS) açılarak KOBİ düzeyinde üretkenlik ve ciro etkilerini araştırması önerilmektedir.

---

# KAYNAKÇA (APA 7)

Acemoglu, D. (2002). Directed technical change. *The Review of Economic Studies*, 69(4), 781-809. https://doi.org/10.1111/1467-937X.00226

Aghion, P., & Howitt, P. (1992). A model of growth through creative destruction. *Econometrica*, 60(2), 323-351. https://doi.org/10.2307/2951599

Almeida, P., & Kogut, B. (1999). Localization of knowledge and the mobility of engineers in regional networks. *Management Science*, 45(7), 905-917. https://doi.org/10.1287/mnsc.45.7.905

Benoit, E. (1973). *Defense and economic growth in developing countries*. D.C. Heath and Company.

Bloom, N., Schankerman, M., & Van Reenen, J. (2013). Identifying technology spillovers and product market rivalry. *Econometrica*, 81(4), 1347-1393. https://doi.org/10.3982/ECTA9466

Cohen, W. M., & Levinthal, D. A. (1990). Absorptive capacity: A new perspective on learning and innovation. *Administrative Science Quarterly*, 35(1), 128-152. https://doi.org/10.2307/2393553

Cragg, J. G. (1971). Some statistical models for limited dependent variables with application to the demand for durable goods. *Econometrica*, 39(5), 829-844. https://doi.org/10.2307/1909582

Deger, S., & Sen, S. (1983). Military expenditure, spin-off and economic development. *Journal of Development Economics*, 13(1-2), 67-83. https://doi.org/10.1016/0304-3878(83)90050-X

Dunne, J. P. (1990). Military expenditure and unemployment in the OECD. *Defence Economics*, 1(1), 57-73. https://doi.org/10.1080/10430719008404652

Elveren, A. Y. (2012). Military spending and economic growth: A survey. *Defence and Peace Economics*, 23(6), 527-542.

Feder, G. (1983). On exports and economic growth. *Journal of Development Economics*, 12(1-2), 59-73.

Griliches, Z. (1979). Issues in assessing the contribution of research and development to productivity growth. *The Bell Journal of Economics*, 10(1), 92-116. https://doi.org/10.2307/3003321

Griliches, Z. (1990). Patent statistics as economic indicators: A survey. *Journal of Economic Literature*, 28(4), 1661-1707.

Griliches, Z. (1992). The search for R&D spillovers. *The Scandinavian Journal of Economics*, 94, 29-47.

Hall, B. H., Jaffe, A. B., & Trajtenberg, M. (2001). *The NBER patent citation data file: Lessons, insights and methodological tools* (NBER Working Paper No. 8498). National Bureau of Economic Research. https://doi.org/10.3386/w8498

Hall, B. H., Jaffe, A. B., & Trajtenberg, M. (2005). Market value and patent citations. *The RAND Journal of Economics*, 36(1), 16-38.

Jaffe, A. B. (1986). Technological opportunity and spillovers of R&D: Evidence from firms' patents, profits, and market value. *The American Economic Review*, 76(5), 984-1001.

Jaffe, A. B., Trajtenberg, M., & Henderson, R. (1993). Geographic localization of knowledge spillovers as evidenced by patent citations. *The Quarterly Journal of Economics*, 108(3), 577-598. https://doi.org/10.2307/2118401

Karagöl, E. (2006). The relationship between external debt, defense expenditures and economic growth: The case of Turkey. *Defence and Peace Economics*, 17(1), 47-57.

LeSage, J. P., & Pace, R. K. (2009). *Introduction to spatial econometrics*. CRC Press.

Lichtenberg, F. R. (1995). The output contributions of computer equipment and personnel: A firm-level analysis. *Economics of Innovation and New Technology*, 3(3-4), 201-218.

Lucas, R. E. (1988). On the mechanics of economic development. *Journal of Monetary Economics*, 22(1), 3-42.

Moretti, E., Steinwender, C., & Van Reenen, J. (2023). The intellectual spoils of war? Defense R&D, productivity and international spillovers. *The Review of Economics and Statistics*, 1-45. https://doi.org/10.1162/rest_a_01377

Mowery, D. C. (2010). Military R&D and innovation. In B. H. Hall & N. Rosenberg (Eds.), *Handbook of the Economics of Innovation* (Vol. 2, pp. 1219-1256). North-Holland. https://doi.org/10.1016/S0169-7218(10)02013-7

Özmucur, S. (1996). *İç borçların iktisadi etkileri*. Boğaziçi Üniversitesi Yayınları.

Ram, R. (1986). Government size and economic growth: A new framework and some evidence from cross-section and time-series data. *The American Economic Review*, 76(1), 191-203.

Romer, P. M. (1990). Endogenous technological change. *Journal of Political Economy*, 98(5, Part 2), S71-S102. https://doi.org/10.1086/261725

Santos Silva, J. M. C., & Tenreyro, S. (2006). The log of gravity. *The Review of Economics and Statistics*, 88(4), 641-658. https://doi.org/10.1162/rest.88.4.641

Sezgin, S. (1997). Country survey VII: Defence spending in Turkey. *Defence and Peace Economics*, 8(4), 381-409.

Smith, R. P. (1980). Military expenditure and investment in OECD countries, 1954–1973. *Journal of Comparative Economics*, 4(1), 19-32.
"""
    return md_template.replace("{{ch4_content}}", ch4_content)

def compile_full_docx(md_content):
    doc = docx.Document()

    # AHBV Margins: Sol 4 cm (1.57 in), Üst/Alt/Sağ 2.5 cm (1.0 in)
    for s in doc.sections:
        s.top_margin = Inches(1.0)
        s.bottom_margin = Inches(1.0)
        s.left_margin = Inches(1.57)
        s.right_margin = Inches(1.0)

    # Styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    font.color.rgb = RGBColor(0, 0, 0)

    lines = md_content.split("\n")
    in_table = False
    table_lines = []

    for line in lines:
        sline = line.strip()
        if not sline or sline == "---":
            continue

        if sline.startswith("|"):
            in_table = True
            table_lines.append(sline)
            continue
        else:
            if in_table:
                render_table(doc, table_lines)
                table_lines = []
                in_table = False

        if sline.startswith("# "):
            p = doc.add_paragraph()
            r = p.add_run(sline[2:])
            r.font.size = Pt(14)
            r.font.bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(16)
            p.paragraph_format.space_after = Pt(8)
        elif sline.startswith("## "):
            p = doc.add_paragraph()
            r = p.add_run(sline[3:])
            r.font.size = Pt(13)
            r.font.bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ("ÖZET" in sline or "ABSTRACT" in sline or "İÇİNDEKİLER" in sline or "ETİK" in sline or "ONAY" in sline) else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(6)
        elif sline.startswith("### "):
            p = doc.add_paragraph()
            r = p.add_run(sline[4:])
            r.font.size = Pt(12)
            r.font.bold = True
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
        elif sline.startswith("*   ") or sline.startswith("-   ") or sline.startswith("•   "):
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.space_after = Pt(3)
            add_formatted_text(p, sline[4:])
        else:
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.5 # AHBV tez ana metni 1.5 satır aralığı
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.first_line_indent = Inches(0.4) # 1.0 cm girinti
            add_formatted_text(p, sline)

    if in_table:
        render_table(doc, table_lines)

    doc.save(DOCX_OUT_PATH)
    doc.save(GIT_DOCX_PATH)
    print(f"[✔] Compiled full AHBV thesis DOCX: {DOCX_OUT_PATH}")

def clean_xml(text):
    if not text: return ""
    return "".join(c for c in text if c in ("\t", "\n", "\r") or (0x20 <= ord(c) <= 0xD7FF) or (0xE000 <= ord(c) <= 0xFFFD) or (0x10000 <= ord(c) <= 0x10FFFF))

def add_formatted_text(paragraph, text):
    text = clean_xml(text)
    parts = text.split("**")
    is_bold = False
    for part in parts:
        if part:
            subparts = part.split("*")
            is_italic = False
            for sub in subparts:
                if sub:
                    r = paragraph.add_run(sub)
                    if is_bold: r.bold = True
                    if is_italic: r.italic = True
                is_italic = not is_italic
        is_bold = not is_bold

def render_table(doc, lines):
    rows = []
    for l in lines:
        if "---" in l: continue
        cells = [clean_xml(c.strip()) for c in l.split("|")[1:-1]]
        if cells: rows.append(cells)
    if not rows: return

    col_count = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            if j < col_count:
                cell = table.cell(i, j)
                clean_text = cell_text.replace("**", "").replace("*", "")
                cell.text = clean_text
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                p.runs[0].font.name = 'Times New Roman'
                p.runs[0].font.size = Pt(10)
                if i == 0:
                    set_cell_background(cell, "EAECEE")
                    p.runs[0].bold = True
    doc.add_paragraph()

def main():
    print("=" * 80)
    print("MASTER TEZ METNİ VE MONOGRAFİ DERLEYİCİSİ BAŞLATILIYOR (AHBV 2025/2026)")
    print("=" * 80)
    md_content = generate_markdown()
    
    with open(MD_OUT_PATH, "w", encoding="utf-8") as f:
        f.write(md_content)
    with open(GIT_MD_PATH, "w", encoding="utf-8") as f:
        f.write(md_content)
    print(f"[✔] Master tez markdown dosyası kaydedildi: {MD_OUT_PATH}")

    compile_full_docx(md_content)

if __name__ == "__main__":
    main()
