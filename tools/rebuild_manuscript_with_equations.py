#!/usr/bin/env python3
"""
DEFINITIVE AHBV MASTER THESIS BUILDER WITH NATIVE WORD EQUATIONS
1. Uses the official authentic AHBV template (`tez_yazim_sablonu_2025_.docx`)
2. Injects all front matters (Dış Kapak, İç Kapak, ONAY, Etik Beyan, İthaf, Teşekkür, Özet, Abstract, Kısaltmalar)
3. Implements strict AHBV heading hierarchy:
   - 1. GİRİŞ (Bölüm Başlıkları stili, alt numarasız)
   - 2. BÖLÜM BAŞLIKLARI (BÖLÜM BAŞLIKLARI stili)
   - 2.1. İKİNCİ BAŞLIK stili
   - 2.1.1. ÜÇÜNCÜ BAŞLIK stili
4. CONVERTS ALL EQUATIONS TO NATIVE WORD EQUATION BUILDER OBJECTS (OMML)
5. Properly formats Tables and Figures per AHBV 2025/2026 standards
"""

import os
import re
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from equation_converter import add_equation_to_paragraph, latex_to_omml_element

OFFICIAL_TEMPLATE_PATH = "/Users/dogukancihanbeyoglu/Desktop/AHBV_Iktisat_Tez_Calismasi/02_Resmi_AHBV_Kilavuzlar_ve_Sablonlar/tez_yazim_sablonu_2025_.docx"
DESKTOP_OUT_PATH = "/Users/dogukancihanbeyoglu/Desktop/AHBV_Iktisat_Tez_Calismasi/01_Tez_Oneri_Formu/AHBV_IKTISAT_TEZI_TAM_METIN.docx"
GIT_OUT_PATH = "/Users/dogukancihanbeyoglu/Gemini/tez_calismasi/01_Tez_Oneri_Formu/AHBV_IKTISAT_TEZI_TAM_METIN.docx"
MD_PATH = "/Users/dogukancihanbeyoglu/Desktop/AHBV_Iktisat_Tez_Calismasi/01_Tez_Oneri_Formu/AHBV_IKTISAT_TEZI_TAM_METIN.md"

def clean_xml(text):
    if not text:
        return ""
    return "".join(c for c in text if c in ("\t", "\n", "\r") or (0x20 <= ord(c) <= 0xD7FF) or (0xE000 <= ord(c) <= 0xFFFD) or (0x10000 <= ord(c) <= 0x10FFFF))

def set_cell_background(cell, fill_color):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_color}"/>')
    tcPr.append(shd)

def add_formatted_text(paragraph, text):
    text = clean_xml(text)
    parts = text.split("**")
    is_bold = False
    for part in parts:
        if part:
            subparts = part.split("*")
            is_italic = False
            for sub in subparts:
                if sub:
                    r = paragraph.add_run(sub)
                    r.font.name = "Times New Roman"
                    r.font.size = Pt(12)
                    if is_bold: r.bold = True
                    if is_italic: r.italic = True
                is_italic = not is_italic
        is_bold = not is_bold

def render_table(doc, lines):
    rows = []
    for l in lines:
        if "---" in l: continue
        cells = [clean_xml(c.strip()) for c in l.split("|")[1:-1]]
        if cells: rows.append(cells)
    if not rows: return

    col_count = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx]
        is_header = (r_idx == 0)
        for c_idx in range(min(col_count, len(row_data))):
            cell = row.cells[c_idx]
            cell.text = row_data[c_idx]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if (is_header or c_idx > 0) else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.line_spacing = 1.0 # AHBV kuralı: tablo tek satır aralığı
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)
                if is_header:
                    run.bold = True
            if is_header:
                set_cell_background(cell, "EAECEE")
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

def build_manuscript():
    print("[1/5] Loading official AHBV template...")
    doc = docx.Document(OFFICIAL_TEMPLATE_PATH)

    thesis_title_tr = "TÜRKİYE SAVUNMA SANAYİİ YAYILMA DİNAMİKLERİNİN İLERİ TEKNOLOJİ PATENT EKOSİSTEMİNE ETKİLERİ: MİKRO-EKONOMETRİK VE MEKÂNSAL BİR ANALİZ (2010–2024)"
    thesis_title_en = "SPILLOVER DYNAMICS OF THE TURKISH DEFENSE INDUSTRY ON THE ADVANCED TECHNOLOGY PATENT ECOSYSTEM: A MICRO-ECONOMETRIC AND SPATIAL ANALYSIS (2010–2024)"
    author_name = "Doğukan CİHANBEYOĞLU"
    supervisor = "Tez Danışmanı: [Unvanı, Adı SOYADI]"
    department_tr = "İKTİSAT ANA BİLİM DALI"
    program_tr = "İKTİSAT TEZLİ YÜKSEK LİSANS PROGRAMI"
    city_year = "Ankara - 2027"

    print("[2/5] Updating metadata, covers, ethics and approval fields...")
    for p in doc.paragraphs[:150]:
        t = p.text.strip()
        if "Tez Adı Baş Harfleri Büyük" in t or "Büyük harflerle ve ortalanmış olarak tez adı" in t:
            p.text = clean_xml(thesis_title_tr)
            for r in p.runs:
                r.font.name = "Times New Roman"
                r.font.size = Pt(14)
                r.bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif t == "Ad SOYAD" or t == "Adı SOYADI" or t == "(Ad ve SOYAD)":
            p.text = clean_xml(author_name)
            for r in p.runs:
                r.font.name = "Times New Roman"
                r.font.size = Pt(12)
                r.bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif "Tez Danışmanı Unvan Ad SOYAD" in t:
            p.text = clean_xml(supervisor)
            for r in p.runs:
                r.font.name = "Times New Roman"
                r.font.size = Pt(12)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif "… ANA BİLİM DALI" in t:
            p.text = clean_xml(department_tr)
            for r in p.runs:
                r.font.name = "Times New Roman"
                r.font.size = Pt(12)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif "… PROGRAMI" in t:
            p.text = clean_xml(program_tr)
            for r in p.runs:
                r.font.name = "Times New Roman"
                r.font.size = Pt(12)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif "YÜKSEK LİSANS TEZİ/DOKTORA" in t:
            p.text = "YÜKSEK LİSANS TEZİ"
            for r in p.runs:
                r.font.name = "Times New Roman"
                r.font.size = Pt(12)
                r.bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif "SAVUNMA YILI" in t:
            p.text = clean_xml(city_year)
            for r in p.runs:
                r.font.name = "Times New Roman"
                r.font.size = Pt(12)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif "(Tezin/Projenin Türkçe Başlığı)" in t:
            p.text = clean_xml(thesis_title_tr)
            for r in p.runs:
                r.font.name = "Times New Roman"
                r.font.size = Pt(12)
                r.bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif "(English Title of Thesis/Project)" in t:
            p.text = clean_xml(thesis_title_en)
            for r in p.runs:
                r.font.name = "Times New Roman"
                r.font.size = Pt(12)
                r.bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif "Name and SURNAME" in t:
            p.text = clean_xml(author_name)
            for r in p.runs:
                r.font.name = "Times New Roman"
                r.font.size = Pt(12)
                r.bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif "Supervisor: Academic Title" in t:
            p.text = "Supervisor: [Title, Name SURNAME]"
            for r in p.runs:
                r.font.name = "Times New Roman"
                r.font.size = Pt(12)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif "Department of …" in t:
            p.text = "Department of Economics, Master's Program"
            for r in p.runs:
                r.font.name = "Times New Roman"
                r.font.size = Pt(12)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif "Year, Ankara" in t:
            p.text = "January 2027, Ankara"
            for r in p.runs:
                r.font.name = "Times New Roman"
                r.font.size = Pt(12)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif "Master’s Thesis, Doctoral Thesis" in t:
            p.text = "Master's Thesis"
            for r in p.runs:
                r.font.name = "Times New Roman"
                r.font.size = Pt(12)
                r.bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif "… Ana Bilim Dalı/…Ana Sanat Dalı" in t:
            p.text = "İktisat Ana Bilim Dalı"
            for r in p.runs:
                r.font.name = "Times New Roman"
                r.font.size = Pt(12)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif "Yıl, Ankara" in t:
            p.text = "Ocak 2027, Ankara"
            for r in p.runs:
                r.font.name = "Times New Roman"
                r.font.size = Pt(12)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Update ONAY text
    for p in doc.paragraphs:
        if "tarih ve saatinde yapılan tez savunma sınavında" in p.text:
            p.text = f"İktisat Ana Bilim Dalı İktisat Tezli Yüksek Lisans Programı öğrencisi {author_name} tarafından hazırlanan '{thesis_title_tr}' başlıklı tez çalışması .../01/2027 tarihinde yapılan tez savunma sınavında aşağıdaki jüri tarafından OY BİRLİĞİ ile YÜKSEK LİSANS TEZİ olarak KABUL edilmiştir."
            for r in p.runs:
                r.font.name = "Times New Roman"
                r.font.size = Pt(11)
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Read markdown source
    with open(MD_PATH, "r", encoding="utf-8") as f:
        md_text = f.read()

    # Extract Özet text
    ozet_match = re.search(r"## ÖZET.*?\n\n(.*?)\n\n\*\*Anahtar Kelimeler:\*\*", md_text, re.DOTALL)
    if ozet_match:
        ozet_text = ozet_match.group(1).replace("\n", " ").strip()
        for p in doc.paragraphs[50:70]:
            if "……………………………………………………" in p.text:
                p.text = clean_xml(ozet_text)
                for r in p.runs:
                    r.font.name = "Times New Roman"
                    r.font.size = Pt(11)
                p.paragraph_format.line_spacing = 1.0 # AHBV kuralı: özet tek satır aralığı
                p.paragraph_format.space_after = Pt(6)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                break

    # Extract Keywords TR
    kw_tr_match = re.search(r"\*\*Anahtar Kelimeler:\*\*\s*(.*)", md_text)
    if kw_tr_match:
        for p in doc.paragraphs[55:75]:
            if "Anahtar Kelimeler:" in p.text:
                p.text = "Anahtar Kelimeler: " + clean_xml(kw_tr_match.group(1).strip())
                for r in p.runs:
                    r.font.name = "Times New Roman"
                    r.font.size = Pt(11)
                p.paragraph_format.line_spacing = 1.0
                break

    # Extract Abstract text EN
    abstract_match = re.search(r"## ABSTRACT.*?\n\n(.*?)\n\n\*\*Keywords:\*\*", md_text, re.DOTALL)
    if abstract_match:
        abstract_text = abstract_match.group(1).replace("\n", " ").strip()
        for p in doc.paragraphs[80:100]:
            if "……………………………………………………" in p.text:
                p.text = clean_xml(abstract_text)
                for r in p.runs:
                    r.font.name = "Times New Roman"
                    r.font.size = Pt(11)
                p.paragraph_format.line_spacing = 1.0 # AHBV kuralı: abstract tek satır aralığı
                p.paragraph_format.space_after = Pt(6)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                break

    # Extract Keywords EN
    kw_en_match = re.search(r"\*\*Keywords:\*\*\s*(.*)", md_text)
    if kw_en_match:
        for p in doc.paragraphs[85:105]:
            if "Keywords:" in p.text:
                p.text = "Keywords: " + clean_xml(kw_en_match.group(1).strip())
                for r in p.runs:
                    r.font.name = "Times New Roman"
                    r.font.size = Pt(11)
                p.paragraph_format.line_spacing = 1.0
                break

    # Update Teşekkür / İthaf
    for p in doc.paragraphs[95:110]:
        if "İthaf; herhangi bir şekil" in p.text:
            p.text = "Bu çalışma; Türkiye'nin bağımsız teknoloji hamlesine ömrünü adayan mühendis ve bilim insanlarına ithaf edilmiştir."
            for r in p.runs:
                r.font.name = "Times New Roman"
                r.font.size = Pt(11)
                r.italic = True
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif "Metin, 12 punto büyüklüğünde, tek satır aralığında yazılır" in p.text:
            p.text = "Bu tez çalışmasının hazırlanma sürecinde değerli bilgi, birikim ve tecrübeleriyle yolumu aydınlatan kıymetli tez danışmanıma, araştırmam sırasında kurumsal ve teknik desteklerini esirgemeyen tüm hocalarıma ve her daim yanımda olan aileme en içten teşekkürlerimi sunarım."
            for r in p.runs:
                r.font.name = "Times New Roman"
                r.font.size = Pt(11)
            p.paragraph_format.line_spacing = 1.0
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Populate Kısaltmalar tablosu (Table 1 in doc)
    abbr_data = [
        ("AHBV", "Ankara Hacı Bayram Veli Üniversitesi"),
        ("ASELSAN", "Askeri Elektronik Sanayi A.Ş."),
        ("BIST", "Borsa İstanbul"),
        ("CAATSA", "Countering America's Adversaries Through Sanctions Act"),
        ("CPC", "Cooperative Patent Classification (Ortak Patent Sınıflandırması)"),
        ("DiD", "Difference-in-Differences (Farkların Farkı Yöntemi)"),
        ("DOCDB", "EPO Worldwide Patent Database"),
        ("EPO", "European Patent Office (Avrupa Patent Ofisi)"),
        ("GSYİH", "Gayri Safi Yurtiçi Hasıla"),
        ("IPC", "International Patent Classification (Uluslararası Patent Sınıflandırması)"),
        ("KAP", "Kamuyu Aydınlatma Platformu"),
        ("KPF", "Knowledge Production Function (Bilgi Üretim Fonksiyonu)"),
        ("OLS", "Ordinary Least Squares (En Küçük Kareler Yöntemi)"),
        ("PPML", "Poisson Pseudo-Maximum Likelihood"),
        ("ROKETSAN", "Roket Sanayii ve Ticaret A.Ş."),
        ("SASAD", "Savunma ve Havacılık Sanayii İmalatçılar Derneği"),
        ("SDM", "Spatial Durbin Model (Mekânsal Durbin Modeli)"),
        ("SSB", "Savunma Sanayii Başkanlığı"),
        ("SSİK", "Savunma Sanayii İcra Komitesi"),
        ("TCMB", "Türkiye Cumhuriyet Merkez Bankası"),
        ("TFP", "Total Factor Productivity (Toplam Faktör Verimliliği)"),
        ("TSKGV", "Türk Silahlı Kuvvetlerini Güçlendirme Vakfı"),
        ("TUSAŞ", "Türk Havacılık ve Uzay Sanayii A.Ş."),
        ("TÜBİTAK", "Türkiye Bilimsel ve Teknolojik Araştırma Kurumu"),
        ("TÜİK", "Türkiye İstatistik Kurumu"),
        ("TÜRKPATENT", "Türk Patent ve Marka Kurumu"),
        ("TWFE", "Two-Way Fixed Effects (Çift Yönlü Sabit Etkiler)"),
        ("WIPO", "World Intellectual Property Organization"),
    ]
    if len(doc.tables) > 1:
        abbr_table = doc.tables[1]
        while len(abbr_table.rows) > 1:
            abbr_table._tbl.remove(abbr_table.rows[-1]._tr)
        for abbr, full in abbr_data:
            row = abbr_table.add_row()
            row.cells[0].text = abbr
            row.cells[1].text = full
            for c_idx, cell in enumerate(row.cells):
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.line_spacing = 1.0
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                for r in p.runs:
                    r.font.name = "Times New Roman"
                    r.font.size = Pt(10)
                    if c_idx == 0: r.bold = True

    print("[3/5] Truncating placeholder body text from official template...")
    start_body_idx = None
    for idx, p in enumerate(doc.paragraphs):
        if p.style.name == "BÖLÜM BAŞLIKLARI" and "GİRİŞ" in p.text:
            start_body_idx = idx
            break
    
    if start_body_idx is not None:
        for p in doc.paragraphs[start_body_idx:]:
            p._element.getparent().remove(p._element)

    print("[4/5] Injecting chapters and CONVERTING EQUATIONS TO NATIVE WORD EQUATIONS (OMML)...")
    
    # We define the clean, definitive equations dictionary mapped to key identifiers
    equations_map = {
        "FEDER_RAM": r"Y = C + D, \quad C = C(K_c, L_c, D), \quad D = D(K_d, L_d)",
        "JAFFE_SIM": r"S_{ij} = \frac{f_i f_j'}{\sqrt{(f_i f_i')(f_j f_j')}}",
        "PPML_MODEL": r"\mathbb{E}[Y_{it} \mid \mathbf{X}_{it}] = \exp\left( \alpha_i + \lambda_t + \beta_1 \ln(\text{Def\_R\&D}_{t-2}) + \beta_2 \text{Jaffe}_i + \beta_3 (\ln(\text{Def\_R\&D}_{t-2}) \times \text{Jaffe}_i) + \gamma \ln(\text{Sales}_{it}) \right)",
        "HURDLE_P1": r"\Pr(Y_{it} > 0 \mid X_{it}) = \Phi(X_{it}' \gamma)",
        "HURDLE_P2": r"\mathbb{E}[Y_{it} \mid Y_{it} > 0, X_{it}] = \exp(X_{it}' \beta) \cdot [1 - \exp(-\exp(X_{it}' \beta))]^{-1}",
        "SPATIAL_DURBIN": r"Y = \rho W Y + X \beta + W X \theta + \mu + \varepsilon",
        "DID_MODEL": r"Y_{ijt} = \alpha + \gamma \text{Treat}_j + \lambda \text{Post2020}_t + \beta_{\text{DiD}} (\text{Treat}_j \times \text{Post2020}_t) + \mathbf{Z}_{it}' \delta + \varepsilon_{ijt}",
        "COX_MODEL": r"h(t \mid X) = h_0(t) \exp(\beta_1 \text{SavunmaAkrabalığı} + \beta_2 \text{Mühendis} + \mathbf{Z}' \boldsymbol{\gamma})",
        "KPF_MODEL": r"\ln(\text{Defense\_Patents}_t) = \alpha + \beta_1 \ln(\text{Def\_R\&D}_{t-2}) + \beta_2 \ln(\text{Engineers}_t) + \varepsilon_t",
        "MARGINAL_DERIV": r"\frac{\partial \mathbb{E}[Y_{it}] / \mathbb{E}[Y_{it}]}{\partial \ln(\text{Def\_R\&D})} = \beta_1 + \beta_3 \cdot \text{Jaffe}_i = -1.2161 + 4.1579 \cdot \text{Jaffe}_i",
        "BREAKEVEN_TAU": r"\tau^* = \frac{-\beta_1}{\beta_3} = \frac{1.2161}{4.1579} \approx 0.2925",
        "DID_PERCENT": r"\% \Delta = (\exp(1.0358) - 1) \times 100 = +181.7\%",
        "COX_HR": r"\text{Hazard Ratio (HR)} = \exp(\hat{\beta}) = 0.684^{***} \quad (z = -2.87, \; p = 0.0041)"
    }

    body_md_idx = md_text.find("# BÖLÜM 1: GİRİŞ")
    thesis_body_md = md_text[body_md_idx:]
    lines = thesis_body_md.split("\n")

    in_table = False
    table_lines = []

    for line in lines:
        sline = line.strip()

        if sline.startswith("|") and sline.endswith("|"):
            in_table = True
            table_lines.append(sline)
            continue
        else:
            if in_table:
                render_table(doc, table_lines)
                in_table = False
                table_lines = []

        if not sline:
            continue

        # 1. DERECE BAŞLIK (BÖLÜM BAŞLIKLARI):
        if sline.startswith("# BÖLÜM 1: GİRİŞ"):
            # AHBV Kuralı: 1. GİRİŞ (alt başlık numarası almaz)
            p = doc.add_paragraph(style='BÖLÜM BAŞLIKLARI')
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(12)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run("1. GİRİŞ")
            r.font.name = "Times New Roman"
            r.font.size = Pt(12)
            r.bold = True
        elif sline.startswith("# BÖLÜM ") or sline.startswith("# KAYNAKÇA"):
            header_text = sline.replace("# ", "").strip()
            # If it's "BÖLÜM 2: ...", transform to AHBV "2. KURAMSAL ÇERÇEVE VE LİTERATÜR TARAMASI"
            m = re.match(r"BÖLÜM\s+(\d+):\s*(.*)", header_text)
            if m:
                header_text = f"{m.group(1)}. {m.group(2).upper()}"
            p = doc.add_paragraph(style='BÖLÜM BAŞLIKLARI')
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(12)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(clean_xml(header_text))
            r.font.name = "Times New Roman"
            r.font.size = Pt(12)
            r.bold = True
        # 2. DERECE BAŞLIK (İKİNCİ BAŞLIK):
        elif sline.startswith("## "):
            h_text = sline[3:].strip()
            # If in Chapter 1, AHBV allows unnumbered italic/plain subheadings
            if h_text.startswith("1."):
                # In AHBV guide: Giriş altında 1.1, 1.2 kullanılmaz, ara başlık olarak italik/düz verilir
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(14)
                p.paragraph_format.space_after = Pt(6)
                p.paragraph_format.first_line_indent = Inches(0.3937)
                r = p.add_run(clean_xml(h_text[h_text.find(' ')+1:]))
                r.font.name = "Times New Roman"
                r.font.size = Pt(12)
                r.bold = True
            else:
                p = doc.add_paragraph(style='İKİNCİ BAŞLIK')
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(6)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                r = p.add_run(clean_xml(h_text))
                r.font.name = "Times New Roman"
                r.font.size = Pt(12)
                r.bold = True
        # 3. DERECE BAŞLIK (ÜÇÜNCÜ BAŞLIK):
        elif sline.startswith("### "):
            h_text = sline[4:].strip()
            # If it is a Table caption like "### Tablo 4.1: ..."
            if h_text.startswith("Tablo ") or h_text.startswith("Şekil "):
                p = doc.add_paragraph(style='Caption')
                p.paragraph_format.space_before = Pt(10)
                p.paragraph_format.space_after = Pt(4)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                r = p.add_run(clean_xml(h_text))
                r.font.name = "Times New Roman"
                r.font.size = Pt(11)
                r.bold = True
            else:
                p = doc.add_paragraph(style='ÜÇÜNCÜ BAŞLIK')
                p.paragraph_format.space_before = Pt(10)
                p.paragraph_format.space_after = Pt(4)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                r = p.add_run(clean_xml(h_text))
                r.font.name = "Times New Roman"
                r.font.size = Pt(12)
                r.bold = True
                r.italic = True
        # EQUATIONS ($$ or known formula strings):
        elif sline.startswith("$$") or "Y = C + D" in sline or "S_{ij} =" in sline or "\\mathbb{E}[Y_{it}" in sline or "Y = \\rho W Y" in sline or "h(t \\mid X)" in sline or "\\ln(\\text{Defense" in sline or "\\frac{\\partial \\mathbb{E}" in sline or "\\tau^* =" in sline or "\\% \\Delta =" in sline:
            # Map to native Word OMML equation
            matched_latex = None
            if "Y = C + D" in sline: matched_latex = equations_map["FEDER_RAM"]
            elif "S_{ij} =" in sline or "S_{{ij}}" in sline: matched_latex = equations_map["JAFFE_SIM"]
            elif "mathbb{E}[Y_{it} \\mid \\mathbf{X}" in sline: matched_latex = equations_map["PPML_MODEL"]
            elif "Pr(Y_{it} > 0" in sline: matched_latex = equations_map["HURDLE_P1"]
            elif "Truncated Poisson" in sline: matched_latex = equations_map["HURDLE_P2"]
            elif "Spatial Durbin" in sline or "Y = \\rho" in sline or "ho W Y" in sline: matched_latex = equations_map["SPATIAL_DURBIN"]
            elif "Y_{ijt} =" in sline or "Treat" in sline: matched_latex = equations_map["DID_MODEL"]
            elif "h(t \\mid X)" in sline: matched_latex = equations_map["COX_MODEL"]
            elif "Defense\\_Patents" in sline: matched_latex = equations_map["KPF_MODEL"]
            elif "partial \\mathbb{E}" in sline: matched_latex = equations_map["MARGINAL_DERIV"]
            elif "\\tau^* =" in sline or "0.2925" in sline: matched_latex = equations_map["BREAKEVEN_TAU"]
            elif "\\Delta =" in sline or "181.7" in sline: matched_latex = equations_map["DID_PERCENT"]
            elif "Hazard Ratio" in sline: matched_latex = equations_map["COX_HR"]
            else:
                raw_eq = sline.replace("$$", "").strip()
                matched_latex = raw_eq

            if matched_latex:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(8)
                p.paragraph_format.space_after = Pt(8)
                success = add_equation_to_paragraph(p, matched_latex)
                if success:
                    print(f"[✔] Rendered native Word Equation: {matched_latex[:40]}...")
        elif sline.startswith("---"):
            continue
        elif sline.startswith("* ") or sline.startswith("- "):
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.left_indent = Inches(0.3937) # 1 cm
            add_formatted_text(p, "• " + sline[2:].strip())
        elif re.match(r"^\d+\.\s", sline):
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.left_indent = Inches(0.3937)
            add_formatted_text(p, sline)
        else:
            # Body text per AHBV: 12 pt Times New Roman, 1.5 satır aralığı, 3nk önce / 6nk sonra, 1.0 cm girinti, iki yana yaslı
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.first_line_indent = Inches(0.3937) # 1.0 cm
            add_formatted_text(p, sline)

    if in_table:
        render_table(doc, table_lines)

    print("[5/5] Saving final manuscript...")
    doc.save(DESKTOP_OUT_PATH)
    doc.save(GIT_OUT_PATH)
    print(f"[✔] Master Thesis DOCX successfully generated with Native Word Equations: {DESKTOP_OUT_PATH}")

if __name__ == "__main__":
    build_manuscript()
